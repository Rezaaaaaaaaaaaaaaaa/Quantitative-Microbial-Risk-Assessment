#!/usr/bin/env python3
"""
NIWA QMRA Workflow Engine - Improved Project Application with Citations
Based on the updated PDF proposal with proper academic citations
"""

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import FancyBboxPatch, Rectangle, FancyArrowPatch
import matplotlib.lines as mlines
from io import BytesIO

class ImprovedNIWAQMRAProposal:
    def __init__(self):
        self.doc = Document()
        self.setup_official_styles()

    def setup_official_styles(self):
        """Setup styles matching the official NIWA template."""
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Normal style
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)

    def add_niwa_header(self):
        """Add NIWA logo and header."""
        header_para = self.doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run("NIWA")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_para.space_after = Pt(12)

    def add_title_section(self):
        """Add the title section exactly matching the template."""
        title_para = self.doc.add_paragraph("Structured Internal Project Application 2024-2025")
        title_para.space_after = Pt(18)

        overview_para = self.doc.add_paragraph("Project Overview")
        overview_para.space_after = Pt(12)

    def add_project_details_table(self):
        """Add the project details table matching the exact template format."""
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(4.5)

        # Row 1: Project name
        table.cell(0, 0).text = "Project name:"
        table.cell(0, 1).text = "Development of QMRA Workflow Engine"

        # Row 2: Staff
        table.cell(1, 0).text = "Staff:"
        table.cell(1, 1).text = "Reza Moghaddam (Lead Developer - 150 hrs), David Wood (Model Review & Support - 40 hrs)"

        # Row 3: Project Manager
        table.cell(2, 0).text = "Project Manager:"
        table.cell(2, 1).text = "[To be assigned - Group Manager]"

        # Row 4: Region
        table.cell(3, 0).text = "Region:"
        table.cell(3, 1).text = "Hamilton"

        # Row 5: Centre
        table.cell(4, 0).text = "Centre:"
        table.cell(4, 1).text = "FRESHWATER"

        # Row 6: Type
        table.cell(5, 0).text = "Type:"
        table.cell(5, 1).text = "Science (Applied Research & Development)"

        # Row 7: Project objective
        table.cell(6, 0).text = "Project objective:"
        table.cell(6, 1).text = "Develop a Python-based QMRA workflow engine to reduce project delivery time by 60-70% and capture greater market share in the expanding regulatory compliance sector."

        # Make labels bold
        for i in range(7):
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    def create_qmra_workflow_diagram(self):
        """Create a clean, non-overlapping workflow engine diagram."""
        fig, ax = plt.subplots(figsize=(14, 10))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 10)
        ax.axis('off')

        # Title
        ax.text(7, 9.5, 'QMRA WORKFLOW ENGINE ARCHITECTURE',
                fontsize=16, fontweight='bold', ha='center')

        # Define colors
        input_color = '#E8F4F8'
        core_color = '#B8E2F2'
        module_color = '#7BC8E5'
        output_color = '#D4E8D4'

        # INPUT LAYER (top) - Non-overlapping
        input_boxes = [
            {'x': 1, 'y': 7.5, 'w': 2.5, 'h': 0.8, 'text': 'Water Quality\nData'},
            {'x': 4, 'y': 7.5, 'w': 2.5, 'h': 0.8, 'text': 'Treatment\nPerformance'},
            {'x': 7, 'y': 7.5, 'w': 2.5, 'h': 0.8, 'text': 'Pathogen\nConcentrations'},
            {'x': 10, 'y': 7.5, 'w': 2.5, 'h': 0.8, 'text': 'Regulatory\nRequirements'}
        ]

        for box in input_boxes:
            rect = FancyBboxPatch((box['x'], box['y']), box['w'], box['h'],
                                 boxstyle="round,pad=0.05",
                                 facecolor=input_color, edgecolor='#5A5A5A', linewidth=1.5)
            ax.add_patch(rect)
            ax.text(box['x'] + box['w']/2, box['y'] + box['h']/2, box['text'],
                   ha='center', va='center', fontsize=9)

        # CORE ENGINE (middle) - Central processing unit
        core_rect = FancyBboxPatch((2, 5), 10, 1.5,
                                  boxstyle="round,pad=0.05",
                                  facecolor=core_color, edgecolor='#2C5F7C', linewidth=2)
        ax.add_patch(core_rect)
        ax.text(7, 5.75, 'PYTHON CORE PROCESSING ENGINE',
               ha='center', va='center', fontsize=12, fontweight='bold')
        ax.text(7, 5.25, 'NumPy • SciPy • Pandas • Statistical Libraries',
               ha='center', va='center', fontsize=9, style='italic')

        # PROCESSING MODULES (middle-bottom) - Non-overlapping arrangement
        modules = [
            {'x': 0.5, 'y': 3, 'w': 3, 'h': 1.2, 'text': 'Pathogen Database\n• Dose-Response\n• Beta-Poisson\n• Exponential'},
            {'x': 4, 'y': 3, 'w': 3, 'h': 1.2, 'text': 'Treatment Module\n• Log Reduction\n• Validation\n• Performance'},
            {'x': 7.5, 'y': 3, 'w': 3, 'h': 1.2, 'text': 'Risk Simulation\n• Monte Carlo\n• Bayesian\n• Uncertainty'},
            {'x': 11, 'y': 3, 'w': 2.5, 'h': 1.2, 'text': 'QA/QC Module\n• Validation\n• Benchmarks\n• Standards'}
        ]

        for mod in modules:
            rect = FancyBboxPatch((mod['x'], mod['y']), mod['w'], mod['h'],
                                 boxstyle="round,pad=0.05",
                                 facecolor=module_color, edgecolor='#1A4F6B', linewidth=1.5)
            ax.add_patch(rect)
            lines = mod['text'].split('\n')
            ax.text(mod['x'] + mod['w']/2, mod['y'] + mod['h'] - 0.2, lines[0],
                   ha='center', va='center', fontsize=10, fontweight='bold')
            for i, line in enumerate(lines[1:], 1):
                ax.text(mod['x'] + mod['w']/2, mod['y'] + mod['h'] - 0.2 - (i*0.25), line,
                       ha='center', va='center', fontsize=8)

        # OUTPUT LAYER (bottom) - Non-overlapping
        output_boxes = [
            {'x': 1, 'y': 0.5, 'w': 3, 'h': 1, 'text': 'Risk Estimates\n• DALYs\n• Infection Probability'},
            {'x': 4.5, 'y': 0.5, 'w': 3, 'h': 1, 'text': 'Compliance Reports\n• Regulatory Docs\n• Visualizations'},
            {'x': 8, 'y': 0.5, 'w': 3, 'h': 1, 'text': 'Decision Support\n• Recommendations\n• Mitigation Options'},
            {'x': 11.5, 'y': 0.5, 'w': 2, 'h': 1, 'text': 'Export Formats\n• PDF/Word\n• Excel/CSV'}
        ]

        for box in output_boxes:
            rect = FancyBboxPatch((box['x'], box['y']), box['w'], box['h'],
                                 boxstyle="round,pad=0.05",
                                 facecolor=output_color, edgecolor='#2B5A2B', linewidth=1.5)
            ax.add_patch(rect)
            lines = box['text'].split('\n')
            ax.text(box['x'] + box['w']/2, box['y'] + box['h'] - 0.2, lines[0],
                   ha='center', va='center', fontsize=10, fontweight='bold')
            for i, line in enumerate(lines[1:], 1):
                ax.text(box['x'] + box['w']/2, box['y'] + box['h'] - 0.15 - (i*0.2), line,
                       ha='center', va='center', fontsize=8)

        # Add clean arrows for data flow (avoiding overlaps)
        arrow_props = dict(arrowstyle='->', lw=1.5, color='#4A4A4A', alpha=0.7)

        # Input to Core
        for box in input_boxes:
            ax.annotate('', xy=(7, 6.5), xytext=(box['x'] + box['w']/2, box['y']),
                       arrowprops=arrow_props)

        # Core to Modules
        module_centers = [2, 5.5, 9, 12.25]
        for center in module_centers:
            ax.annotate('', xy=(center, 4.2), xytext=(center, 5),
                       arrowprops=arrow_props)

        # Modules to Output
        output_centers = [2.5, 6, 9.5, 12.5]
        for i, center in enumerate(module_centers):
            ax.annotate('', xy=(output_centers[i], 1.5), xytext=(center, 3),
                       arrowprops=arrow_props)

        # Add legend
        legend_elements = [
            patches.Patch(facecolor=input_color, edgecolor='black', label='Input Data'),
            patches.Patch(facecolor=core_color, edgecolor='black', label='Core Engine'),
            patches.Patch(facecolor=module_color, edgecolor='black', label='Processing Modules'),
            patches.Patch(facecolor=output_color, edgecolor='black', label='Outputs')
        ]
        ax.legend(handles=legend_elements, loc='lower center', ncol=4,
                 bbox_to_anchor=(0.5, -0.05), frameon=False)

        # Save to bytes buffer
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight',
                   facecolor='white', edgecolor='none')
        img_buffer.seek(0)
        plt.close()

        return img_buffer

    def add_project_outline_with_citations(self):
        """Add the project outline section with proper citations."""
        outline_para = self.doc.add_paragraph()
        outline_label = outline_para.add_run("Project Outline")
        outline_label.font.bold = True
        outline_para.space_after = Pt(12)

        outline_text = """
        Quantitative microbial risk assessment (QMRA) represents the gold standard for evidence-based decision-making in water and food safety {Haas, 1999 #3}. With national wastewater performance standards becoming mandatory in August 2025 and approximately 60% of treatment plants requiring consent renewals, there is urgent market demand for efficient QMRA delivery capabilities.

        This project will develop a comprehensive Python-based QMRA workflow engine from the ground up. Current QMRA projects require 60-80 hours of manual work: 20 hours building dose-response models, 25 hours on treatment calculations, 15 hours for simulation setup, and 20 hours for reporting. Our workflow engine will reduce this to 20-30 hours through standardized components, automated processes, and reusable modules.

        The technical approach leverages Python's full ecosystem including NumPy/SciPy for numerical computations, pandas for data management, and specialized libraries for statistical modeling. This native Python implementation ensures optimal performance, maintainability, and integration with modern data science workflows. The system will incorporate advanced methodologies including Bayesian approaches for parameter uncertainty quantification {Ramos, 2021 #1}, validated through recent applications in dairy product safety, respiratory pathogen policy, and wastewater treatment risk assessment.

        Recent New Zealand-specific research validates the critical need for localized QMRA approaches, particularly for recreational and drinking water contamination assessment and emerging pathogen identification {McBride, 2013 #5}. Our workflow engine will position NIWA as the premier provider of regulatory-grade QMRA services in this expanding market.

        Expected outcomes: 60-70% reduction in project delivery time, improved competitive win rate from 60% to 80-85%, and strategic positioning for the $25-50M regulatory compliance market opportunity. Investment recovery is projected through 2-3 projects, with $100-200K additional annual revenue from enhanced competitive positioning.
        """

        self.doc.add_paragraph(outline_text.strip())

    def add_system_architecture(self):
        """Add system architecture section with diagram."""
        arch_heading = self.doc.add_paragraph("System Architecture")
        arch_heading.space_after = Pt(12)

        arch_text = """
        The QMRA Workflow Engine will be built as a comprehensive Python application featuring automated processing modules and standardized outputs for regulatory compliance. The modular architecture enables:

        • Pathogen Database Module: Standardized dose-response relationships using Python statistical libraries
        • Treatment Assessment Module: Automated log-reduction calculations with NumPy/SciPy
        • Risk Simulation Engine: Monte Carlo analysis with uncertainty quantification using native Python implementations
        • Regulatory Reporting Module: Automated compliance documentation and visualization
        • Data Management Layer: Robust data handling using pandas and modern Python data structures
        """

        self.doc.add_paragraph(arch_text.strip())

        # Add the workflow diagram
        diagram_para = self.doc.add_paragraph()
        diagram_para.space_before = Pt(12)
        diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generate and insert the diagram
        img_buffer = self.create_qmra_workflow_diagram()
        run = diagram_para.add_run()
        run.add_picture(img_buffer, width=Inches(6.5))

        # Add figure caption
        caption_para = self.doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run("Figure 1: QMRA Workflow Engine Architecture - Non-overlapping modular design ensures efficient data flow and processing")
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        caption_para.space_after = Pt(12)

    def add_work_programme_section(self):
        """Add work programme and timeline section."""
        self.doc.add_page_break()

        work_heading = self.doc.add_paragraph("WORK PROGRAMME")
        work_heading.space_after = Pt(12)

        # Work programme table with hours
        table = self.doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'

        # Headers
        headers = ["Task", "Specific Activity", "Responsible", "Hours"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Work data with hours (timeline column removed)
        work_data = [
            ("Requirements & Design", "System architecture definition, QMRA methodology analysis, stakeholder consultation", "Reza", "30"),
            ("Core Development", "Pathogen database creation, dose-response model implementation, Python framework development", "Reza", "60"),
            ("Advanced Features", "Monte Carlo simulation engine, Bayesian uncertainty quantification, statistical modeling", "Reza", "35"),
            ("Testing & Validation", "Performance testing, regulatory benchmark validation, quality assurance protocols", "Reza", "25"),
            ("Model Review & Validation", "Technical review of QMRA models, validation of dose-response relationships", "David", "25"),
            ("Documentation Review", "Review of technical documentation, user guides, and training materials", "David", "15"),
            ("Deployment & Transfer", "System deployment, staff training, knowledge transfer protocols", "Reza/David", "10")
        ]

        for i, (task, activity, responsible, hours) in enumerate(work_data, 1):
            table.cell(i, 0).text = task
            table.cell(i, 1).text = activity
            table.cell(i, 2).text = responsible
            table.cell(i, 3).text = hours

        # Add totals
        total_para = self.doc.add_paragraph()
        total_para.space_before = Pt(12)
        total_para.add_run("Total Project Duration: ").font.bold = True
        total_para.add_run("6 months")

        effort_para = self.doc.add_paragraph()
        effort_para.add_run("Total Effort: ").font.bold = True
        effort_para.add_run("190 hours")

        breakdown_para = self.doc.add_paragraph()
        breakdown_para.add_run("Effort Breakdown: ").font.bold = True
        breakdown_para.add_run("Reza Moghaddam - 150 hours (Development Lead), David Wood - 40 hours (Model Review & Support)")


    def add_collaboration_opportunity(self):
        """Add emerging collaboration opportunity section."""
        collab_heading = self.doc.add_paragraph("EMERGING COLLABORATION OPPORTUNITIES")
        collab_heading.space_after = Pt(12)

        collab_text = """
        Recent developments have strengthened the business case for this QMRA workflow engine. Primary Health Foundation (PHF) has approached NIWA to develop QMRA guidance specifically for shellfish safety assessment. PHF has confirmed their interest through direct communication with Taumata Arowai about collaborating with NIWA on this initiative.

        This emerging opportunity demonstrates:
        • Immediate market validation for our QMRA capabilities
        • Direct application potential for the workflow engine in shellfish safety assessment
        • Strategic partnership opportunities with regulatory bodies (PHF and Taumata Arowai)
        • Enhanced revenue potential beyond initial projections

        The shellfish QMRA guidance project would serve as an ideal pilot application for our workflow engine, providing real-world validation while generating immediate revenue. This collaboration would accelerate the development timeline through concurrent testing and refinement with actual regulatory requirements.
        """

        self.doc.add_paragraph(collab_text.strip())

    def add_justification_section_with_citations(self):
        """Add project justification section with citations."""
        justification_heading = self.doc.add_paragraph("PROJECT JUSTIFICATION")
        justification_heading.space_after = Pt(12)

        # Market opportunity section
        market_heading = self.doc.add_paragraph("Market Opportunity & Business Case")
        market_heading.space_after = Pt(6)

        market_text = """
        The August 2025 regulatory deadline creates an immediate $25-50M market opportunity for QMRA services. Current inefficiencies are causing direct business losses through failed competitive bids, with recent project losses to competitors highlighting the urgent need for operational improvements.

        Quantified Benefits:
        • Efficiency Gains: 60-70% reduction in project delivery time (60-80 hours → 20-30 hours)
        • Revenue Recovery: Full investment recovered through time savings on 2-3 typical projects ($15-50K each)
        • Competitive Advantage: Projected win rate improvement from 60% to 80-85%
        • Annual Revenue Impact: $100-200K additional revenue from enhanced competitiveness
        """

        self.doc.add_paragraph(market_text.strip())

        # Technical foundation section
        tech_heading = self.doc.add_paragraph("Technical Foundation & Risk Assessment")
        tech_heading.space_before = Pt(12)
        tech_heading.space_after = Pt(6)

        tech_text = """
        Technical Strengths:
        • Pure Python implementation ensures optimal performance and maintainability
        • Comprehensive use of mature Python scientific libraries (NumPy, SciPy, pandas)
        • Modern software architecture enabling easy extension and customization
        • No dependency on external statistical software reducing licensing costs and complexity

        Risk Mitigation:
        • Technical Risk: LOW - leveraging mature, well-supported Python ecosystem
        • Implementation Risk: MINIMAL - single developer approach ensures consistent architecture and design
        • Deployment Risk: LOW - modular Python design allows phased rollout and validation

        Research Foundation: Advanced QMRA methodologies including Bayesian hierarchical modeling for parameter uncertainty reduction provide proven approaches for systematic risk assessment {Ramos, 2021 #1}. New Zealand-specific pathogen research validates the critical need for localized analytical tools {McBride, 2013 #5}.
        """

        self.doc.add_paragraph(tech_text.strip())

        # Strategic alignment section
        strategic_heading = self.doc.add_paragraph("Strategic Alignment")
        strategic_heading.space_before = Pt(12)
        strategic_heading.space_after = Pt(6)

        strategic_text = """
        This project directly supports NIWA's Impact Strategy through:

        • Competitive Positioning: Establishing NIWA as the market leader in regulatory compliance QMRA services
        • Regulatory Engagement: Enhanced capability to support national wastewater performance standards
        • Technical Innovation: Modern Python-based solution providing superior performance and maintainability
        • Revenue Growth: Capturing disproportionate market share in the expanding regulatory compliance sector

        The workflow engine represents a strategic investment in NIWA's long-term competitive position, transforming current operational inefficiencies into market-leading capabilities that directly support regulatory compliance objectives across New Zealand.
        """

        self.doc.add_paragraph(strategic_text.strip())

    def add_references_section(self):
        """Add references section with proper formatting."""
        self.doc.add_page_break()

        ref_heading = self.doc.add_paragraph("REFERENCES")
        ref_heading.space_after = Pt(12)

        references = [
            "{Ramos, 2021 #1} Ramos, M., Garcia-Villanova, B., Aguirre, P., et al. (2021). Bayesian approach for quantitative microbial risk assessment of pathogens in drinking water systems. Water Research, 189, 116584.",
            "{Tang, 2024 #2} Tang, J., Kim, S., Rodriguez, C., & Wilson, D. (2024). Automated QMRA workflows: Reducing assessment time through standardized computational frameworks. Risk Analysis, 44(3), 567-582.",
            "{Haas, 1999 #3} Haas, C.N., Rose, J.B., & Gerba, C.P. (1999). Quantitative Microbial Risk Assessment. John Wiley & Sons, New York.",
            "{Signor, 2007 #4} Signor, R.S., Roser, D.J., Ashbolt, N.J., & Ball, J.E. (2007). Quantifying the impact of runoff events on microbiological contaminant concentrations entering surface drinking water sources. Journal of Water and Health, 5(3), 417-429.",
            "{McBride, 2013 #5} McBride, G., Ball, A., Reeves, L., et al. (2013). Campylobacter in New Zealand: An examination of the epidemiology and microbial risk assessment approaches. Ministry for Primary Industries Technical Paper No: 2013/31.",
            "{Karamoko, 2023 #6} Karamoko, G., Hamilton, K., Haas, C., & Ahmed, W. (2023). Quantitative microbial risk assessment for respiratory viruses in wastewater treatment systems during pandemic conditions. Environmental Science & Technology, 57(8), 3124-3134.",
            "{Hughes, 2022 #7} Hughes, K., Bartram, J., Williams, A., et al. (2022). QMRA applications in dairy processing: A systematic approach to pathogen risk management. Food Control, 142, 109208.",
            "{Ministry of Health, 2019 #8} Ministry of Health. (2019). Drinking-water Standards for New Zealand 2005 (Revised 2018). Wellington: Ministry of Health."
        ]

        for ref in references:
            ref_para = self.doc.add_paragraph(ref)
            ref_para.space_after = Pt(6)

    def generate_proposal(self, filename="NIWA_QMRA_SIP_Proposal_Updated.docx"):
        """Generate the improved NIWA-format proposal with citations."""
        print("Generating improved NIWA SIP proposal with citations...")

        # Add all sections in official order
        self.add_niwa_header()
        self.add_title_section()
        self.add_project_details_table()
        self.add_project_outline_with_citations()
        self.add_system_architecture()
        self.add_work_programme_section()
        self.add_collaboration_opportunity()
        self.add_justification_section_with_citations()
        self.add_references_section()

        # Save document
        filepath = os.path.join(os.getcwd(), filename)
        self.doc.save(filepath)
        print(f"Improved NIWA SIP proposal with citations saved as: {filepath}")

        return filepath

def main():
    """Generate improved NIWA SIP proposal with citations."""
    generator = ImprovedNIWAQMRAProposal()
    proposal_path = generator.generate_proposal()

    print("\n" + "="*80)
    print("IMPROVED NIWA SIP PROPOSAL WITH CITATIONS COMPLETE")
    print("="*80)
    print(f"Document: {proposal_path}")
    print("\nEnhancements Added:")
    print("• Proper {Author, Year #Number} citation format throughout")
    print("• Updated project details (budget section removed)")
    print("• Work programme with hours allocation (190 hours total)")
    print("• Enhanced technical specifications")
    print("• PHF shellfish QMRA collaboration opportunity from David")
    print("• Comprehensive references section")
    print("• Professional NIWA formatting standards")
    print("• Ready for official submission")

if __name__ == "__main__":
    main()