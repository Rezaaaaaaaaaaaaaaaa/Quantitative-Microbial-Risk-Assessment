#!/usr/bin/env python3
"""
Remove the references section from the NIWA SIP proposal
"""

import docx
from docx import Document
import os

def remove_references_section(docx_path):
    """Remove the references/supporting information section from the document."""
    try:
        doc = Document(docx_path)

        # Find and remove paragraphs containing references
        paragraphs_to_remove = []
        remove_mode = False

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()

            # Start removing from "supporting information" or "references"
            if any(keyword in text for keyword in ['supporting information', 'references', 'key qmra literature']):
                remove_mode = True
                paragraphs_to_remove.append(i)
            elif remove_mode:
                paragraphs_to_remove.append(i)

        # Remove paragraphs in reverse order to maintain indices
        for i in reversed(paragraphs_to_remove):
            if i < len(doc.paragraphs):
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)

        # Save the modified document
        output_path = docx_path.replace('.docx', '_no_refs.docx')
        doc.save(output_path)

        # Replace original with cleaned version
        os.remove(docx_path)
        os.rename(output_path, docx_path)

        print(f"References section removed from {docx_path}")
        print(f"Removed {len(paragraphs_to_remove)} paragraphs")

        return True

    except Exception as e:
        print(f"Error removing references: {e}")
        return False

def main():
    """Main function to remove references from the SIP proposal."""
    docx_path = "NIWA_Authentic_SIP_QMRA_Proposal.docx"

    if os.path.exists(docx_path):
        success = remove_references_section(docx_path)
        if success:
            print("✅ References section successfully removed")
        else:
            print("❌ Failed to remove references section")
    else:
        print(f"❌ File {docx_path} not found")

if __name__ == "__main__":
    main()