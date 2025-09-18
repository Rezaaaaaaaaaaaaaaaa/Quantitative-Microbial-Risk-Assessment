#!/usr/bin/env python3
"""
Complete NIWA SIP Proposal Generator for QMRA Workflow Engine
Includes citations, system schematic, and all required elements
"""

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Rectangle
import os

class CompleteNIWAQMRAProposal:
    def __init__(self):
        self.doc = Document()
        self.setup_official_styles()

    def setup_official_styles(self):
        """Setup styles matching the official NIWA template."""
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Normal style
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(10)

    def create_qmra_schematic(self):
        """Create professional QMRA system schematic."""
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        ax.set_xlim(0, 12)
        ax.set_ylim(0, 8)
        ax.axis('off')

        # NIWA colors
        niwa_blue = '#0066CC'
        niwa_green = '#66B266'
        accent_orange = '#FF8C00'
        light_gray = '#F0F0F0'

        # Title
        ax.text(6, 7.5, 'QMRA Workflow Engine System Architecture',
                fontsize=16, fontweight='bold', ha='center')

        # Input Data Layer
        input_box = FancyBboxPatch((0.5, 5.5), 2.5, 1.5, boxstyle="round,pad=0.1",
                                  facecolor='lightblue', edgecolor=niwa_blue, linewidth=2)
        ax.add_patch(input_box)
        ax.text(1.75, 6.6, 'DATA INPUTS', fontsize=12, fontweight='bold', ha='center')
        ax.text(1.75, 6.1, '• Pathogen characteristics\n• Treatment parameters\n• Exposure scenarios\n• Population data',
                fontsize=9, ha='center')

        # Existing R Libraries
        r_box = FancyBboxPatch((4, 5.5), 2.5, 1.5, boxstyle="round,pad=0.1",
                              facecolor='lightyellow', edgecolor=accent_orange, linewidth=2)
        ax.add_patch(r_box)
        ax.text(5.25, 6.6, 'EXISTING R LIBRARIES', fontsize=12, fontweight='bold', ha='center')
        ax.text(5.25, 6.1, '• NIWA dose-response models\n• Treatment efficacy code\n• Validation datasets\n• Legacy algorithms',
                fontsize=9, ha='center')

        # Python Framework
        python_box = FancyBboxPatch((7.5, 5.5), 3.5, 1.5, boxstyle="round,pad=0.1",
                                   facecolor='lightgreen', edgecolor=niwa_green, linewidth=2)
        ax.add_patch(python_box)
        ax.text(9.25, 6.6, 'PYTHON FRAMEWORK', fontsize=12, fontweight='bold', ha='center')
        ax.text(9.25, 6.1, '• NumPy/SciPy scientific computing\n• rpy2 R integration interface\n• Pandas data management\n• Automated workflows',
                fontsize=9, ha='center')

        # Core Processing Engine
        core_box = FancyBboxPatch((2, 3.5), 8, 1.5, boxstyle="round,pad=0.1",
                                 facecolor=light_gray, edgecolor=niwa_blue, linewidth=3)
        ax.add_patch(core_box)
        ax.text(6, 4.7, 'QMRA CORE PROCESSING ENGINE', fontsize=14, fontweight='bold', ha='center')

        # Processing modules
        modules = [
            ('Pathogen\nDatabase', 2.8, 4.2),
            ('Dose-Response\nModeling', 4.8, 4.2),
            ('Monte Carlo\nSimulation', 6.8, 4.2),
            ('Risk\nCharacterization', 8.8, 4.2)
        ]

        for name, x, y in modules:
            module_box = Rectangle((x-0.4, y-0.3), 0.8, 0.6,
                                 facecolor='white', edgecolor=niwa_blue)
            ax.add_patch(module_box)
            ax.text(x, y, name, fontsize=9, fontweight='bold', ha='center', va='center')

        # Output Layer
        output_box = FancyBboxPatch((1, 1.5), 10, 1.5, boxstyle="round,pad=0.1",
                                   facecolor='lavender', edgecolor=niwa_blue, linewidth=2)
        ax.add_patch(output_box)
        ax.text(6, 2.7, 'AUTOMATED OUTPUTS', fontsize=12, fontweight='bold', ha='center')
        ax.text(6, 2.2, '• Regulatory compliance reports • Risk assessment summaries • Sensitivity analysis\n• Uncertainty quantification • Executive dashboards • Technical documentation',
                fontsize=10, ha='center')

        # Benefits box
        benefits_box = FancyBboxPatch((2, 0.2), 8, 0.8, boxstyle="round,pad=0.1",
                                     facecolor='lightcyan', edgecolor=accent_orange, linewidth=2)
        ax.add_patch(benefits_box)
        ax.text(6, 0.6, 'EXPECTED BENEFITS: 60-70% faster delivery • Standardized methodology • Enhanced competitiveness',
                fontsize=11, fontweight='bold', ha='center')

        # Arrows
        arrow_props = dict(arrowstyle='->', lw=2, color='black')

        # Input flows to core
        ax.annotate('', xy=(2.8, 4.9), xytext=(1.75, 5.5), arrowprops=arrow_props)
        ax.annotate('', xy=(6, 4.9), xytext=(5.25, 5.5), arrowprops=arrow_props)
        ax.annotate('', xy=(8.2, 4.9), xytext=(9.25, 5.5), arrowprops=arrow_props)

        # Core to output
        ax.annotate('', xy=(6, 3), xytext=(6, 3.5), arrowprops=arrow_props)

        plt.tight_layout()
        diagram_path = os.path.join(os.getcwd(), 'qmra_system_schematic.png')
        plt.savefig(diagram_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        return diagram_path

    def add_niwa_header(self):
        """Add NIWA logo and header."""
        header_para = self.doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run("NIWA")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_para.space_after = Pt(12)

    def add_title_section(self):
        """Add the title section exactly matching the template."""
        title_para = self.doc.add_paragraph("Structured internal project application 2024-2025")
        title_para.space_after = Pt(18)

        overview_para = self.doc.add_paragraph("Project Overview")
        overview_para.space_after = Pt(12)

    def add_project_details_table(self):
        """Add the project details table with all information."""
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(4.5)

        # Project details with proper field names
        table.cell(0, 0).text = "Project name: (Short title)"
        table.cell(0, 1).text = "Development of QMRA Workflow Engine"

        table.cell(1, 0).text = "Staff: (who will be completing the work?)"
        table.cell(1, 1).text = "Reza Moghaddam"

        table.cell(2, 0).text = "Project Manager: (usually a Group Manager)"
        table.cell(2, 1).text = "David Wood (with support from Andrew Hughes as Project Director)"

        table.cell(3, 0).text = "Region:"
        table.cell(3, 1).text = "National"

        table.cell(4, 0).text = "Centre:"
        table.cell(4, 1).text = "Environmental Information"

        table.cell(5, 0).text = "Type: (science, operations activity, or other - explain)"
        table.cell(5, 1).text = "Science"

        table.cell(6, 0).text = "Budget: (attach costing prepared by your project coordinator)"
        table.cell(6, 1).text = "$25,000"

        table.cell(7, 0).text = "Project objective: (30 words max)"
        table.cell(7, 1).text = "Develop Python-based QMRA workflow engine to reduce project delivery time by 60-70% and strengthen competitive position for regulatory compliance market."

        # Make labels bold
        for i in range(8):
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    def add_project_outline_with_citations(self):
        """Add the project outline section with proper citations."""
        outline_para = self.doc.add_paragraph()
        outline_label = outline_para.add_run("Project outline: (150-300 words max)")
        outline_label.font.bold = True
        outline_para.space_after = Pt(6)

        outline_text = """
        NIWA is losing competitive bids because our manual QMRA process takes 60% longer than competitors using automated tools. Quantitative microbial risk assessment represents the gold standard for evidence-based decision-making in water and food safety {Tang, 2024 #2; Ramos, 2021 #1}. With national wastewater performance standards mandatory from August 2025 and 60% of treatment plants requiring consent renewals, this creates urgent need for efficiency improvements.

        This SIP will develop a Python-based QMRA workflow engine integrating our existing R code libraries. Recent applications demonstrate QMRA's expanding utility in dairy product safety {Ramos, 2021 #1}, respiratory pathogen policy {Tang, 2024 #2}, and wastewater treatment risk assessment {Dada, 2021 #4}. New Zealand-specific research validates the need for localized QMRA approaches, particularly for drinking water contamination {Phiri, 2021 #6} and novel pathogen identification {Bloomfield, 2020 #9}.

        Current projects require 60-80 hours: 20 hours building dose-response models, 25 hours treatment calculations, 15 hours simulation setup, 20 hours reporting. The workflow engine will reduce this to 20-30 hours through standardized components and automated processes.

        Technical approach uses Python with rpy2 interfaces for R integration, leveraging NumPy/SciPy for performance. Advanced methodologies including Bayesian approaches for parameter uncertainty reduction {Seis, 2020 #8} provide technical direction. David Wood will handle R code integration while Reza leads Python development.

        Expected outcomes: 60-70% reduction in delivery time, improved win rate from 60% to 80-85%, positioning for $25-50M regulatory compliance market. Investment recovery through 2-3 projects, with $100-200k additional annual revenue from enhanced competitiveness.
        """

        outline_content = self.doc.add_paragraph(outline_text.strip())

    def add_system_schematic(self):
        """Add the system schematic diagram."""
        schematic_heading = self.doc.add_paragraph("System Architecture Schematic")
        schematic_heading.space_before = Pt(18)
        schematic_heading.space_after = Pt(12)

        # Create and insert diagram
        diagram_path = self.create_qmra_schematic()

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_picture(diagram_path, width=Inches(6.5))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Caption
        caption = self.doc.add_paragraph("Figure 1: QMRA Workflow Engine showing integration of existing R libraries with new Python framework, automated processing modules, and standardized outputs for regulatory compliance.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.space_after = Pt(12)

        # Clean up diagram file
        if os.path.exists(diagram_path):
            os.remove(diagram_path)

    def add_work_programme_section(self):
        """Add work programme and timeline section."""
        self.doc.add_page_break()

        work_heading = self.doc.add_paragraph("WORK PROGRAMME AND TIMELINE")
        work_heading.space_after = Pt(12)

        instruction = self.doc.add_paragraph("Outline the tasks to be done, who will do what and by when. Be as specific as possible.")
        instruction.space_after = Pt(12)

        # Work programme table
        table = self.doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'

        # Headers
        headers = ["Task", "Specific activity (who, what)", "By when", "Hours"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Timeline data
        work_data = [
            ("Requirements & Design", "System architecture, R code audit, stakeholder consultation (Reza, David)", "Month 1", "40"),
            ("Core Development", "Pathogen database creation, dose-response model integration (Reza)", "Month 2-3", "60"),
            ("R Integration", "rpy2 implementation, legacy code integration, validation (David, Reza)", "Month 4", "30"),
            ("Testing & Validation", "Performance testing, benchmark validation, quality assurance (Reza)", "Month 5", "25"),
            ("Documentation & Training", "User guides, technical documentation, training materials (Reza, David)", "Month 6", "20"),
            ("Deployment", "System deployment, staff training, knowledge transfer (Team)", "Month 6", "15")
        ]

        for i, (task, activity, when, hours) in enumerate(work_data, 1):
            table.cell(i, 0).text = task
            table.cell(i, 1).text = activity
            table.cell(i, 2).text = when
            table.cell(i, 3).text = hours

    def add_budget_details(self):
        """Add detailed budget breakdown."""
        budget_heading = self.doc.add_paragraph("BUDGET BREAKDOWN")
        budget_heading.space_after = Pt(12)

        # Budget table
        table = self.doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        headers = ["Item", "Rate/Cost", "Total"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        budget_data = [
            ("Reza Moghaddam (150 hours)", "$125/hour", "$18,750"),
            ("David Wood (40 hours)", "$125/hour", "$5,000"),
            ("Computing resources & software", "Fixed cost", "$750"),
            ("Documentation & training materials", "Fixed cost", "$500")
        ]

        for i, (item, rate, total) in enumerate(budget_data, 1):
            table.cell(i, 0).text = item
            table.cell(i, 1).text = rate
            table.cell(i, 2).text = total

        # Total
        total_para = self.doc.add_paragraph()
        total_para.space_before = Pt(12)
        total_run = total_para.add_run("TOTAL PROJECT COST: $25,000")
        total_run.font.bold = True

    def add_justification_with_citations(self):
        """Add project justification with supporting citations."""
        justification_heading = self.doc.add_paragraph("PROJECT JUSTIFICATION")
        justification_heading.space_after = Pt(12)

        justification_text = """
        Current inefficiencies are causing direct business losses through failed competitive bids. Recent project losses to competitors demonstrate immediate need for operational improvements. International food safety risk assessment systems show standardized approaches provide competitive advantages {Ng, 2022 #3}.

        Market opportunity: August 2025 regulatory deadline creates $25-50M market for QMRA services. Organizations with standardized tools will capture disproportionate market share, as demonstrated in Australian campylobacter risk assessment applications {Habib, 2020 #10}.

        Return on investment: Time savings on 2-3 typical projects ($15-50k each) will recover full investment. Enhanced competitiveness projected to increase win rate by 25-30%, generating $100-200k additional annual revenue.

        Technical foundation: Advanced QMRA methodologies including Bayesian hierarchical modeling for parameter uncertainty reduction {Seis, 2020 #8} provide proven approaches for systematic risk assessment. New Zealand-specific pathogen research {Bloomfield, 2020 #9} validates need for localized tools.

        Risk mitigation: Low technical risk due to proven technologies (Python, rpy2) and experienced team. David Wood's R expertise minimizes integration challenges. Modular design allows incremental deployment.

        Strategic importance: Positions NIWA as leader in regulatory compliance market while preserving existing intellectual property through R code integration. This aligns with NIWA's Impact Strategy on competitive positioning and regulatory engagement.
        """

        self.doc.add_paragraph(justification_text.strip())

    def add_references_section(self):
        """Add supporting references section."""
        ref_heading = self.doc.add_paragraph("SUPPORTING REFERENCES")
        ref_heading.space_before = Pt(18)
        ref_heading.space_after = Pt(12)

        references = [
            "{Bloomfield, 2020 #9} Bloomfield, S., et al. Campylobacter novaezeelandiae sp. nov., isolated from birds and water in New Zealand. Int J Syst Evol Microbiol. 2020;70(6):3775-3784.",

            "{Dada, 2021 #4} Dada, A.C., Gyawali, P. Quantitative microbial risk assessment (QMRA) of occupational exposure to SARS-CoV-2 in wastewater treatment plants. Sci Total Environ. 2021;763:142989.",

            "{Habib, 2020 #10} Habib, I., et al. Human campylobacteriosis related to cross-contamination during handling of raw chicken meat: Application of quantitative risk assessment to guide intervention scenarios analysis in the Australian context. Int J Food Microbiol. 2020;332:108775.",

            "{Ng, 2022 #3} Ng, S., Shao, S., Ling, N. Food safety risk-assessment systems utilized by China, Australia/New Zealand, Canada, and the United States. J Food Sci. 2022;87(11):4780-4795.",

            "{Phiri, 2021 #6} Phiri, B.J., et al. Microbial contamination in drinking water at public outdoor recreation facilities in New Zealand. J Appl Microbiol. 2021;130(1):302-312.",

            "{Ramos, 2021 #1} Ramos, G.L.P.A., et al. Quantitative microbiological risk assessment in dairy products: Concepts and applications. Trends Food Sci Technol. 2021;111:610-616.",

            "{Seis, 2020 #8} Seis, W., Rouault, P., Medema, G. Addressing and reducing parameter uncertainty in quantitative microbial risk assessment by incorporating external information via Bayesian hierarchical modeling. Water Res. 2020;185:116202.",

            "{Tang, 2024 #2} Tang, L., et al. Applications of quantitative microbial risk assessment to respiratory pathogens and implications for uptake in policy: A state-of-the-science review. Environ Health Perspect. 2024;132(5):56001."
        ]

        for ref in references:
            ref_para = self.doc.add_paragraph(ref)
            ref_para.paragraph_format.hanging_indent = Inches(0.5)
            ref_para.space_after = Pt(3)

    def generate_complete_proposal(self, filename="Complete_NIWA_QMRA_SIP_Proposal.docx"):
        """Generate the complete proposal with all elements."""
        print("Generating complete NIWA SIP proposal with citations and schematic...")

        # Add all sections
        self.add_niwa_header()
        self.add_title_section()
        self.add_project_details_table()
        self.add_project_outline_with_citations()
        self.add_system_schematic()
        self.add_work_programme_section()
        self.add_budget_details()
        self.add_justification_with_citations()
        self.add_references_section()

        # Save document
        filepath = os.path.join(os.getcwd(), filename)
        self.doc.save(filepath)
        print(f"Complete NIWA SIP proposal saved as: {filepath}")

        return filepath

def main():
    """Generate complete NIWA SIP proposal with all elements."""
    generator = CompleteNIWAQMRAProposal()
    proposal_path = generator.generate_complete_proposal()

    print("\n" + "="*80)
    print("COMPLETE NIWA SIP PROPOSAL WITH CITATIONS & SCHEMATIC")
    print("="*80)
    print(f"Document: {proposal_path}")
    print("\nComplete Elements Included:")
    print("✅ Official NIWA template format and structure")
    print("✅ Proper citations throughout text {Author, Year #Number}")
    print("✅ Professional system architecture schematic")
    print("✅ Work programme with detailed timeline")
    print("✅ Budget breakdown and justification")
    print("✅ Supporting references section")
    print("✅ Ready for official submission")

if __name__ == "__main__":
    main()