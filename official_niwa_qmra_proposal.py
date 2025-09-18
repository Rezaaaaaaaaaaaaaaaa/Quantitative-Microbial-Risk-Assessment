#!/usr/bin/env python3
"""
Official NIWA SIP Proposal Generator for QMRA Workflow Engine
Matches the exact NIWA template format from the provided image
"""

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

class OfficialNIWAQMRAProposal:
    def __init__(self):
        self.doc = Document()
        self.setup_official_styles()

    def setup_official_styles(self):
        """Setup styles matching the official NIWA template."""
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Normal style
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(10)

    def add_niwa_header(self):
        """Add NIWA logo and header."""
        # Add space for logo (would be inserted manually or via template)
        header_para = self.doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run("NIWA")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_para.space_after = Pt(12)

    def add_title_section(self):
        """Add the title section exactly matching the template."""
        title_para = self.doc.add_paragraph("Structured internal project application 2024-2025")
        title_para.space_after = Pt(18)

        overview_para = self.doc.add_paragraph("Project Overview")
        overview_para.space_after = Pt(12)

    def add_project_details_table(self):
        """Add the project details table matching the exact template format."""
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'

        # Set column widths (approximate)
        for row in table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(4.5)

        # Row 1: Project name
        table.cell(0, 0).text = "Project name: (Short title)"
        table.cell(0, 1).text = "Development of QMRA Workflow Engine"

        # Row 2: Staff
        table.cell(1, 0).text = "Staff: (who will be completing the work?)"
        table.cell(1, 1).text = "Reza Moghaddam"

        # Row 3: Project Manager
        table.cell(2, 0).text = "Project Manager: (usually a Group Manager)"
        table.cell(2, 1).text = "David Wood (with support from Andrew Hughes as Project Director)"

        # Row 4: Region
        table.cell(3, 0).text = "Region:"
        table.cell(3, 1).text = "National"

        # Row 5: Centre
        table.cell(4, 0).text = "Centre:"
        table.cell(4, 1).text = "Environmental Information"

        # Row 6: Type
        table.cell(5, 0).text = "Type: (science, operations activity, or other - explain)"
        table.cell(5, 1).text = "Science"

        # Row 7: Budget
        table.cell(6, 0).text = "Budget: (attach costing prepared by your project coordinator)"
        table.cell(6, 1).text = "$25,000"

        # Row 8: Project objective
        table.cell(7, 0).text = "Project objective: (30 words max)"
        table.cell(7, 1).text = "Develop Python-based QMRA workflow engine to reduce project delivery time by 60-70% and strengthen competitive position for regulatory compliance market."

        # Make labels bold
        for i in range(8):
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    def add_project_outline(self):
        """Add the project outline section."""
        outline_para = self.doc.add_paragraph()
        outline_label = outline_para.add_run("Project outline: (150-300 words max)")
        outline_label.font.bold = True
        outline_para.space_after = Pt(6)

        outline_text = """
        NIWA is losing competitive bids because our manual QMRA process takes 60% longer than competitors using automated tools. With national wastewater performance standards mandatory from August 2025 and 60% of treatment plants requiring consent renewals, this creates urgent need for efficiency improvements.

        This SIP will develop a Python-based QMRA workflow engine integrating our existing R code libraries. The system will standardize pathogen dose-response relationships, treatment efficacy models, and Monte Carlo simulations while automating regulatory-compliant report generation.

        Current projects require 60-80 hours: 20 hours building dose-response models, 25 hours treatment calculations, 15 hours simulation setup, 20 hours reporting. The workflow engine will reduce this to 20-30 hours through standardized components and automated processes.

        Technical approach uses Python with rpy2 interfaces for R integration, leveraging NumPy/SciPy for performance. David Wood will handle R code integration while Reza leads Python development and system architecture.

        Expected outcomes: 60-70% reduction in delivery time, improved win rate from 60% to 80-85%, positioning for $25-50M regulatory compliance market. Investment recovery through 2-3 projects, with $100-200k additional annual revenue from enhanced competitiveness.

        The August 2025 deadline makes this time-critical for capturing market opportunity. Standardized tools will enable scaling to multi-site assessments and ensure consistent quality across projects.
        """

        outline_content = self.doc.add_paragraph(outline_text.strip())

    def add_work_programme_section(self):
        """Add work programme and timeline section."""
        self.doc.add_page_break()

        work_heading = self.doc.add_paragraph("WORK PROGRAMME AND TIMELINE")
        work_heading.space_after = Pt(12)

        instruction = self.doc.add_paragraph("Outline the tasks to be done, who will do what and by when. Be as specific as possible.")
        instruction.space_after = Pt(12)

        # Work programme table
        table = self.doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'

        # Headers
        headers = ["Task", "Specific activity (who, what)", "By when", "Hours"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Timeline data
        work_data = [
            ("Requirements & Design", "System architecture, R code audit (Reza, David)", "Month 1", "40"),
            ("Core Development", "Pathogen database, simulation engine (Reza)", "Month 2-3", "60"),
            ("R Integration", "rpy2 implementation, legacy code integration (David, Reza)", "Month 4", "30"),
            ("Testing & Validation", "Performance testing, benchmark validation (Reza)", "Month 5", "25"),
            ("Documentation", "User guides, technical documentation (Reza, David)", "Month 6", "20"),
            ("Deployment", "System deployment, staff training (Team)", "Month 6", "15")
        ]

        for i, (task, activity, when, hours) in enumerate(work_data, 1):
            table.cell(i, 0).text = task
            table.cell(i, 1).text = activity
            table.cell(i, 2).text = when
            table.cell(i, 3).text = hours

    def add_budget_details(self):
        """Add detailed budget breakdown."""
        budget_heading = self.doc.add_paragraph("BUDGET BREAKDOWN")
        budget_heading.space_after = Pt(12)

        # Budget table
        table = self.doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        headers = ["Item", "Rate/Cost", "Total"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        budget_data = [
            ("Reza Moghaddam (150 hours)", "$125/hour", "$18,750"),
            ("David Wood (40 hours)", "$125/hour", "$5,000"),
            ("Computing resources", "Fixed cost", "$750"),
            ("Documentation/training", "Fixed cost", "$500")
        ]

        for i, (item, rate, total) in enumerate(budget_data, 1):
            table.cell(i, 0).text = item
            table.cell(i, 1).text = rate
            table.cell(i, 2).text = total

        # Total
        total_para = self.doc.add_paragraph()
        total_para.space_before = Pt(12)
        total_run = total_para.add_run("TOTAL PROJECT COST: $25,000")
        total_run.font.bold = True

    def add_justification_section(self):
        """Add project justification section."""
        justification_heading = self.doc.add_paragraph("PROJECT JUSTIFICATION")
        justification_heading.space_after = Pt(12)

        justification_text = """
        Current inefficiencies are causing direct business losses through failed competitive bids. Recent project losses to competitors demonstrate immediate need for operational improvements.

        Market opportunity: August 2025 regulatory deadline creates $25-50M market for QMRA services. Organizations with standardized tools will capture disproportionate market share.

        Return on investment: Time savings on 2-3 typical projects ($15-50k each) will recover full investment. Enhanced competitiveness projected to increase win rate by 25-30%, generating $100-200k additional annual revenue.

        Risk mitigation: Low technical risk due to proven technologies (Python, rpy2) and experienced team. David Wood's R expertise minimizes integration challenges. Modular design allows incremental deployment.

        Strategic importance: Positions NIWA as leader in regulatory compliance market while preserving existing intellectual property through R code integration.
        """

        self.doc.add_paragraph(justification_text.strip())

    def generate_proposal(self, filename="Official_NIWA_QMRA_SIP_Proposal.docx"):
        """Generate the official NIWA-format proposal."""
        print("Generating official NIWA SIP proposal matching template format...")

        # Add all sections in official order
        self.add_niwa_header()
        self.add_title_section()
        self.add_project_details_table()
        self.add_project_outline()
        self.add_work_programme_section()
        self.add_budget_details()
        self.add_justification_section()

        # Save document
        filepath = os.path.join(os.getcwd(), filename)
        self.doc.save(filepath)
        print(f"Official NIWA SIP proposal saved as: {filepath}")

        return filepath

def main():
    """Generate official NIWA SIP proposal."""
    generator = OfficialNIWAQMRAProposal()
    proposal_path = generator.generate_proposal()

    print("\n" + "="*70)
    print("OFFICIAL NIWA SIP PROPOSAL COMPLETE")
    print("="*70)
    print(f"Document: {proposal_path}")
    print("\nOfficial Template Features:")
    print("• Exact field structure from NIWA template")
    print("• Proper table formatting and layout")
    print("• Official terminology and style")
    print("• Structured sections matching template")
    print("• Professional NIWA formatting standards")
    print("• Ready for official submission")

if __name__ == "__main__":
    main()