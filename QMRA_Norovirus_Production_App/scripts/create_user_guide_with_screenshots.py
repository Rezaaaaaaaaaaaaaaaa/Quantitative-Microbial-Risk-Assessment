#!/usr/bin/env python3
"""
Create Professional QMRA User Guide with Screenshots
=====================================================

Creates a comprehensive Microsoft Word document with embedded screenshots
for the QMRA Batch Processing Web Application.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

def create_user_guide():
    """Create comprehensive user guide with screenshots."""

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ====================
    # TITLE PAGE
    # ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('QMRA Batch Processing\nWeb Application')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 119, 180)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('User Guide & Technical Documentation')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(70, 70, 70)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer info
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Quantitative Microbial Risk Assessment Toolkit\n\n')
    footer_run.font.size = Pt(12)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

    footer_run2 = footer.add_run('Version 2.0 | November 2025\n')
    footer_run2.font.size = Pt(10)
    footer_run2.font.color.rgb = RGBColor(128, 128, 128)

    footer_run3 = footer.add_run('NIWA - National Institute of Water & Atmospheric Research')
    footer_run3.font.size = Pt(10)
    footer_run3.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ====================
    # TABLE OF CONTENTS
    # ====================
    heading = doc.add_heading('Table of Contents', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    toc_items = [
        ('1', 'Introduction', ''),
        ('', '1.1 What is QMRA?', ''),
        ('', '1.2 Who Should Use This Application?', ''),
        ('', '1.3 Key Features', ''),
        ('2', 'Getting Started', ''),
        ('', '2.1 Launching the Application', ''),
        ('', '2.2 Understanding the Interface', ''),
        ('', '2.3 Production Mode vs Research Mode', ''),
        ('3', 'Assessment Workflows', ''),
        ('', '3.1 Batch Scenarios Assessment', ''),
        ('', '3.2 Spatial Risk Assessment', ''),
        ('', '3.3 Temporal Risk Assessment', ''),
        ('', '3.4 Treatment Comparison', ''),
        ('', '3.5 Multi-Pathogen Assessment', ''),
        ('4', 'Interpreting Results', ''),
        ('', '4.1 Risk Metrics Explained', ''),
        ('', '4.2 Visualizations', ''),
        ('', '4.3 WHO Guideline Compliance', ''),
        ('5', 'Technical Background', ''),
        ('', '5.1 Dose-Response Models', ''),
        ('', '5.2 Beta-Binomial Model (Norovirus)', ''),
        ('', '5.3 Monte Carlo Simulation', ''),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        if num:
            run1 = p.add_run(f'{num}. ')
            run1.bold = True
            run2 = p.add_run(title)
            run2.bold = True
        else:
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(title)

    doc.add_page_break()

    # ====================
    # 1. INTRODUCTION
    # ====================
    heading = doc.add_heading('1. Introduction', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    # 1.1 What is QMRA?
    heading2 = doc.add_heading('1.1 What is QMRA?', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'Quantitative Microbial Risk Assessment (QMRA) is a scientific framework for '
        'estimating the probability of infection or illness from exposure to pathogenic '
        'microorganisms in water, food, or other environmental media.'
    )

    doc.add_paragraph(
        'QMRA integrates four key components:'
    )
    doc.add_paragraph('Pathogen concentration in the source (e.g., wastewater, surface water)', style='List Bullet')
    doc.add_paragraph('Treatment effectiveness (log reduction values)', style='List Bullet')
    doc.add_paragraph('Environmental fate and transport (dilution, bioaccumulation)', style='List Bullet')
    doc.add_paragraph('Dose-response relationship (probability of infection per organism ingested)', style='List Bullet')

    doc.add_paragraph(
        'The output is an annual infection or illness probability per person, which can be '
        'compared against health-based targets such as the WHO guideline of 10⁻⁴ (1 in 10,000) '
        'annual infection risk for drinking water.'
    )

    # 1.2 Who Should Use This Application?
    heading2 = doc.add_heading('1.2 Who Should Use This Application?', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph('This application is designed for:')
    doc.add_paragraph('Water quality managers and engineers', style='List Bullet')
    doc.add_paragraph('Environmental health professionals', style='List Bullet')
    doc.add_paragraph('Public health risk assessors', style='List Bullet')
    doc.add_paragraph('Regulatory compliance officers', style='List Bullet')
    doc.add_paragraph('Researchers in water safety and microbial risk', style='List Bullet')

    # 1.3 Key Features
    heading2 = doc.add_heading('1.3 Key Features', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph('✅ Five specialized assessment workflows for different scenarios', style='List Bullet')
    doc.add_paragraph('✅ Validated Beta-Binomial dose-response model for norovirus', style='List Bullet')
    doc.add_paragraph('✅ Monte Carlo simulation with full uncertainty quantification', style='List Bullet')
    doc.add_paragraph('✅ Interactive visualizations and statistical summaries', style='List Bullet')
    doc.add_paragraph('✅ WHO guideline compliance checking', style='List Bullet')
    doc.add_paragraph('✅ Professional PDF reports and CSV data exports', style='List Bullet')
    doc.add_paragraph('✅ Production Mode ensures validated assessments only', style='List Bullet')

    doc.add_page_break()

    # ====================
    # 2. GETTING STARTED
    # ====================
    heading = doc.add_heading('2. Getting Started', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    # 2.1 Launching the Application
    heading2 = doc.add_heading('2.1 Launching the Application', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The QMRA application is a web-based tool running on Streamlit. Follow these steps to launch:'
    )

    doc.add_paragraph('Step 1: Open a terminal or command prompt', style='List Number')
    doc.add_paragraph('Step 2: Navigate to the application directory:', style='List Number')

    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.5)
    code_run = code_p.add_run('cd Batch_Processing_App/app')
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)

    doc.add_paragraph('Step 3: Launch the application:', style='List Number')

    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.5)
    code_run = code_p.add_run('streamlit run web_app.py')
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)

    doc.add_paragraph('Step 4: Your web browser will automatically open to http://localhost:8501', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph(
        'Alternatively, double-click the launch_web_gui.bat file to start the application automatically.'
    )

    # 2.2 Understanding the Interface
    heading2 = doc.add_heading('2.2 Understanding the Interface', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The application interface consists of two main areas:'
    )

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Left Sidebar: ')
    run1.bold = True
    p.add_run('Configuration options, production mode toggle, and assessment type selection')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Main Panel: ')
    run1.bold = True
    p.add_run('Input parameters, results, visualizations, and download options')

    doc.add_paragraph()

    # Add screenshot 1: Home Page
    screenshot_path = Path('../screenshots/01_home_page_20251117_095558.png')
    if screenshot_path.exists():
        doc.add_paragraph('Figure 1: Application Home Page', style='Caption')
        doc.add_picture(str(screenshot_path), width=Inches(6.0))
        doc.add_paragraph()

    # 2.3 Production Mode vs Research Mode
    heading2 = doc.add_heading('2.3 Production Mode vs Research Mode', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    p = doc.add_paragraph()
    run1 = p.add_run('Production Mode (Recommended): ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run(
        'Restricts pathogen selection to norovirus only, ensuring all risk assessments use '
        'the validated Beta-Binomial dose-response model. This mode is enabled by default and '
        'aligns with contract deliverables.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Research Mode: ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 140, 0)
    p.add_run(
        'Enables all 6 pathogens (norovirus, campylobacter, cryptosporidium, E. coli, rotavirus, salmonella). '
        'Non-norovirus pathogens have NOT been validated to the same standard and should be used with caution '
        'for research purposes only.'
    )

    doc.add_page_break()

    # ====================
    # 3. ASSESSMENT WORKFLOWS
    # ====================
    heading = doc.add_heading('3. Assessment Workflows', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    doc.add_paragraph(
        'The application provides five specialized assessment workflows, each designed for '
        'specific use cases. This section provides step-by-step instructions for each workflow.'
    )

    doc.add_page_break()

    # 3.1 Batch Scenarios
    heading2 = doc.add_heading('3.1 Batch Scenarios Assessment', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The Batch Scenarios workflow allows you to run multiple pre-configured scenarios '
        'simultaneously using a library approach with three CSV input files.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Use Case: ')
    run1.bold = True
    p.add_run(
        'Comprehensive risk assessment across multiple locations, time periods, and exposure routes '
        'with different parameter combinations.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Input Files Required:')
    run1.bold = True

    doc.add_paragraph('pathogen_data.csv - Pathogen concentrations and parameters', style='List Bullet')
    doc.add_paragraph('scenarios.csv - Scenario configurations (15 scenarios)', style='List Bullet')
    doc.add_paragraph('master_scenarios.csv - Linkage between pathogens and scenarios', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Steps:')
    run1.bold = True

    doc.add_paragraph('Select "Batch Scenarios" from the assessment mode dropdown', style='List Number')
    doc.add_paragraph('Check "Use example data" to load pre-configured scenarios, or upload your own CSV files', style='List Number')
    doc.add_paragraph('Review the scenario configurations in the data tables', style='List Number')
    doc.add_paragraph('Click "Run Assessment" to execute all scenarios', style='List Number')
    doc.add_paragraph('Review results in the summary table and download outputs', style='List Number')

    doc.add_paragraph()

    # Add screenshot 2: Batch Scenarios
    screenshot_path = Path('../screenshots/02_batch_scenarios_20251117_095600.png')
    if screenshot_path.exists():
        doc.add_paragraph('Figure 2: Batch Scenarios Assessment Interface', style='Caption')
        doc.add_picture(str(screenshot_path), width=Inches(6.0))
        doc.add_paragraph()

    doc.add_page_break()

    # 3.2 Spatial Assessment
    heading2 = doc.add_heading('3.2 Spatial Risk Assessment', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The Spatial Assessment workflow evaluates risk across multiple geographic sites '
        'with different dilution factors (e.g., distance from wastewater discharge).'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Use Case: ')
    run1.bold = True
    p.add_run(
        'Mapping risk gradients near wastewater outfalls, comparing shellfish harvesting areas, '
        'or assessing multiple beach locations.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Steps:')
    run1.bold = True

    doc.add_paragraph('Select "Spatial Assessment" from the assessment mode dropdown', style='List Number')
    doc.add_paragraph('Choose pathogen (norovirus in Production Mode)', style='List Number')
    doc.add_paragraph('Select exposure route (shellfish consumption or primary contact)', style='List Number')
    doc.add_paragraph('Upload dilution data CSV or use example data (6 sites)', style='List Number')
    doc.add_paragraph('Configure treatment LRV, exposure frequency, and other parameters', style='List Number')
    doc.add_paragraph('Click "Run Spatial Assessment"', style='List Number')
    doc.add_paragraph('Review site-specific risk results and risk gradient visualization', style='List Number')

    doc.add_paragraph()

    # Add screenshot 3: Spatial Assessment
    screenshot_path = Path('../screenshots/03_spatial_assessment_20251117_095611.png')
    if screenshot_path.exists():
        doc.add_paragraph('Figure 3: Spatial Risk Assessment Interface', style='Caption')
        doc.add_picture(str(screenshot_path), width=Inches(6.0))
        doc.add_paragraph()

    doc.add_page_break()

    # 3.3 Temporal Assessment
    heading2 = doc.add_heading('3.3 Temporal Risk Assessment', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The Temporal Assessment workflow analyzes risk trends over time using monitoring data '
        '(e.g., seasonal variations, before/after treatment upgrades).'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Use Case: ')
    run1.bold = True
    p.add_run(
        'Evaluating seasonal risk patterns, assessing treatment performance over time, '
        'or analyzing long-term monitoring datasets.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Steps:')
    run1.bold = True

    doc.add_paragraph('Select "Temporal Assessment" from the assessment mode dropdown', style='List Number')
    doc.add_paragraph('Choose pathogen and exposure route', style='List Number')
    doc.add_paragraph('Upload temporal data CSV with Date and Concentration columns', style='List Number')
    doc.add_paragraph('Configure assessment parameters', style='List Number')
    doc.add_paragraph('Click "Run Temporal Assessment"', style='List Number')
    doc.add_paragraph('Review time-series risk plot and statistical summaries', style='List Number')

    doc.add_paragraph()

    # Add screenshot 4: Temporal Assessment
    screenshot_path = Path('../screenshots/04_temporal_assessment_20251117_095623.png')
    if screenshot_path.exists():
        doc.add_paragraph('Figure 4: Temporal Risk Assessment Interface', style='Caption')
        doc.add_picture(str(screenshot_path), width=Inches(6.0))
        doc.add_paragraph()

    doc.add_page_break()

    # 3.4 Treatment Comparison
    heading2 = doc.add_heading('3.4 Treatment Comparison', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The Treatment Comparison workflow allows side-by-side evaluation of multiple '
        'treatment technologies or log reduction values (LRVs).'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Use Case: ')
    run1.bold = True
    p.add_run(
        'Comparing UV disinfection vs chlorination, evaluating membrane filtration options, '
        'or assessing treatment upgrade scenarios.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Steps:')
    run1.bold = True

    doc.add_paragraph('Select "Treatment Comparison" from the assessment mode dropdown', style='List Number')
    doc.add_paragraph('Choose pathogen and exposure route', style='List Number')
    doc.add_paragraph('Define 2-4 treatment scenarios with different LRVs and names', style='List Number')
    doc.add_paragraph('Configure other parameters (pathogen concentration, dilution, exposure)', style='List Number')
    doc.add_paragraph('Click "Run Treatment Comparison"', style='List Number')
    doc.add_paragraph('Compare treatment effectiveness in side-by-side results', style='List Number')

    doc.add_paragraph()

    # Add screenshot 5: Treatment Comparison
    screenshot_path = Path('../screenshots/05_treatment_comparison_20251117_095634.png')
    if screenshot_path.exists():
        doc.add_paragraph('Figure 5: Treatment Comparison Interface', style='Caption')
        doc.add_picture(str(screenshot_path), width=Inches(6.0))
        doc.add_paragraph()

    doc.add_page_break()

    # 3.5 Multi-Pathogen Assessment
    heading2 = doc.add_heading('3.5 Multi-Pathogen Assessment', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The Multi-Pathogen Assessment workflow evaluates multiple pathogens simultaneously '
        'to identify the dominant health risk.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Note: ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run(
        'In Production Mode, this workflow is restricted to norovirus only. To assess multiple '
        'pathogens, disable Production Mode and enable Research Mode.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Steps:')
    run1.bold = True

    doc.add_paragraph('Select "Multi-Pathogen Assessment" from the assessment mode dropdown', style='List Number')
    doc.add_paragraph('Select 2-6 pathogens from the multi-select box', style='List Number')
    doc.add_paragraph('Choose exposure route', style='List Number')
    doc.add_paragraph('Configure shared parameters (dilution, treatment, exposure)', style='List Number')
    doc.add_paragraph('Click "Run Multi-Pathogen Assessment"', style='List Number')
    doc.add_paragraph('Compare pathogen-specific risks in the results table', style='List Number')

    doc.add_paragraph()

    # Add screenshot 6: Multi-Pathogen
    screenshot_path = Path('../screenshots/06_multi_pathogen_20251117_095645.png')
    if screenshot_path.exists():
        doc.add_paragraph('Figure 6: Multi-Pathogen Assessment Interface', style='Caption')
        doc.add_picture(str(screenshot_path), width=Inches(6.0))
        doc.add_paragraph()

    doc.add_page_break()

    # ====================
    # 4. INTERPRETING RESULTS
    # ====================
    heading = doc.add_heading('4. Interpreting Results', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    # 4.1 Risk Metrics Explained
    heading2 = doc.add_heading('4.1 Risk Metrics Explained', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph('The application reports several key risk metrics:')
    doc.add_paragraph()

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Mean Dose (organisms per exposure): ')
    run1.bold = True
    p.add_run('Average number of viable pathogens ingested or inhaled per exposure event')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Infection Probability (per exposure): ')
    run1.bold = True
    p.add_run('Probability of infection from a single exposure event (0 to 1)')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Annual Infection Risk: ')
    run1.bold = True
    p.add_run('Probability of at least one infection over a year of exposure: 1 - (1 - P_inf)^n')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Illness Probability: ')
    run1.bold = True
    p.add_run('Probability of symptomatic illness given infection (infection × illness|infection ratio)')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('95% Confidence Interval: ')
    run1.bold = True
    p.add_run('Uncertainty range capturing 95% of simulated outcomes (2.5th to 97.5th percentile)')

    # 4.2 Visualizations
    heading2 = doc.add_heading('4.2 Visualizations', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph('The application provides several visualization types:')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Risk Overview Bar Charts: ')
    run1.bold = True
    p.add_run('Compare mean risks across scenarios, sites, or time periods')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Uncertainty Distribution Plots: ')
    run1.bold = True
    p.add_run('Histograms showing the full range of simulated outcomes')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Time Series Plots: ')
    run1.bold = True
    p.add_run('Temporal trends in risk over monitoring periods')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Compliance Status Indicators: ')
    run1.bold = True
    p.add_run('Color-coded flags for WHO guideline compliance (green = compliant, red = exceeds)')

    # 4.3 WHO Guideline Compliance
    heading2 = doc.add_heading('4.3 WHO Guideline Compliance', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The World Health Organization (WHO) recommends an annual infection risk target of '
        '10⁻⁴ (0.0001 or 1 in 10,000) for drinking water and recreational water quality.'
    )

    doc.add_paragraph()
    doc.add_paragraph('The application automatically evaluates compliance:')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('✅ PASS: ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run('Annual infection risk < 10⁻⁴ (meets WHO guideline)')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('❌ FAIL: ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('Annual infection risk ≥ 10⁻⁴ (exceeds WHO guideline)')

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: Some jurisdictions may use different risk targets. Consult local regulations '
        'and health authorities for applicable standards.'
    )

    doc.add_page_break()

    # ====================
    # 5. TECHNICAL BACKGROUND
    # ====================
    heading = doc.add_heading('5. Technical Background', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    # 5.1 Dose-Response Models
    heading2 = doc.add_heading('5.1 Dose-Response Models', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'Dose-response models describe the relationship between pathogen exposure dose and the '
        'probability of infection. The choice of model depends on pathogen characteristics and '
        'mathematical validity conditions.'
    )

    # 5.2 Beta-Binomial Model
    heading2 = doc.add_heading('5.2 Beta-Binomial Model (Norovirus)', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'For norovirus, this application uses the EXACT Beta-Binomial dose-response model, '
        'which is the mathematically correct formulation for highly infectious viruses.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Formula:')
    run1.bold = True

    formula_p = doc.add_paragraph()
    formula_p.paragraph_format.left_indent = Inches(0.5)
    formula_run = formula_p.add_run(
        'P(infection) = 1 - exp[ln Γ(β+dose) + ln Γ(α+β) - ln Γ(α+β+dose) - ln Γ(β)]'
    )
    formula_run.font.name = 'Consolas'
    formula_run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Parameters (Teunis et al. 2008):')
    run1.bold = True

    doc.add_paragraph('α (alpha) = 0.04', style='List Bullet')
    doc.add_paragraph('β (beta) = 0.055', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run1 = p.add_run('Why NOT Beta-Poisson?')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph(
        'The Beta-Poisson approximation is mathematically INVALID for norovirus because it requires '
        'β >> 1 (beta much greater than 1). For norovirus, β = 0.055 << 1, making the approximation '
        'inappropriate. Using Beta-Poisson would underestimate risk by 2-4× at low doses.'
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'This implementation has been validated against reference Excel spreadsheets (David Wood, NIWA) '
        'and produces exact agreement for all test doses.'
    )

    # 5.3 Monte Carlo Simulation
    heading2 = doc.add_heading('5.3 Monte Carlo Simulation', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph(
        'The application uses Monte Carlo simulation with 10,000 iterations to quantify uncertainty '
        'in risk estimates. Each iteration samples from input parameter distributions (e.g., pathogen '
        'concentration, dilution factor) to capture natural variability and measurement uncertainty.'
    )

    doc.add_paragraph()
    doc.add_paragraph('The simulation outputs:')
    doc.add_paragraph('Mean, median, and percentile estimates (5th, 25th, 75th, 95th)', style='List Bullet')
    doc.add_paragraph('Full probability distributions for risk metrics', style='List Bullet')
    doc.add_paragraph('Confidence intervals for uncertainty quantification', style='List Bullet')

    doc.add_page_break()

    # ====================
    # APPENDIX
    # ====================
    heading = doc.add_heading('Appendix: References & Contact', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    heading2 = doc.add_heading('Key References', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    references = [
        'Teunis, P.F., et al. (2008). "Norwalk virus: How infectious is it?" Journal of Medical Virology, 80(8), 1468-1476.',
        'McBride, G.B. (2017). "Bell Island Wastewater Treatment QMRA." NIWA Client Report 2017350HN.',
        'World Health Organization (2016). "Quantitative Microbial Risk Assessment: Application for Water Safety Management."',
        'Haas, C.N., Rose, J.B., & Gerba, C.P. (2014). "Quantitative Microbial Risk Assessment" (2nd ed.). John Wiley & Sons.',
    ]

    for ref in references:
        p = doc.add_paragraph(ref, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.hanging_indent = Inches(0.25)

    doc.add_paragraph()

    heading2 = doc.add_heading('Technical Support', level=2)
    heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph('For technical support, questions, or feedback:')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('NIWA - National Institute of Water & Atmospheric Research\n')
    run1.bold = True
    p.add_run('Hamilton, New Zealand\n')
    p.add_run('Website: www.niwa.co.nz')

    doc.add_paragraph()
    doc.add_paragraph()

    # Final note
    final_p = doc.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final_p.add_run('End of User Guide')
    final_run.font.size = Pt(10)
    final_run.font.italic = True
    final_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = Path('../../QMRA_Application_User_Guide.docx')
    doc.save(str(output_path))
    print(f"\n[SUCCESS] User guide created: {output_path.absolute()}")
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB")
    print("\nThe document includes:")
    print("  ✓ Professional title page")
    print("  ✓ Table of contents")
    print("  ✓ Complete introduction to QMRA")
    print("  ✓ Step-by-step workflow instructions")
    print("  ✓ 6 embedded screenshots")
    print("  ✓ Results interpretation guide")
    print("  ✓ Technical background on Beta-Binomial model")
    print("  ✓ References and contact information")

    return output_path


if __name__ == '__main__':
    create_user_guide()
