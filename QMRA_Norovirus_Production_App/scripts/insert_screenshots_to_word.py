#!/usr/bin/env python3
"""
Insert Screenshots into QMRA Word Document
===========================================

Inserts captured app screenshots into the Word report.
Creates a new "App Screenshots" section with all images.

Usage:
    python insert_screenshots_to_word.py
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def insert_screenshots_to_word(
    word_doc_path='QMRA_Application_Overview.docx',
    screenshots_dir='screenshots',
    manifest_path='screenshots/manifest.json'
):
    """Insert screenshots into Word document."""

    # Load manifest
    manifest_file = Path(manifest_path)
    if not manifest_file.exists():
        print(f"[ERROR] Manifest not found: {manifest_path}")
        print("First run: python capture_app_screenshots.py")
        return False

    with open(manifest_file, 'r') as f:
        manifest = json.load(f)

    screenshots = manifest.get('screenshots', [])
    if not screenshots:
        print("[WARN] No screenshots found in manifest")
        return False

    print(f"\n[LOAD] Found {len(screenshots)} screenshots in manifest")

    # Open Word document
    doc = Document(word_doc_path)
    print(f"[LOAD] Opened Word document: {word_doc_path}")

    # Add new page
    doc.add_page_break()

    # Add title
    heading = doc.add_heading('Appendix: Application Screenshots', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

    doc.add_paragraph(
        f'This appendix contains {len(screenshots)} screenshots from the QMRA Batch Processing '
        f'Web Application, demonstrating key interface pages and features.'
    )

    doc.add_page_break()

    # Add each screenshot
    screenshots_dir_path = Path(screenshots_dir)

    for i, screenshot in enumerate(screenshots, 1):
        filename = screenshot.get('filename', '')
        page = screenshot.get('page', '').replace('_', ' ').title()
        description = screenshot.get('description', '')

        # Screenshot heading
        heading = doc.add_heading(f'{i}. {page}', level=2)
        heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)

        # Description
        if description:
            doc.add_paragraph(description, style='Normal')

        # Try to insert image
        image_path = screenshots_dir_path / filename
        if image_path.exists():
            try:
                # Add image with caption
                doc.add_paragraph()  # Spacing
                doc.add_picture(str(image_path), width=Inches(6.0))

                # Add caption
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                caption = doc.add_paragraph(f'Figure {i}: {page}')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.runs[0]
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(100, 100, 100)

                print(f"[OK] Inserted image {i}/{len(screenshots)}: {page}")

            except Exception as e:
                print(f"[WARN] Failed to insert image: {filename} - {e}")
                doc.add_paragraph(f"[Image could not be inserted: {filename}]")
        else:
            print(f"[WARN] Image file not found: {image_path}")
            doc.add_paragraph(f"[Image file not found: {filename}]")

        doc.add_paragraph()  # Spacing between screenshots

    # Save updated document
    output_path = word_doc_path.replace('.docx', '_with_screenshots.docx')
    doc.save(output_path)

    print(f"\n[OK] Document saved: {output_path}")
    print(f"[OK] {len(screenshots)} screenshots inserted")

    return True


def main():
    """Main entry point."""
    print("\n" + "="*80)
    print("INSERT SCREENSHOTS INTO WORD DOCUMENT")
    print("="*80)

    # Check if screenshots exist
    manifest_path = Path('screenshots/manifest.json')
    if not manifest_path.exists():
        print("\n[ERROR] No screenshots found!")
        print("\nFirst, capture screenshots with:")
        print("  1. Start app: streamlit run web_app.py")
        print("  2. In another terminal: python capture_app_screenshots.py")
        print("  3. Then run: python insert_screenshots_to_word.py")
        return

    # Insert screenshots
    success = insert_screenshots_to_word(
        word_doc_path='QMRA_Application_Overview.docx',
        screenshots_dir='screenshots',
        manifest_path='screenshots/manifest.json'
    )

    if success:
        print("\n" + "="*80)
        print("SUCCESS: Document updated with screenshots!")
        print("="*80)
        print("\nNew file: QMRA_Application_Overview_with_screenshots.docx")
        print("\nYou can now:")
        print("  1. Open the updated Word document")
        print("  2. Review the new 'App Screenshots' appendix")
        print("  3. Add captions or rearrange as needed")
        print("  4. Share the complete report!")
    else:
        print("\n[FAILED] Could not insert screenshots")


if __name__ == '__main__':
    main()
