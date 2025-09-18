#!/usr/bin/env python3
"""
Generate QMRA Research Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_qmra_document():
    """Create a comprehensive QMRA research document."""

    # Create new document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "Quantitative Microbial Risk Assessment (QMRA) Research Document"
    doc.core_properties.author = "NIWA Research Team"
    doc.core_properties.subject = "QMRA Research and Methodology"

    # Add title
    title = doc.add_heading('Quantitative Microbial Risk Assessment (QMRA)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.add_run(f'Research Document - {datetime.now().strftime("%B %Y")}').bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'This document provides a comprehensive overview of Quantitative Microbial Risk Assessment (QMRA) '
        'methodology, its applications in New Zealand context, and current research developments. '
        'QMRA is a systematic approach to estimating the risk of infection, illness, or death resulting '
        'from exposure to pathogenic microorganisms in water, food, and environmental sources.'
    )
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Introduction to QMRA',
        '2. QMRA Framework and Methodology',
        '3. New Zealand Regulatory Context',
        '4. Water Quality Risk Assessment',
        '5. Food Safety Applications',
        '6. Environmental Health Considerations',
        '7. Case Studies',
        '8. Future Research Directions',
        '9. References and Resources'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # Section 1: Introduction
    doc.add_heading('1. Introduction to QMRA', 1)
    doc.add_paragraph(
        'Quantitative Microbial Risk Assessment (QMRA) is a scientific framework used to estimate '
        'the risk of adverse health effects from exposure to pathogenic microorganisms. '
        'It integrates information from various disciplines including microbiology, epidemiology, '
        'environmental science, and public health.'
    )

    doc.add_heading('1.1 Key Components', 2)
    components = [
        'Hazard Identification: Identifying pathogenic microorganisms of concern',
        'Exposure Assessment: Estimating the magnitude, frequency, and duration of exposure',
        'Dose-Response Assessment: Characterizing the relationship between dose and health effects',
        'Risk Characterization: Integrating exposure and dose-response to estimate risk'
    ]
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')
    doc.add_paragraph()

    # Section 2: Framework and Methodology
    doc.add_heading('2. QMRA Framework and Methodology', 1)
    doc.add_paragraph(
        'The QMRA framework follows a systematic approach aligned with international guidelines '
        'including WHO Water Safety Plans and Codex Alimentarius food safety principles.'
    )

    doc.add_heading('2.1 Mathematical Models', 2)
    doc.add_paragraph(
        'QMRA employs various mathematical models to quantify risk:'
    )
    models = [
        'Exponential model: P(inf) = 1 - exp(-r × dose)',
        'Beta-Poisson model: P(inf) = 1 - (1 + dose/β)^(-α)',
        'Monte Carlo simulation for uncertainty analysis',
        'Bayesian approaches for parameter estimation'
    ]
    for model in models:
        doc.add_paragraph(model, style='List Bullet')
    doc.add_paragraph()

    # Section 3: NZ Regulatory Context
    doc.add_heading('3. New Zealand Regulatory Context', 1)
    doc.add_paragraph(
        'New Zealand has established comprehensive regulatory frameworks for microbial risk management '
        'in water and food systems.'
    )

    doc.add_heading('3.1 Key Regulatory Bodies', 2)
    bodies = [
        'Taumata Arowai - Water Services Regulator',
        'Ministry of Health - Public Health Guidelines',
        'Ministry for Primary Industries (MPI) - Food Safety',
        'Regional Councils - Environmental Management'
    ]
    for body in bodies:
        doc.add_paragraph(body, style='List Bullet')

    doc.add_heading('3.2 Drinking Water Standards', 2)
    doc.add_paragraph(
        'The Drinking Water Standards for New Zealand (DWSNZ) specify maximum acceptable values '
        'for microbiological determinands including:'
    )
    standards = [
        'E. coli: <1 per 100 mL',
        'Total coliforms: Monitored as process indicator',
        'Protozoa: 3-log removal/inactivation requirement',
        'Viruses: 4-log removal/inactivation requirement'
    ]
    for standard in standards:
        doc.add_paragraph(standard, style='List Bullet')
    doc.add_paragraph()

    # Section 4: Water Quality Risk Assessment
    doc.add_heading('4. Water Quality Risk Assessment', 1)
    doc.add_paragraph(
        'Water quality risk assessment is critical for ensuring safe drinking water supply '
        'and protecting public health from waterborne pathogens.'
    )

    doc.add_heading('4.1 Priority Pathogens', 2)
    pathogens = [
        'Bacteria: Campylobacter, Salmonella, pathogenic E. coli, Legionella',
        'Viruses: Norovirus, Rotavirus, Adenovirus, Hepatitis A',
        'Protozoa: Cryptosporidium, Giardia, Naegleria fowleri',
        'Cyanobacteria and associated toxins'
    ]
    for pathogen in pathogens:
        doc.add_paragraph(pathogen, style='List Bullet')

    doc.add_heading('4.2 Source Water Protection', 2)
    doc.add_paragraph(
        'Implementing multiple barriers from catchment to consumer:'
    )
    barriers = [
        'Catchment management and source protection zones',
        'Treatment processes (coagulation, filtration, disinfection)',
        'Distribution system integrity',
        'Monitoring and verification programs'
    ]
    for barrier in barriers:
        doc.add_paragraph(barrier, style='List Bullet')
    doc.add_paragraph()

    # Section 5: Food Safety Applications
    doc.add_heading('5. Food Safety Applications', 1)
    doc.add_paragraph(
        'QMRA is extensively used in food safety to assess risks throughout the food chain '
        'from production to consumption.'
    )

    doc.add_heading('5.1 Food Production Systems', 2)
    systems = [
        'Primary production (farms, aquaculture)',
        'Processing and manufacturing',
        'Distribution and retail',
        'Food service and consumer handling'
    ]
    for system in systems:
        doc.add_paragraph(system, style='List Bullet')

    doc.add_heading('5.2 Critical Control Points', 2)
    doc.add_paragraph(
        'Hazard Analysis and Critical Control Points (HACCP) integration with QMRA:'
    )
    ccps = [
        'Temperature control (cooking, cooling, storage)',
        'pH and water activity management',
        'Cross-contamination prevention',
        'Time limits for pathogen growth'
    ]
    for ccp in ccps:
        doc.add_paragraph(ccp, style='List Bullet')
    doc.add_paragraph()

    # Section 6: Environmental Health
    doc.add_heading('6. Environmental Health Considerations', 1)
    doc.add_paragraph(
        'Environmental factors significantly influence microbial risk in both natural '
        'and built environments.'
    )

    doc.add_heading('6.1 Climate Change Impacts', 2)
    impacts = [
        'Increased water temperature affecting pathogen survival',
        'Extreme weather events and flooding risks',
        'Changes in precipitation patterns affecting water sources',
        'Sea level rise and saltwater intrusion'
    ]
    for impact in impacts:
        doc.add_paragraph(impact, style='List Bullet')

    doc.add_heading('6.2 Recreational Water Quality', 2)
    doc.add_paragraph(
        'Managing microbial risks in recreational waters including beaches, '
        'rivers, and lakes through monitoring and public advisories.'
    )
    doc.add_paragraph()

    # Section 7: Case Studies
    doc.add_heading('7. Case Studies', 1)

    doc.add_heading('7.1 Havelock North Water Crisis (2016)', 2)
    doc.add_paragraph(
        'The Havelock North campylobacteriosis outbreak affected over 5,000 people '
        'and highlighted the importance of source water protection, treatment barriers, '
        'and proactive risk assessment. Key lessons learned include the need for '
        'mandatory treatment of all drinking water supplies and improved regulatory oversight.'
    )

    doc.add_heading('7.2 New Zealand Shellfish QMRA', 2)
    doc.add_paragraph(
        'Development of risk-based management strategies for shellfish harvesting areas '
        'considering norovirus and bacterial contamination from wastewater discharges '
        'and agricultural runoff.'
    )
    doc.add_paragraph()

    # Section 8: Future Research
    doc.add_heading('8. Future Research Directions', 1)
    directions = [
        'Integration of genomic techniques for pathogen detection and characterization',
        'Machine learning applications in risk prediction',
        'Climate change adaptation strategies',
        'One Health approaches linking human, animal, and environmental health',
        'Real-time monitoring and early warning systems',
        'Community-specific risk assessment tools'
    ]
    for direction in directions:
        doc.add_paragraph(direction, style='List Bullet')
    doc.add_paragraph()

    # Section 9: References
    doc.add_heading('9. References and Resources', 1)

    doc.add_heading('Key Publications', 2)
    pubs = [
        'WHO (2016) Quantitative Microbial Risk Assessment: Application for Water Safety Management',
        'Haas, C.N., Rose, J.B., & Gerba, C.P. (2014) Quantitative Microbial Risk Assessment, 2nd Edition',
        'Taumata Arowai (2022) Drinking Water Quality Assurance Rules',
        'Ministry of Health (2018) Drinking Water Standards for New Zealand',
        'MPI (2021) Microbiological Reference Criteria for Food'
    ]
    for pub in pubs:
        doc.add_paragraph(pub, style='List Bullet')

    doc.add_heading('Useful Websites', 2)
    sites = [
        'Taumata Arowai: www.taumataarowai.govt.nz',
        'Ministry of Health: www.health.govt.nz',
        'MPI Food Safety: www.mpi.govt.nz/food-safety',
        'ESR Public Health: www.esr.cri.nz',
        'Water New Zealand: www.waternz.org.nz'
    ]
    for site in sites:
        doc.add_paragraph(site, style='List Bullet')

    # Add appendices
    doc.add_page_break()
    doc.add_heading('Appendix A: Glossary of Terms', 1)
    terms = {
        'QMRA': 'Quantitative Microbial Risk Assessment',
        'DALYs': 'Disability Adjusted Life Years',
        'Log Reduction': 'Logarithmic reduction in pathogen concentration',
        'MPN': 'Most Probable Number',
        'CFU': 'Colony Forming Units',
        'HACCP': 'Hazard Analysis and Critical Control Points',
        'WSP': 'Water Safety Plan'
    }
    for term, definition in terms.items():
        para = doc.add_paragraph()
        para.add_run(f'{term}: ').bold = True
        para.add_run(definition)

    doc.add_page_break()
    doc.add_heading('Appendix B: Common Dose-Response Parameters', 1)

    # Add a table for dose-response parameters
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Pathogen'
    header_cells[1].text = 'Model'
    header_cells[2].text = 'Parameters'
    header_cells[3].text = 'Reference'

    pathogens_data = [
        ('Campylobacter jejuni', 'Beta-Poisson', 'α=0.145, β=7.59', 'Teunis et al. 2005'),
        ('Cryptosporidium parvum', 'Exponential', 'r=0.0042', 'Haas et al. 1996'),
        ('Rotavirus', 'Beta-Poisson', 'α=0.253, β=0.42', 'Regli et al. 1991'),
        ('E. coli O157:H7', 'Beta-Poisson', 'α=0.49, β=1.81×10^5', 'Strachan et al. 2005')
    ]

    for pathogen, model, params, ref in pathogens_data:
        row_cells = table.add_row().cells
        row_cells[0].text = pathogen
        row_cells[1].text = model
        row_cells[2].text = params
        row_cells[3].text = ref

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f'Document generated on {datetime.now().strftime("%d %B %Y")}').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    filename = f'QMRA_Research_Document_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f'Document created: {filename}')
    return filename

if __name__ == '__main__':
    create_qmra_document()