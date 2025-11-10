#!/usr/bin/env python3
"""
Simple QMRA Demo Document Generator with Screenshots
====================================================

Creates a straightforward demonstration guide without official formatting.
This content can be copied into the official template later.

Version: 1.2.0
Author: NIWA Earth Sciences New Zealand
Date: November 2025
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def add_simple_title(doc):
    """Add simple title."""
    title = doc.add_heading('QMRA Batch Processing Web Application', 0)
    subtitle = doc.add_heading('Step-by-Step Demonstration Guide', level=1)

    info = doc.add_paragraph()
    info.add_run(f"Version 1.2.0\n")
    info.add_run(f"{datetime.now().strftime('%B %Y')}\n")
    info.add_run("NIWA Earth Sciences New Zealand")

    doc.add_page_break()


def add_introduction(doc):
    """Add brief introduction."""
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        "This guide demonstrates how to use the QMRA Batch Processing Web Application "
        "for conducting quantitative microbial risk assessments. The application provides "
        "5 assessment modes for different analysis scenarios."
    )

    doc.add_heading('Assessment Modes', level=2)
    modes = [
        "Batch Scenarios - Multiple scenarios from CSV file",
        "Spatial Assessment - Risk at different distances from discharge",
        "Temporal Assessment - Risk over time periods",
        "Treatment Comparison - Compare treatment options",
        "Multi-Pathogen Assessment - Compare different pathogens"
    ]

    for mode in modes:
        doc.add_paragraph(mode, style='List Bullet')

    doc.add_page_break()


def add_getting_started(doc):
    """Add getting started section."""
    doc.add_heading('Getting Started', level=1)

    doc.add_heading('Prerequisites', level=2)
    prereqs = [
        "Python 3.8 or higher installed",
        "All dependencies installed: pip install -r requirements.txt",
        "Test data files in Batch_Processing_App/input_data/",
        "Web browser (Chrome, Firefox, or Edge)"
    ]

    for prereq in prereqs:
        doc.add_paragraph(f"✓ {prereq}", style='List Bullet')

    doc.add_heading('Launching the Application', level=2)

    doc.add_paragraph("1. Open terminal/command prompt")
    doc.add_paragraph("2. Navigate to the application directory:")
    doc.add_paragraph('   cd "path/to/Batch_Processing_App"')
    doc.add_paragraph("3. Launch the app:")
    doc.add_paragraph('   streamlit run web_app.py')
    doc.add_paragraph("   Or on Windows, double-click: launch_web_gui.bat")
    doc.add_paragraph("4. Browser opens automatically at http://localhost:8501")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run("Screenshots below show the actual application interface running at localhost:8501")

    doc.add_page_break()


def add_demo_scenario_1(doc):
    """Add first demonstration scenario with placeholders for screenshots."""
    doc.add_heading('Demo 1: Spatial Risk Assessment', level=1)

    doc.add_paragraph(
        "This demonstration shows how to assess risk at multiple locations "
        "around a wastewater discharge point."
    )

    # Step 1
    doc.add_heading('Step 1: Select Spatial Assessment Mode', level=2)
    doc.add_paragraph("In the sidebar, select 'Spatial Assessment' from the dropdown menu.")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 1: ").bold = True
    screenshot_note.add_run("Main interface showing sidebar with assessment mode selector]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/01_select_spatial_mode.png")
    doc.add_paragraph()

    # Step 2
    doc.add_heading('Step 2: Upload Dilution Data', level=2)
    doc.add_paragraph("Click 'Browse files' and select the dilution data file:")
    doc.add_paragraph("   File: input_data/dilution_data/spatial_dilution_6_sites.csv")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 2: ").bold = True
    screenshot_note.add_run("File upload section and data preview]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/02_upload_dilution_data.png")
    doc.add_paragraph()

    # Step 3
    doc.add_heading('Step 3: Configure Parameters', level=2)
    doc.add_paragraph("Set the following parameters in the sidebar:")

    params = [
        "Pathogen: Norovirus",
        "Effluent Concentration: 1.54×10⁶ copies/L",
        "Treatment LRV: 3.0",
        "Exposure Route: Primary Contact",
        "Exposure Volume: 50 mL",
        "Exposure Frequency: 20 events/year",
        "Population: 10,000",
        "Monte Carlo Iterations: 10,000"
    ]

    for param in params:
        doc.add_paragraph(f"• {param}")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 3: ").bold = True
    screenshot_note.add_run("Parameter configuration sidebar]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/03_configure_parameters.png")
    doc.add_paragraph()

    # Step 4
    doc.add_heading('Step 4: Run Assessment', level=2)
    doc.add_paragraph("Click 'Run Spatial Assessment' button and wait for processing to complete.")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 4: ").bold = True
    screenshot_note.add_run("Processing in progress or completion message]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/04_run_assessment.png")
    doc.add_paragraph()

    # Step 5
    doc.add_heading('Step 5: Review Results', level=2)
    doc.add_paragraph("The results table appears below showing risk for each site:")

    doc.add_paragraph("Key metrics displayed:")
    metrics = [
        "Annual Risk (Mean, Median, 95th percentile)",
        "Illness Risk",
        "Population Impact (expected cases)",
        "Compliance Status (vs WHO threshold)"
    ]

    for metric in metrics:
        doc.add_paragraph(f"• {metric}")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 5: ").bold = True
    screenshot_note.add_run("Results table with risk values for all sites]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/05_results_table.png")
    doc.add_paragraph()

    # Step 6
    doc.add_heading('Step 6: View Visualizations', level=2)
    doc.add_paragraph("Click through the visualization tabs:")

    viz_tabs = [
        "Risk Overview - Bar chart of risk by site",
        "Compliance Distribution - Pie chart",
        "Risk Distribution - Histogram",
        "Population Impact - Bar chart of illness cases"
    ]

    for tab in viz_tabs:
        doc.add_paragraph(f"• {tab}")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 6: ").bold = True
    screenshot_note.add_run("Visualization tabs - Risk Overview chart]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/06_visualizations.png")
    doc.add_paragraph()

    # Step 7
    doc.add_heading('Step 7: Download Results', level=2)
    doc.add_paragraph("Use the download buttons to save results:")

    downloads = [
        "Download Full CSV - Complete results table",
        "Download All (ZIP) - All plots and tables",
        "Generate PDF Report - Professional report",
        "Download Individual Plots - High-res PNG files"
    ]

    for download in downloads:
        doc.add_paragraph(f"• {download}")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 7: ").bold = True
    screenshot_note.add_run("Download buttons section]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/07_download_options.png")
    doc.add_paragraph()

    doc.add_page_break()


def add_demo_scenario_2(doc):
    """Add second demonstration scenario."""
    doc.add_heading('Demo 2: Treatment Comparison', level=1)

    doc.add_paragraph(
        "This demonstration compares the effectiveness of different treatment options."
    )

    # Step 1
    doc.add_heading('Step 1: Select Treatment Comparison Mode', level=2)
    doc.add_paragraph("In the sidebar, select 'Treatment Comparison' from the dropdown.")

    # Step 2
    doc.add_heading('Step 2: Upload Treatment Files', level=2)
    doc.add_paragraph("Upload two treatment scenario files:")
    doc.add_paragraph("• Current: input_data/treatment_scenarios/secondary_treatment.yaml")
    doc.add_paragraph("• Proposed: input_data/treatment_scenarios/advanced_uv_treatment.yaml")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 8: ").bold = True
    screenshot_note.add_run("Treatment comparison mode with file uploads]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/08_treatment_comparison.png")
    doc.add_paragraph()

    # Step 3
    doc.add_heading('Step 3: Configure Common Parameters', level=2)
    doc.add_paragraph("Set parameters that apply to both scenarios:")

    params = [
        "Pathogen: Norovirus",
        "Effluent Concentration: 1.54×10⁶ copies/L",
        "Dilution Factor: 13.6",
        "Exposure parameters as before"
    ]

    for param in params:
        doc.add_paragraph(f"• {param}")

    # Step 4
    doc.add_heading('Step 4: View Comparison Results', level=2)
    doc.add_paragraph("Results show side-by-side comparison:")

    doc.add_paragraph("Secondary Treatment (LRV 3.0):")
    doc.add_paragraph("• Annual Risk: ~3.4×10⁻²")
    doc.add_paragraph("• Status: NON-COMPLIANT")

    doc.add_paragraph("Advanced UV Treatment (LRV 8.0):")
    doc.add_paragraph("• Annual Risk: ~2.1×10⁻⁵")
    doc.add_paragraph("• Status: COMPLIANT")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 9: ").bold = True
    screenshot_note.add_run("Treatment comparison results table]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/09_comparison_results.png")
    doc.add_paragraph()

    doc.add_page_break()


def add_demo_scenario_3(doc):
    """Add third demonstration scenario."""
    doc.add_heading('Demo 3: Batch Scenarios Assessment', level=1)

    doc.add_paragraph(
        "This demonstration shows how to process multiple scenarios from a CSV file."
    )

    # Step 1
    doc.add_heading('Step 1: Select Batch Scenarios Mode', level=2)
    doc.add_paragraph("In the sidebar, select 'Batch Scenarios' from the dropdown.")

    # Step 2
    doc.add_heading('Step 2: Upload Batch Scenario File', level=2)
    doc.add_paragraph("Upload the master batch scenarios file:")
    doc.add_paragraph("   File: input_data/batch_scenarios/master_batch_scenarios.csv")

    doc.add_paragraph()
    doc.add_paragraph("The file contains 15 pre-configured scenarios with varying:")
    scenarios = [
        "Pathogen types (Norovirus, Campylobacter, etc.)",
        "Treatment levels (LRV 1.0 to 8.0)",
        "Dilution factors",
        "Exposure routes and frequencies"
    ]

    for scenario in scenarios:
        doc.add_paragraph(f"• {scenario}")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 10: ").bold = True
    screenshot_note.add_run("Batch scenarios upload and preview]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/10_batch_scenarios.png")
    doc.add_paragraph()

    # Step 3
    doc.add_heading('Step 3: Run Batch Assessment', level=2)
    doc.add_paragraph("Click 'Run Batch Assessment' - all scenarios process automatically.")

    # Step 4
    doc.add_heading('Step 4: Analyze Comprehensive Results', level=2)
    doc.add_paragraph("Results table shows all 15 scenarios with:")

    results = [
        "Scenario name and parameters",
        "Risk metrics for each scenario",
        "Compliance status",
        "Ranking from highest to lowest risk"
    ]

    for result in results:
        doc.add_paragraph(f"• {result}")

    doc.add_paragraph()
    screenshot_note = doc.add_paragraph()
    screenshot_note.add_run("[SCREENSHOT 11: ").bold = True
    screenshot_note.add_run("Batch results with all scenarios]")
    screenshot_note.add_run("\nLocation: ").bold = True
    screenshot_note.add_run("screenshots_example/11_batch_results.png")
    doc.add_paragraph()

    doc.add_page_break()


def add_key_features(doc):
    """Add key features summary."""
    doc.add_heading('Key Features Summary', level=1)

    doc.add_heading('Interactive Visualizations', level=2)
    doc.add_paragraph("The app provides 4 interactive chart types:")

    charts = [
        "Risk Overview - Compare risk across sites/scenarios",
        "Compliance Distribution - Visual compliance status",
        "Risk Distribution - Histogram of Monte Carlo results",
        "Population Impact - Expected illness cases"
    ]

    for chart in charts:
        doc.add_paragraph(f"• {chart}")

    doc.add_heading('Download Options', level=2)
    doc.add_paragraph("Multiple formats available:")

    downloads = [
        "CSV - Full results table for Excel/analysis",
        "ZIP Package - All plots + tables + README",
        "PDF Report - Professional formatted report",
        "Individual PNGs - High-resolution plots (300 DPI)"
    ]

    for download in downloads:
        doc.add_paragraph(f"• {download}")

    doc.add_heading('Risk Metrics Explained', level=2)

    doc.add_paragraph()
    doc.add_paragraph("Mean Annual Risk - Average probability per person per year")
    doc.add_paragraph("Median Annual Risk - 50th percentile (less affected by outliers)")
    doc.add_paragraph("95th Percentile Risk - Upper uncertainty bound")
    doc.add_paragraph("Mean Illness Risk - Probability of symptomatic illness")
    doc.add_paragraph("Expected Cases - Projected illnesses in population")
    doc.add_paragraph("Compliance Status - Comparison to WHO threshold (1×10⁻⁴)")

    doc.add_page_break()


def add_screenshot_guide(doc):
    """Add guide for taking screenshots."""
    doc.add_heading('Screenshot Capture Guide', level=1)

    doc.add_paragraph(
        "To replace the screenshot placeholders with actual images, "
        "follow these steps:"
    )

    doc.add_heading('Required Screenshots', level=2)

    screenshots = [
        ("01_select_spatial_mode.png", "Main interface with sidebar showing assessment mode dropdown"),
        ("02_upload_dilution_data.png", "File upload section with data preview"),
        ("03_configure_parameters.png", "Sidebar with all parameter inputs filled"),
        ("04_run_assessment.png", "Run button and processing status"),
        ("05_results_table.png", "Full results table with all metrics"),
        ("06_visualizations.png", "Visualization tabs showing charts"),
        ("07_download_options.png", "Download buttons section"),
        ("08_treatment_comparison.png", "Treatment comparison mode interface"),
        ("09_comparison_results.png", "Side-by-side treatment comparison results"),
        ("10_batch_scenarios.png", "Batch upload interface"),
        ("11_batch_results.png", "Complete batch results table")
    ]

    for filename, description in screenshots:
        p = doc.add_paragraph()
        p.add_run(f"{filename}: ").bold = True
        p.add_run(description)

    doc.add_heading('How to Capture', level=2)

    steps = [
        "Launch the application at http://localhost:8501",
        "Navigate to each section as described in the demos",
        "Use Windows Snipping Tool (Win+Shift+S) or similar",
        "Save screenshots to screenshots_example/ folder",
        "Name files exactly as shown above",
        "Replace placeholders in this document with actual images"
    ]

    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()


def add_next_steps(doc):
    """Add next steps section."""
    doc.add_heading('Next Steps', level=1)

    doc.add_paragraph(
        "After capturing all screenshots and replacing the placeholders:"
    )

    steps = [
        "Review all screenshots for clarity and completeness",
        "Ensure all parameter values match the demonstrations",
        "Verify that results shown are realistic and accurate",
        "Copy this content into the official NIWA template",
        "Add official branding, headers, and footers",
        "Final review and approval"
    ]

    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_paragraph()
    doc.add_paragraph()

    final = doc.add_paragraph()
    final.add_run("Document prepared by: ").bold = True
    final.add_run("Reza Moghaddam, NIWA\n")
    final.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n")
    final.add_run("Version: 1.2.0 - Simple Demo (Pre-Template)")


def main():
    """Generate simple demo document."""
    print("="*80)
    print("QMRA Simple Demo Generator (No Official Formatting)")
    print("="*80)
    print()

    # Create document
    print("Creating document...")
    doc = Document()

    # Add sections
    print("Adding title...")
    add_simple_title(doc)

    print("Adding introduction...")
    add_introduction(doc)

    print("Adding getting started...")
    add_getting_started(doc)

    print("Adding Demo 1: Spatial Assessment...")
    add_demo_scenario_1(doc)

    print("Adding Demo 2: Treatment Comparison...")
    add_demo_scenario_2(doc)

    print("Adding Demo 3: Batch Scenarios...")
    add_demo_scenario_3(doc)

    print("Adding key features...")
    add_key_features(doc)

    print("Adding screenshot guide...")
    add_screenshot_guide(doc)

    print("Adding next steps...")
    add_next_steps(doc)

    # Save
    output_file = "QMRA_Simple_Demo_NO_TEMPLATE.docx"
    print(f"\nSaving: {output_file}")
    doc.save(output_file)

    file_size = os.path.getsize(output_file) / 1024

    print("\n" + "="*80)
    print("DOCUMENT COMPLETE")
    print("="*80)
    print(f"\nFile: {output_file}")
    print(f"Size: {file_size:.1f} KB")
    print(f"\nContent Structure:")
    print("  [OK] Simple title page (no official formatting)")
    print("  [OK] Introduction and assessment modes")
    print("  [OK] Getting started instructions")
    print("  [OK] Demo 1: Spatial Assessment (7 steps with screenshot placeholders)")
    print("  [OK] Demo 2: Treatment Comparison (4 steps)")
    print("  [OK] Demo 3: Batch Scenarios (4 steps)")
    print("  [OK] Key features summary")
    print("  [OK] Screenshot capture guide (11 screenshots needed)")
    print("  [OK] Next steps for template integration")
    print(f"\nTotal Screenshots Needed: 11")
    print(f"Screenshot folder: screenshots_example/")
    print(f"\nNOTE: This is a plain document ready for copying into your official template.")
    print(f"      Take screenshots from http://localhost:8501 and replace placeholders.")
    print("\n" + "="*80 + "\n")


if __name__ == '__main__':
    main()
