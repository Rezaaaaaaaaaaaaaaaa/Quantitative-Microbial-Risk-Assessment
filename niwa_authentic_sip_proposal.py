#!/usr/bin/env python3
"""
Authentic NIWA SIP Proposal Generator for QMRA Workflow Engine
Matches the authentic NIWA tone, style, and format from SIP_example.docx
"""

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import os

class AuthenticNIWASIPProposal:
    def __init__(self):
        self.doc = Document()
        self.setup_authentic_styles()

    def setup_authentic_styles(self):
        """Setup styles matching NIWA's actual format."""
        # Simple, professional styles like the example
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)

        # Heading styles - simple and direct
        heading1 = self.doc.styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(14)
        heading1.font.bold = True

        heading3 = self.doc.styles['Heading 3']
        heading3.font.name = 'Calibri'
        heading3.font.size = Pt(12)
        heading3.font.bold = True

    def create_simple_diagram(self):
        """Create a simple, clear system diagram matching NIWA style."""
        fig, ax = plt.subplots(1, 1, figsize=(10, 6))
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.axis('off')

        # Simple colors
        blue = '#4472C4'
        green = '#70AD47'
        orange = '#FFC000'

        # Title
        ax.text(5, 5.5, 'QMRA Workflow Engine', fontsize=14, fontweight='bold', ha='center')

        # Input box
        input_box = FancyBboxPatch((0.5, 3.5), 2, 1, boxstyle="round,pad=0.1",
                                  facecolor='lightblue', edgecolor=blue, linewidth=1)
        ax.add_patch(input_box)
        ax.text(1.5, 4, 'Data Inputs\n• Pathogen data\n• Treatment data', fontsize=10, ha='center')

        # Processing box
        process_box = FancyBboxPatch((4, 3.5), 2, 1, boxstyle="round,pad=0.1",
                                    facecolor='lightgreen', edgecolor=green, linewidth=1)
        ax.add_patch(process_box)
        ax.text(5, 4, 'Python Engine\n• R integration\n• Monte Carlo', fontsize=10, ha='center')

        # Output box
        output_box = FancyBboxPatch((7.5, 3.5), 2, 1, boxstyle="round,pad=0.1",
                                   facecolor='lightyellow', edgecolor=orange, linewidth=1)
        ax.add_patch(output_box)
        ax.text(8.5, 4, 'Reports\n• Risk assessment\n• Compliance docs', fontsize=10, ha='center')

        # Arrows
        ax.annotate('', xy=(3.9, 4), xytext=(2.6, 4), arrowprops=dict(arrowstyle='->', lw=1.5))
        ax.annotate('', xy=(7.4, 4), xytext=(6.1, 4), arrowprops=dict(arrowstyle='->', lw=1.5))

        # Benefits
        ax.text(5, 2.5, 'Benefits: 60-70% faster delivery, standardized approach, regulatory compliance',
                fontsize=11, ha='center', style='italic')

        plt.tight_layout()
        diagram_path = os.path.join(os.getcwd(), 'niwa_qmra_diagram.png')
        plt.savefig(diagram_path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        return diagram_path

    def add_title_section(self):
        """Add title section matching NIWA format."""
        title = self.doc.add_paragraph("Structured internal project application 2024-2025")
        title.style = 'Normal'

        project_title = self.doc.add_paragraph("Development of QMRA Workflow Engine")
        project_title.style = 'Normal'

    def add_project_overview(self):
        """Add project overview in NIWA style."""
        overview_heading = self.doc.add_paragraph("Project Overview")
        overview_heading.style = 'Normal'

        overview_text = """
        NIWA will develop a Python-based quantitative microbial risk assessment (QMRA) workflow engine to address current competitive disadvantages and position for regulatory opportunities. We are losing bids because our manual process takes 60% longer than competitors who use standardized tools.

        The August 2025 wastewater performance standards and upcoming consent renewals for 60% of NZ treatment plants create significant market opportunity. Our current approach rebuilds models from scratch for each project, making us slower and less competitive.

        This tool will integrate our existing R code with Python frameworks, standardize pathogen databases and treatment models, and automate report generation. Expected outcome is 60-70% reduction in project delivery time while maintaining our technical quality.

        Budget requested: $25,000 over 6 months. Expected ROI: 400-800% within 18 months through faster delivery and increased win rate.
        """

        para = self.doc.add_paragraph(overview_text.strip())

    def add_work_programme(self):
        """Add work programme section matching NIWA format."""
        work_heading = self.doc.add_paragraph("Work programme and timeline")
        work_heading.style = 'Heading 3'

        instruction = self.doc.add_paragraph(
            "Tasks to be done, who will do what and by when."
        )

        # Simple table like the example
        table = self.doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'

        # Headers
        headers = ["Phase", "Who", "When", "Deliverable"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        # Timeline data - concise like NIWA style
        timeline_data = [
            ("Design & setup", "Reza, David", "Month 1", "System architecture, R code audit"),
            ("Core development", "Reza", "Months 2-3", "Pathogen database, simulation engine"),
            ("Integration", "Reza, David", "Month 4", "R integration, treatment modules"),
            ("Testing & reports", "Reza", "Month 5", "Validation, automated reporting"),
            ("Documentation", "Reza, David", "Month 6", "User guides, training materials"),
            ("Deployment", "Team", "Month 6", "Live system, staff training")
        ]

        for i, (phase, who, when, deliverable) in enumerate(timeline_data, 1):
            table.cell(i, 0).text = phase
            table.cell(i, 1).text = who
            table.cell(i, 2).text = when
            table.cell(i, 3).text = deliverable

    def add_technical_details(self):
        """Add technical section with NIWA's practical approach."""
        tech_heading = self.doc.add_paragraph("Technical approach")
        tech_heading.style = 'Heading 3'

        tech_text = """
        Our current manual QMRA process is inefficient and inconsistent. We rebuild dose-response models, treatment calculations, and Monte Carlo simulations for every project. This takes 40-80 hours per project vs 20-30 hours for competitors with automated tools.

        The new system will use Python with our existing R libraries integrated through rpy2. We'll build standardized modules for:
        - Pathogen dose-response relationships with uncertainty distributions
        - Treatment efficacy models for common NZ processes
        - Automated Monte Carlo simulation (target >1000 iterations/second)
        - Report templates meeting regulatory requirements

        David Wood will lead R integration based on his work with Mike Hickford on moving away from @Risk. I'll handle Python development and system architecture.

        We'll validate against existing projects and benchmark performance. Target is 60-70% time reduction while maintaining our analytical quality.
        """

        self.doc.add_paragraph(tech_text.strip())

    def add_system_diagram(self):
        """Add simple system diagram."""
        diagram_heading = self.doc.add_paragraph("System design")
        diagram_heading.style = 'Heading 3'

        # Create and insert diagram
        diagram_path = self.create_simple_diagram()

        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_picture(diagram_path, width=Inches(6))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Clean up
        if os.path.exists(diagram_path):
            os.remove(diagram_path)

    def add_budget_section(self):
        """Add budget section in NIWA style."""
        budget_heading = self.doc.add_paragraph("Budget and justification")
        budget_heading.style = 'Heading 3'

        # Simple budget table
        table = self.doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        headers = ["Item", "Hours/Cost", "Total"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        budget_data = [
            ("Reza Moghaddam (development)", "150 hrs @ $125", "$18,750"),
            ("David Wood (R integration)", "40 hrs @ $125", "$5,000"),
            ("Computing/documentation", "-", "$1,250")
        ]

        for i, (item, hours, total) in enumerate(budget_data, 1):
            table.cell(i, 0).text = item
            table.cell(i, 1).text = hours
            table.cell(i, 2).text = total

        budget_text = """
        Total: $25,000

        Justification: Current inefficiency is costing us projects. Recent losses to faster competitors demonstrate direct business impact. Time savings on 2-3 projects will recover the investment. Enhanced competitiveness expected to increase win rate by 25-30%, generating $100-200k additional annual revenue.

        The August 2025 regulatory deadline creates time-critical opportunity. Early deployment positions NIWA for significant market share in the $25-50M regulatory compliance market.
        """

        self.doc.add_paragraph(budget_text.strip())

    def add_impact_section(self):
        """Add impact section with specific examples like NIWA style."""
        impact_heading = self.doc.add_paragraph("Expected impact and outcomes")
        impact_heading.style = 'Heading 3'

        impact_text = """
        Our current approach for a typical QMRA project (e.g., wastewater discharge assessment) takes 60-80 hours: 20 hours building dose-response models, 25 hours treatment calculations, 15 hours Monte Carlo setup, 20 hours report writing. Competitors do this in 30-40 hours using automated tools.

        The workflow engine will reduce this to 20-30 hours: 5 hours parameter input, 10 hours scenario setup, 5 hours validation, 10 hours report review. This 60-70% improvement makes us competitive again.

        Specific benefits:
        - Win rate improvement from 60% to 80-85% (4-6 additional projects annually)
        - Faster turnaround enables more concurrent projects
        - Standardized approach reduces QA time and improves consistency
        - Positions NIWA for August 2025 regulatory opportunity ($25-50M market)

        Risk mitigation: Low technical risk due to proven technologies and team experience. David's R expertise minimizes integration challenges. Modular design allows incremental deployment if needed.
        """

        self.doc.add_paragraph(impact_text.strip())

    def add_references_section(self):
        """Add references in NIWA's practical style."""
        ref_heading = self.doc.add_paragraph("Supporting information")
        ref_heading.style = 'Heading 3'

        ref_text = """
        Key QMRA literature supporting this approach:

        {Ramos, 2021 #1} - Quantitative microbiological risk assessment in dairy products: demonstrates standardized approach benefits
        {Tang, 2024 #2} - QMRA applications to respiratory pathogens: shows policy relevance and rapid deployment potential
        {Phiri, 2021 #6} - Microbial contamination in NZ drinking water: validates need for NZ-specific QMRA tools
        {Dada, 2021 #4} - QMRA for wastewater treatment: demonstrates methodology for our target applications
        {Seis, 2020 #8} - Bayesian approaches in QMRA: provides advanced methods for uncertainty analysis

        For reference: NIWA Impact Strategy on competitive positioning and regulatory engagement. See Statement of Corporate Intent on strategic investments for capability development.
        """

        self.doc.add_paragraph(ref_text.strip())

    def generate_proposal(self, filename="NIWA_Authentic_SIP_QMRA_Proposal.docx"):
        """Generate the authentic NIWA-style proposal."""
        print("Generating authentic NIWA SIP proposal...")

        # Add sections in NIWA order
        self.add_title_section()
        self.add_project_overview()
        self.add_work_programme()
        self.add_technical_details()
        self.add_system_diagram()
        self.add_budget_section()
        self.add_impact_section()
        self.add_references_section()

        # Save document
        filepath = os.path.join(os.getcwd(), filename)
        self.doc.save(filepath)
        print(f"Authentic NIWA SIP proposal saved as: {filepath}")

        return filepath

def main():
    """Generate authentic NIWA SIP proposal."""
    generator = AuthenticNIWASIPProposal()
    proposal_path = generator.generate_proposal()

    print("\n" + "="*60)
    print("AUTHENTIC NIWA SIP PROPOSAL COMPLETE")
    print("="*60)
    print(f"Document: {proposal_path}")
    print("\nAuthentic NIWA Features:")
    print("• Concise, practical writing style")
    print("• Direct problem-solution approach")
    print("• Specific examples and metrics")
    print("• Simple tables and clear timelines")
    print("• Practical budget justification")
    print("• Focus on business impact over academic theory")
    print("• NIWA's actual tone and terminology")

if __name__ == "__main__":
    main()