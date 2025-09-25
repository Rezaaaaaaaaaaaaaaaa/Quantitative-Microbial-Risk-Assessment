"""
Generate MS Word versions of the QMRA Toolkit documentation
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_word_document():
    """Create comprehensive Word documentation for QMRA Toolkit."""

    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(24)

    # Heading styles
    heading1_style = styles['Heading 1']
    heading1_style.font.name = 'Calibri'
    heading1_style.font.size = Pt(16)

    heading2_style = styles['Heading 2']
    heading2_style.font.name = 'Calibri'
    heading2_style.font.size = Pt(14)

    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    # Add title
    title = doc.add_heading('QMRA Assessment Toolkit', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('User Guide and Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    # Add metadata
    doc.add_paragraph()
    metadata = doc.add_paragraph()
    metadata.add_run('Developed by: ').bold = True
    metadata.add_run('NIWA Earth Sciences New Zealand\n')
    metadata.add_run('Version: ').bold = True
    metadata.add_run('1.0\n')
    metadata.add_run('Date: ').bold = True
    metadata.add_run(datetime.now().strftime('%B %Y'))
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading('Table of Contents', level=1)
    toc = doc.add_paragraph('1. Overview\n2. Installation\n3. Quick Start\n4. Command-Line Interface\n'
                           '5. Python API\n6. Exposure Routes\n7. Pathogen Database\n'
                           '8. Treatment Modeling\n9. Report Generation\n10. Examples\n'
                           '11. Troubleshooting\n12. Best Practices')

    doc.add_page_break()

    # 1. Overview
    doc.add_heading('1. Overview', level=1)

    doc.add_paragraph(
        'The QMRA Assessment Toolkit is a comprehensive Python-based solution for Quantitative '
        'Microbial Risk Assessment, developed by NIWA Earth Sciences New Zealand. This toolkit '
        'replaces @Risk Excel functionality with automated, reproducible workflows for regulatory '
        'compliance QMRA assessments.'
    )

    # Key Features
    doc.add_heading('Key Features', level=2)

    features = [
        'Comprehensive Pathogen Database - Validated dose-response models for key pathogens',
        'Multiple Exposure Routes - Primary contact, shellfish consumption, drinking water',
        'Dilution Modeling Integration - NIWA\'s key differentiator with engineer-provided LRVs',
        'Monte Carlo Simulation - Advanced uncertainty analysis replacing @Risk functionality',
        'Risk Characterization - Complete risk metrics calculation (infection, illness, DALYs)',
        'Automated Reporting - Generate regulatory compliance reports in Word format',
        'Command-Line Interface - Easy-to-use CLI for common workflows'
    ]

    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')

    # Benefits over @Risk
    doc.add_heading('Benefits Over @Risk Excel', level=2)

    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = '@Risk'
    hdr_cells[2].text = 'QMRA Toolkit'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    benefits_data = [
        ['Platform', 'Excel-dependent', 'Native Python'],
        ['Security', 'Firewall conflicts', 'No external dependencies'],
        ['Automation', 'Manual processes', 'Fully automated'],
        ['Reproducibility', 'Limited', 'Complete version control'],
        ['Cost', 'Commercial license', 'Open source'],
        ['Integration', 'Limited', 'NIWA dilution modeling']
    ]

    for row_data in benefits_data:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        row.cells[2].text = row_data[2]

    doc.add_page_break()

    # 2. Installation
    doc.add_heading('2. Installation', level=1)

    doc.add_heading('Requirements', level=2)
    requirements = [
        'Python 3.8 or higher',
        'NumPy and SciPy for numerical computations',
        'Pandas for data handling',
        'Matplotlib for plotting',
        'python-docx for report generation',
        'PyYAML for configuration files',
        'Click for command-line interface'
    ]

    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('Setup Instructions', level=2)

    doc.add_paragraph('1. Navigate to the toolkit directory:', style='List Number')
    code1 = doc.add_paragraph('cd qmra_toolkit')
    code1.style = 'Intense Quote'

    doc.add_paragraph('2. Install dependencies:', style='List Number')
    code2 = doc.add_paragraph('pip install -r requirements.txt')
    code2.style = 'Intense Quote'

    doc.add_paragraph('3. Verify installation:', style='List Number')
    code3 = doc.add_paragraph('python tests/run_all_tests.py')
    code3.style = 'Intense Quote'

    doc.add_paragraph('4. Test command-line interface:', style='List Number')
    code4 = doc.add_paragraph('python src/qmra_toolkit.py --help')
    code4.style = 'Intense Quote'

    doc.add_page_break()

    # 3. Quick Start
    doc.add_heading('3. Quick Start', level=1)

    doc.add_heading('Basic Risk Assessment Example', level=2)
    doc.add_paragraph(
        'Run a complete QMRA assessment for recreational water exposure to norovirus:'
    )

    basic_cmd = doc.add_paragraph(
        'python src/qmra_toolkit.py assess \\\n'
        '  --pathogen norovirus \\\n'
        '  --exposure-route primary_contact \\\n'
        '  --concentration 10.0 \\\n'
        '  --volume 50.0 \\\n'
        '  --frequency 10 \\\n'
        '  --population 10000 \\\n'
        '  --report'
    )
    basic_cmd.style = 'Intense Quote'

    doc.add_paragraph('This command will:')
    actions = [
        'Assess norovirus exposure from recreational water',
        'Use 10 organisms per 100mL water concentration',
        'Assume 50 mL water ingestion per swimming event',
        'Calculate risk for 10 swimming events per year',
        'Apply results to a population of 10,000',
        'Generate a comprehensive compliance report in Word format'
    ]

    for action in actions:
        doc.add_paragraph(action, style='List Bullet')

    doc.add_heading('Example Output', level=2)

    example_output = doc.add_paragraph(
        'QMRA Assessment Results: Norovirus\n'
        '============================================================\n\n'
        'Infection Probability:\n'
        '  Mean: 1.65e-01\n'
        '  Median: 1.65e-01\n'
        '  95th Percentile: 1.65e-01\n\n'
        'Annual Risk:\n'
        '  Mean: 8.36e-01\n'
        '  Median: 8.36e-01\n'
        '  95th Percentile: 8.36e-01\n'
        '  Expected cases per year: 8361\n\n'
        'Regulatory Compliance:\n'
        '  recreational_water_risk: FAIL\n'
        '  acceptable_annual_risk: FAIL'
    )
    example_output.style = 'Intense Quote'

    doc.add_page_break()

    # 4. Command-Line Interface
    doc.add_heading('4. Command-Line Interface', level=1)

    doc.add_paragraph(
        'The QMRA toolkit provides a comprehensive command-line interface for common workflows.'
    )

    doc.add_heading('Main Commands', level=2)

    # assess command
    doc.add_heading('assess - Complete Risk Assessment', level=3)
    doc.add_paragraph('Performs a complete QMRA assessment for specified pathogen and exposure scenario.')

    doc.add_paragraph('Required Parameters:', style='List Bullet')
    assess_params = [
        '--pathogen / -p: Pathogen name (norovirus, campylobacter, cryptosporidium)',
        '--exposure-route / -e: Exposure route (primary_contact, shellfish_consumption)',
        '--concentration / -c: Pathogen concentration (organisms per unit)'
    ]

    for param in assess_params:
        doc.add_paragraph(param, style='List Bullet 2')

    doc.add_paragraph('Optional Parameters:', style='List Bullet')
    assess_optional = [
        '--volume: Exposure volume (mL for water, grams for food)',
        '--frequency: Exposure frequency (events per year)',
        '--population: Population size for population risk calculations',
        '--iterations / -n: Number of Monte Carlo iterations (default: 10,000)',
        '--output / -o: Output file for results (JSON format)',
        '--report / -r: Generate comprehensive Word report'
    ]

    for param in assess_optional:
        doc.add_paragraph(param, style='List Bullet 2')

    # Other commands
    doc.add_heading('list-pathogens - Available Pathogens', level=3)
    doc.add_paragraph('Lists all available pathogens in the database.')

    list_cmd = doc.add_paragraph('python src/qmra_toolkit.py list-pathogens')
    list_cmd.style = 'Intense Quote'

    doc.add_heading('pathogen-info - Pathogen Details', level=3)
    doc.add_paragraph('Get detailed information about a specific pathogen.')

    info_cmd = doc.add_paragraph('python src/qmra_toolkit.py pathogen-info --pathogen norovirus')
    info_cmd.style = 'Intense Quote'

    doc.add_page_break()

    # 5. Python API
    doc.add_heading('5. Python API', level=1)

    doc.add_paragraph(
        'For advanced users, the toolkit provides a comprehensive Python API for integration '
        'into custom workflows and applications.'
    )

    doc.add_heading('Basic API Usage Example', level=2)

    api_example = doc.add_paragraph(
        'from pathogen_database import PathogenDatabase\n'
        'from exposure_assessment import create_exposure_assessment, ExposureRoute\n'
        'from risk_characterization import RiskCharacterization\n\n'
        '# Initialize components\n'
        'pathogen_db = PathogenDatabase()\n'
        'risk_calc = RiskCharacterization(pathogen_db)\n\n'
        '# Configure exposure\n'
        'exposure_params = {\n'
        '    "water_ingestion_volume": 50.0,  # mL per event\n'
        '    "exposure_frequency": 10         # events per year\n'
        '}\n\n'
        'exposure_model = create_exposure_assessment(\n'
        '    ExposureRoute.PRIMARY_CONTACT,\n'
        '    exposure_params\n'
        ')\n'
        'exposure_model.set_pathogen_concentration(10.0)\n\n'
        '# Run assessment\n'
        'results = risk_calc.run_comprehensive_assessment(\n'
        '    pathogen_name="norovirus",\n'
        '    exposure_assessment=exposure_model,\n'
        '    population_size=10000\n'
        ')\n\n'
        '# Print results\n'
        'for metric_name, result in results.items():\n'
        '    print(f"{metric_name}: {result.statistics[\'mean\']:.2e}")'
    )
    api_example.style = 'Intense Quote'

    doc.add_page_break()

    # 6. Exposure Routes
    doc.add_heading('6. Exposure Routes', level=1)

    doc.add_heading('Primary Contact (Recreational Water)', level=2)
    doc.add_paragraph(
        'Models inadvertent water ingestion during swimming, surfing, or other water activities.'
    )

    doc.add_paragraph('Required Parameters:')
    primary_params = [
        'water_ingestion_volume: Volume ingested per event (mL)',
        'exposure_frequency: Number of exposure events per year'
    ]
    for param in primary_params:
        doc.add_paragraph(param, style='List Bullet')

    doc.add_paragraph('Typical Concentrations:')
    primary_conc = [
        'Beach water: 1-100 organisms per 100 mL',
        'River water: 10-1000 organisms per 100 mL',
        'Treated wastewater discharge: 0.1-10 organisms per 100 mL'
    ]
    for conc in primary_conc:
        doc.add_paragraph(conc, style='List Bullet')

    doc.add_heading('Shellfish Consumption', level=2)
    doc.add_paragraph(
        'Models pathogen exposure through shellfish consumption (oysters, mussels, clams).'
    )

    doc.add_paragraph('Required Parameters:')
    shellfish_params = [
        'shellfish_consumption: Mass consumed per serving (grams)',
        'consumption_frequency: Number of servings per year'
    ]
    for param in shellfish_params:
        doc.add_paragraph(param, style='List Bullet')

    doc.add_paragraph('Typical Concentrations:')
    shellfish_conc = [
        'Raw oysters: 100-10,000 organisms per 100g',
        'Cooked shellfish: 1-100 organisms per 100g'
    ]
    for conc in shellfish_conc:
        doc.add_paragraph(conc, style='List Bullet')

    doc.add_page_break()

    # 7. Pathogen Database
    doc.add_heading('7. Pathogen Database', level=1)

    doc.add_heading('Available Pathogens', level=2)

    doc.add_paragraph('The toolkit includes dose-response models for:')

    # Norovirus
    doc.add_paragraph('Norovirus', style='List Bullet')
    doc.add_paragraph('Model: Beta-Poisson (α=0.04, β=0.055)', style='List Bullet 2')
    doc.add_paragraph('Source: Teunis et al. (2008)', style='List Bullet 2')
    doc.add_paragraph('Illness-to-infection ratio: 0.7', style='List Bullet 2')

    # Campylobacter
    doc.add_paragraph('Campylobacter jejuni', style='List Bullet')
    doc.add_paragraph('Model: Beta-Poisson (α=0.145, β=7.59)', style='List Bullet 2')
    doc.add_paragraph('Source: Teunis et al. (2005)', style='List Bullet 2')
    doc.add_paragraph('Illness-to-infection ratio: 0.3', style='List Bullet 2')

    # Cryptosporidium
    doc.add_paragraph('Cryptosporidium parvum', style='List Bullet')
    doc.add_paragraph('Model: Exponential (r=0.0042)', style='List Bullet 2')
    doc.add_paragraph('Source: Haas et al. (1996)', style='List Bullet 2')
    doc.add_paragraph('Illness-to-infection ratio: 1.0', style='List Bullet 2')

    doc.add_page_break()

    # 8. Treatment Modeling
    doc.add_heading('8. Treatment and Dilution Modeling', level=1)

    doc.add_paragraph(
        'The toolkit includes sophisticated dilution modeling capabilities that integrate '
        'engineer-provided Log Reduction Values (LRVs), multi-barrier treatment trains, '
        'and uncertainty quantification - NIWA\'s key differentiator.'
    )

    doc.add_heading('Treatment Configuration', level=2)
    doc.add_paragraph('Create a YAML configuration file for treatment trains:')

    treatment_yaml = doc.add_paragraph(
        'treatment_barriers:\n'
        '  - name: "Primary Settling"\n'
        '    type: "physical"\n'
        '    lrv: 0.5\n'
        '    variability: 0.2\n\n'
        '  - name: "Activated Sludge"\n'
        '    type: "biological"\n'
        '    lrv: 2.0\n'
        '    variability: 0.5\n\n'
        '  - name: "UV Disinfection"\n'
        '    type: "uv_disinfection"\n'
        '    lrv: 3.0\n'
        '    variability: 0.3'
    )
    treatment_yaml.style = 'Intense Quote'

    doc.add_heading('Python API Usage', level=2)

    treatment_api = doc.add_paragraph(
        'from dilution_model import DilutionModel, TreatmentBarrier, TreatmentType\n\n'
        '# Create dilution model\n'
        'model = DilutionModel()\n\n'
        '# Add treatment barriers\n'
        'model.add_treatment_barrier(TreatmentBarrier(\n'
        '    name="UV Treatment",\n'
        '    treatment_type=TreatmentType.UV,\n'
        '    log_reduction_value=3.0,\n'
        '    variability=0.2\n'
        '))\n\n'
        '# Apply to concentration\n'
        'initial_conc = 1e6  # organisms/L\n'
        'results = model.apply_complete_scenario(initial_conc)'
    )
    treatment_api.style = 'Intense Quote'

    doc.add_page_break()

    # 9. Report Generation
    doc.add_heading('9. Automated Report Generation', level=1)

    doc.add_paragraph(
        'The toolkit automatically generates comprehensive regulatory compliance reports in '
        'professional Word format.'
    )

    doc.add_heading('Report Contents', level=2)

    report_contents = [
        'Executive Summary with key findings',
        'Introduction and Assessment Objectives',
        'Methodology and Model Selection',
        'Risk Assessment Results with statistical analysis',
        'Uncertainty Analysis and Monte Carlo results',
        'Regulatory Compliance Assessment',
        'Conclusions and Recommendations',
        'Technical Appendices with detailed data'
    ]

    for content in report_contents:
        doc.add_paragraph(content, style='List Number')

    doc.add_heading('Generating Reports', level=2)
    doc.add_paragraph('Reports are automatically generated when using the --report flag:')

    report_cmd = doc.add_paragraph(
        'python src/qmra_toolkit.py assess \\\n'
        '  --pathogen norovirus \\\n'
        '  --exposure-route primary_contact \\\n'
        '  --concentration 10.0 \\\n'
        '  --report'
    )
    report_cmd.style = 'Intense Quote'

    doc.add_page_break()

    # 10. Examples
    doc.add_heading('10. Advanced Examples', level=1)

    doc.add_paragraph(
        'The toolkit includes advanced example scripts demonstrating complex analysis workflows.'
    )

    doc.add_heading('Multi-Pathogen Risk Comparison', level=2)
    doc.add_paragraph('Compare risk levels across different pathogens under identical conditions:')

    example1_cmd = doc.add_paragraph('cd examples\npython pathogen_comparison.py')
    example1_cmd.style = 'Intense Quote'

    doc.add_paragraph('This example demonstrates:')
    example1_features = [
        'Automatic model selection for different pathogens',
        'Risk ranking and comparative analysis',
        'Population-level risk calculations',
        'Regulatory compliance assessment across pathogens'
    ]
    for feature in example1_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Treatment Scenario Comparison', level=2)
    doc.add_paragraph('Evaluate effectiveness of different treatment levels:')

    example2_cmd = doc.add_paragraph('cd examples\npython scenario_comparison.py')
    example2_cmd.style = 'Intense Quote'

    doc.add_paragraph('This example demonstrates:')
    example2_features = [
        'Treatment barrier modeling with log reduction values',
        'Risk reduction analysis across treatment scenarios',
        'Cost-benefit evaluation of treatment investments'
    ]
    for feature in example2_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 11. Troubleshooting
    doc.add_heading('11. Troubleshooting', level=1)

    doc.add_heading('Common Issues', level=2)

    issues = [
        ('Import Errors', 'Ensure all dependencies are installed: pip install -r requirements.txt'),
        ('Pathogen Not Found', 'Check available pathogens with: python src/qmra_toolkit.py list-pathogens'),
        ('Configuration Issues', 'Verify YAML syntax in treatment configuration files'),
        ('Memory Issues', 'Reduce Monte Carlo iterations for large simulations using --iterations parameter'),
        ('Report Generation Fails', 'Ensure write permissions for output directory and python-docx is installed')
    ]

    for issue, solution in issues:
        doc.add_paragraph(f'{issue}:', style='List Bullet').runs[0].bold = True
        doc.add_paragraph(solution, style='List Bullet 2')

    doc.add_heading('Getting Help', level=2)
    help_options = [
        'Run the test suite: python tests/run_all_tests.py',
        'View CLI help: python src/qmra_toolkit.py --help',
        'Check pathogen information: python src/qmra_toolkit.py pathogen-info --pathogen [name]',
        'Review example scripts in the examples/ directory'
    ]

    for option in help_options:
        doc.add_paragraph(option, style='List Bullet')

    doc.add_page_break()

    # 12. Best Practices
    doc.add_heading('12. Best Practices', level=1)

    doc.add_heading('Model Selection', level=2)
    model_practices = [
        'Use Beta-Poisson models for most viral and bacterial pathogens',
        'Use Exponential models for highly infectious pathogens or conservative estimates',
        'Validate parameters against published literature',
        'Document model selection rationale in reports'
    ]
    for practice in model_practices:
        doc.add_paragraph(practice, style='List Bullet')

    doc.add_heading('Uncertainty Analysis', level=2)
    uncertainty_practices = [
        'Use at least 10,000 iterations for robust uncertainty estimates',
        'Include parameter uncertainty in LRV values',
        'Document assumptions clearly in reports',
        'Consider sensitivity analysis for key parameters'
    ]
    for practice in uncertainty_practices:
        doc.add_paragraph(practice, style='List Bullet')

    doc.add_heading('Quality Assurance', level=2)
    qa_practices = [
        'Run validation tests regularly',
        'Compare results with previous assessments',
        'Review reports before submission',
        'Archive model inputs with results',
        'Use version control for project files'
    ]
    for practice in qa_practices:
        doc.add_paragraph(practice, style='List Bullet')

    doc.add_page_break()

    # Appendix
    doc.add_heading('Appendix: Technical Specifications', level=1)

    doc.add_heading('System Requirements', level=2)
    sys_req = [
        'Operating System: Windows, macOS, or Linux',
        'Python: Version 3.8 or higher',
        'Memory: Minimum 4GB RAM (8GB recommended for large simulations)',
        'Disk Space: 100MB for toolkit + space for output reports',
        'Network: Internet connection for initial package installation'
    ]
    for req in sys_req:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('File Structure', level=2)
    file_structure = doc.add_paragraph(
        'qmra_toolkit/\n'
        '├── src/                   # Core modules\n'
        '│   ├── pathogen_database.py\n'
        '│   ├── dose_response.py\n'
        '│   ├── exposure_assessment.py\n'
        '│   ├── dilution_model.py\n'
        '│   ├── monte_carlo.py\n'
        '│   ├── risk_characterization.py\n'
        '│   ├── report_generator.py\n'
        '│   └── qmra_toolkit.py    # Main CLI\n'
        '├── data/                  # Pathogen database\n'
        '├── tests/                 # Test suite\n'
        '├── examples/              # Usage examples\n'
        '├── docs/                  # Documentation\n'
        '└── requirements.txt       # Dependencies'
    )
    file_structure.style = 'Intense Quote'

    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('QMRA Assessment Toolkit v1.0\n').font.size = Pt(10)
    footer_para.add_run('© NIWA Earth Sciences New Zealand\n').font.size = Pt(10)
    footer_para.add_run('Strategic Investment Programme 2025-2026').font.size = Pt(10)

    return doc

def main():
    """Generate Word documentation files."""
    print("Generating MS Word documentation...")

    # Create comprehensive user guide
    doc = create_word_document()

    # Save the document
    output_file = 'QMRA_Toolkit_User_Guide.docx'
    doc.save(output_file)

    print(f"[OK] Generated comprehensive user guide: {output_file}")

    # Create quick reference card
    quick_ref = Document()

    # Quick reference title
    quick_ref.add_heading('QMRA Toolkit - Quick Reference', 0)

    # Installation
    quick_ref.add_heading('Installation', level=1)
    quick_ref.add_paragraph('pip install -r requirements.txt', style='Intense Quote')

    # Basic commands
    quick_ref.add_heading('Basic Commands', level=1)

    commands = [
        ('List pathogens', 'python src/qmra_toolkit.py list-pathogens'),
        ('Basic assessment', 'python src/qmra_toolkit.py assess -p norovirus -e primary_contact -c 10.0'),
        ('With report', 'python src/qmra_toolkit.py assess -p norovirus -e primary_contact -c 10.0 --report'),
        ('Get help', 'python src/qmra_toolkit.py --help')
    ]

    for desc, cmd in commands:
        quick_ref.add_paragraph(f'{desc}:')
        quick_ref.add_paragraph(cmd, style='Intense Quote')

    # Available pathogens
    quick_ref.add_heading('Available Pathogens', level=1)
    pathogens = [
        'norovirus (Beta-Poisson model)',
        'campylobacter (Beta-Poisson model)',
        'cryptosporidium (Exponential model)'
    ]
    for pathogen in pathogens:
        quick_ref.add_paragraph(pathogen, style='List Bullet')

    # Exposure routes
    quick_ref.add_heading('Exposure Routes', level=1)
    routes = [
        'primary_contact (recreational water)',
        'shellfish_consumption (oysters, mussels, clams)'
    ]
    for route in routes:
        quick_ref.add_paragraph(route, style='List Bullet')

    quick_ref_file = 'QMRA_Toolkit_Quick_Reference.docx'
    quick_ref.save(quick_ref_file)

    print(f"[OK] Generated quick reference guide: {quick_ref_file}")
    print("\nMS Word documentation generated successfully!")
    print(f"Files created:")
    print(f"  - {output_file} (Comprehensive user guide)")
    print(f"  - {quick_ref_file} (Quick reference)")

if __name__ == "__main__":
    main()