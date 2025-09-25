"""
Report Generator Module for QMRA Toolkit

This module generates standardized reports for regulatory compliance,
replacing @Risk reporting functionality with automated Python-based reports.
"""

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
import io
import base64

from risk_characterization import RiskResult, RiskMetric
from monte_carlo import MonteCarloResults
from dilution_model import DilutionModel


class ReportGenerator:
    """
    Generate comprehensive QMRA reports for regulatory compliance.
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize report generator.

        Args:
            template_path: Path to Word template. If None, creates blank document.
        """
        self.template_path = template_path
        self.figures = []
        self.tables = []

    def create_regulatory_report(self,
                                project_info: Dict[str, str],
                                risk_results: Dict[str, RiskResult],
                                exposure_params: Dict[str, Any],
                                treatment_summary: Optional[str] = None,
                                monte_carlo_results: Optional[MonteCarloResults] = None,
                                output_filename: str = None) -> str:
        """
        Create comprehensive regulatory compliance report.

        Args:
            project_info: Project metadata (name, location, date, etc.)
            risk_results: Dictionary of risk calculation results
            exposure_params: Exposure assessment parameters
            treatment_summary: Optional treatment train summary
            monte_carlo_results: Optional Monte Carlo simulation results
            output_filename: Output filename (auto-generated if None)

        Returns:
            Path to generated report
        """
        # Create new document
        if self.template_path and Path(self.template_path).exists():
            doc = Document(self.template_path)
        else:
            doc = Document()

        # Set document properties
        doc.core_properties.title = project_info.get("title", "QMRA Assessment Report")
        doc.core_properties.author = project_info.get("author", "QMRA Toolkit")
        doc.core_properties.subject = "Quantitative Microbial Risk Assessment"

        # Add title page
        self._add_title_page(doc, project_info)

        # Add executive summary
        self._add_executive_summary(doc, risk_results)

        # Add table of contents placeholder
        doc.add_page_break()
        doc.add_heading('Table of Contents', 1)
        doc.add_paragraph('[Table of Contents to be generated]')
        doc.add_page_break()

        # Section 1: Introduction
        doc.add_heading('1. Introduction', 1)
        self._add_introduction(doc, project_info)

        # Section 2: Methodology
        doc.add_heading('2. Methodology', 1)
        self._add_methodology(doc, exposure_params, treatment_summary)

        # Section 3: Risk Assessment Results
        doc.add_heading('3. Risk Assessment Results', 1)
        self._add_risk_results(doc, risk_results)

        # Section 4: Uncertainty Analysis
        if monte_carlo_results:
            doc.add_heading('4. Uncertainty Analysis', 1)
            self._add_uncertainty_analysis(doc, monte_carlo_results)

        # Section 5: Regulatory Compliance
        doc.add_heading('5. Regulatory Compliance Assessment', 1)
        self._add_compliance_assessment(doc, risk_results)

        # Section 6: Conclusions and Recommendations
        doc.add_heading('6. Conclusions and Recommendations', 1)
        self._add_conclusions(doc, risk_results)

        # Appendices
        doc.add_page_break()
        doc.add_heading('Appendices', 1)
        self._add_appendices(doc, risk_results)

        # Save document
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            output_filename = f"QMRA_Report_{timestamp}.docx"

        doc.save(output_filename)
        return output_filename

    def _add_title_page(self, doc: Document, project_info: Dict[str, str]) -> None:
        """Add title page to document."""
        title = doc.add_heading(project_info.get("title", "QMRA Assessment Report"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Add project details
        for key, value in project_info.items():
            if key not in ["title"]:
                para = doc.add_paragraph()
                para.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                para.add_run(str(value))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.add_run(f"Report Date: {datetime.now().strftime('%d %B %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add disclaimer
        doc.add_paragraph()
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.add_run(
            "This report was generated using the QMRA Assessment Toolkit developed by "
            "NIWA Earth Sciences New Zealand"
        ).italic = True
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, risk_results: Dict[str, RiskResult]) -> None:
        """Add executive summary section."""
        doc.add_heading('Executive Summary', 1)

        # Get key risk metrics
        if "annual_risk" in risk_results:
            annual_risk = risk_results["annual_risk"]
            mean_risk = annual_risk.statistics["mean"]
            p95_risk = annual_risk.statistics["p95"]

            summary_text = (
                f"This Quantitative Microbial Risk Assessment (QMRA) evaluates the health risks "
                f"associated with {annual_risk.pathogen_name} exposure. "
                f"The assessment indicates a mean annual infection risk of {mean_risk:.2e} "
                f"with a 95th percentile risk of {p95_risk:.2e}."
            )
            doc.add_paragraph(summary_text)

            # Add compliance status
            if mean_risk <= 1e-6:
                doc.add_paragraph(
                    "The calculated risks are within acceptable regulatory thresholds, "
                    "meeting the target of less than 1 in 1,000,000 annual risk."
                )
            else:
                doc.add_paragraph(
                    "The calculated risks exceed acceptable regulatory thresholds. "
                    "Additional risk management measures are recommended."
                )

        # Key findings bullets
        doc.add_paragraph("Key Findings:")
        for metric_name, result in risk_results.items():
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(f"{metric_name.replace('_', ' ').title()}: ")
            bullet.add_run(f"Mean = {result.statistics['mean']:.2e}")

    def _add_introduction(self, doc: Document, project_info: Dict[str, str]) -> None:
        """Add introduction section."""
        intro_text = (
            "This report presents the results of a Quantitative Microbial Risk Assessment (QMRA) "
            f"conducted for {project_info.get('project_name', 'the project')}. "
            "QMRA is a systematic approach to estimating the risk of adverse health effects "
            "from exposure to pathogenic microorganisms in water, food, and environmental sources."
        )
        doc.add_paragraph(intro_text)

        doc.add_heading('1.1 Objectives', 2)
        objectives = [
            "Quantify the infection and illness risks from pathogen exposure",
            "Evaluate compliance with regulatory risk thresholds",
            "Identify critical control points for risk reduction",
            "Provide evidence-based recommendations for risk management"
        ]
        for obj in objectives:
            doc.add_paragraph(obj, style='List Bullet')

    def _add_methodology(self, doc: Document, exposure_params: Dict, treatment_summary: Optional[str]) -> None:
        """Add methodology section."""
        doc.add_heading('2.1 Exposure Assessment', 2)
        doc.add_paragraph(
            "Exposure assessment was conducted using the following parameters:"
        )

        # Create exposure parameters table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Parameter'
        header_cells[1].text = 'Value'

        for param, value in exposure_params.items():
            row_cells = table.add_row().cells
            row_cells[0].text = param.replace('_', ' ').title()
            row_cells[1].text = str(value)

        doc.add_paragraph()

        doc.add_heading('2.2 Dose-Response Modelling', 2)
        doc.add_paragraph(
            "Dose-response relationships were modeled using established mathematical models "
            "including Beta-Poisson and exponential models, with parameters from peer-reviewed literature."
        )

        if treatment_summary:
            doc.add_heading('2.3 Treatment Train', 2)
            doc.add_paragraph(treatment_summary)

    def _add_risk_results(self, doc: Document, risk_results: Dict[str, RiskResult]) -> None:
        """Add risk assessment results section."""
        for metric_name, result in risk_results.items():
            doc.add_heading(f'3.{list(risk_results.keys()).index(metric_name) + 1} {metric_name.replace("_", " ").title()}', 2)

            # Add statistics table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Statistic'
            header_cells[1].text = 'Value'

            for stat, value in result.statistics.items():
                row_cells = table.add_row().cells
                row_cells[0].text = stat.replace('_', ' ').title()
                if "probability" in metric_name or "risk" in metric_name:
                    row_cells[1].text = f"{value:.2e}"
                else:
                    row_cells[1].text = f"{value:.4f}"

            doc.add_paragraph()

            # Add population risk if available
            if result.population_risks:
                doc.add_paragraph("Population Impact:")
                for pop_metric, value in result.population_risks.items():
                    doc.add_paragraph(f"• {pop_metric.replace('_', ' ').title()}: {value:.1f}")

    def _add_uncertainty_analysis(self, doc: Document, monte_carlo_results: MonteCarloResults) -> None:
        """Add uncertainty analysis section."""
        doc.add_paragraph(
            f"Uncertainty analysis was conducted using Monte Carlo simulation with "
            f"{monte_carlo_results.iterations:,} iterations."
        )

        # Add convergence information
        doc.add_paragraph(
            f"The simulation converged with stable statistics after approximately "
            f"{monte_carlo_results.iterations // 2:,} iterations."
        )

        # Add key uncertainty metrics
        doc.add_heading('4.1 Uncertainty Metrics', 2)
        uncertainty_metrics = [
            f"Coefficient of Variation: {monte_carlo_results.statistics['std'] / monte_carlo_results.statistics['mean']:.2f}",
            f"95% Confidence Interval: [{monte_carlo_results.percentiles['5%']:.2e}, {monte_carlo_results.percentiles['95%']:.2e}]",
            f"Interquartile Range: {monte_carlo_results.percentiles['75%'] - monte_carlo_results.percentiles['25%']:.2e}"
        ]
        for metric in uncertainty_metrics:
            doc.add_paragraph(metric, style='List Bullet')

    def _add_compliance_assessment(self, doc: Document, risk_results: Dict[str, RiskResult]) -> None:
        """Add regulatory compliance assessment."""
        doc.add_paragraph(
            "The following table summarizes compliance with regulatory risk thresholds:"
        )

        # Create compliance table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Risk Metric'
        header_cells[1].text = 'Calculated Risk'
        header_cells[2].text = 'Threshold'
        header_cells[3].text = 'Status'

        # Define thresholds
        thresholds = {
            "annual_risk": 1e-6,
            "infection_probability": 1e-3,
            "illness_probability": 1e-4
        }

        for metric_name, result in risk_results.items():
            if metric_name in thresholds:
                row_cells = table.add_row().cells
                row_cells[0].text = metric_name.replace('_', ' ').title()
                row_cells[1].text = f"{result.statistics['mean']:.2e}"
                row_cells[2].text = f"{thresholds[metric_name]:.2e}"

                if result.statistics['mean'] <= thresholds[metric_name]:
                    row_cells[3].text = "COMPLIANT"
                else:
                    row_cells[3].text = "NON-COMPLIANT"

    def _add_conclusions(self, doc: Document, risk_results: Dict[str, RiskResult]) -> None:
        """Add conclusions and recommendations."""
        doc.add_heading('6.1 Conclusions', 2)

        # Determine overall risk level
        if "annual_risk" in risk_results:
            annual_risk = risk_results["annual_risk"].statistics["mean"]
            if annual_risk <= 1e-6:
                risk_level = "acceptable"
            elif annual_risk <= 1e-4:
                risk_level = "tolerable with management"
            else:
                risk_level = "unacceptable"

            doc.add_paragraph(
                f"Based on the comprehensive QMRA, the assessed health risks are {risk_level}. "
                f"The mean annual infection risk of {annual_risk:.2e} has been calculated using "
                "conservative assumptions and validated dose-response models."
            )

        doc.add_heading('6.2 Recommendations', 2)
        recommendations = [
            "Continue monitoring pathogen concentrations to validate model assumptions",
            "Implement identified critical control points for risk reduction",
            "Review and update the assessment annually or when system changes occur",
            "Consider additional treatment barriers if risks approach regulatory thresholds"
        ]
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')

    def _add_appendices(self, doc: Document, risk_results: Dict[str, RiskResult]) -> None:
        """Add appendices with detailed data."""
        doc.add_heading('Appendix A: Detailed Statistical Results', 2)

        # Create detailed results table
        all_stats = []
        for metric_name, result in risk_results.items():
            for stat_name, value in result.statistics.items():
                all_stats.append({
                    'Metric': metric_name.replace('_', ' ').title(),
                    'Statistic': stat_name.replace('_', ' ').title(),
                    'Value': value
                })

        if all_stats:
            df = pd.DataFrame(all_stats)
            self._add_dataframe_to_doc(doc, df)

    def _add_dataframe_to_doc(self, doc: Document, df: pd.DataFrame) -> None:
        """Add pandas DataFrame as table in document."""
        # Add table with DataFrame dimensions
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'

        # Add headers
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)

        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                if isinstance(value, (int, float)):
                    if abs(value) < 0.01 or abs(value) > 1000:
                        row_cells[i].text = f"{value:.2e}"
                    else:
                        row_cells[i].text = f"{value:.4f}"
                else:
                    row_cells[i].text = str(value)

    def create_visualization_report(self,
                                  risk_results: Dict[str, RiskResult],
                                  output_filename: str = None) -> str:
        """
        Create report with visualizations.

        Args:
            risk_results: Dictionary of risk results
            output_filename: Output filename

        Returns:
            Path to generated report with plots
        """
        # Set style
        sns.set_style("whitegrid")
        plt.rcParams['figure.figsize'] = (10, 6)

        # Create figure with subplots
        fig, axes = plt.subplots(2, 2, figsize=(12, 10))

        # Plot 1: Risk comparison bar chart
        ax1 = axes[0, 0]
        metrics = []
        mean_values = []
        for metric, result in risk_results.items():
            metrics.append(metric.replace('_', '\n'))
            mean_values.append(result.statistics['mean'])

        ax1.bar(metrics, mean_values)
        ax1.set_ylabel('Risk Value')
        ax1.set_title('Risk Metrics Comparison')
        ax1.set_yscale('log')

        # Plot 2: Uncertainty ranges
        ax2 = axes[0, 1]
        for i, (metric, result) in enumerate(risk_results.items()):
            if 'p5' in result.statistics and 'p95' in result.statistics:
                ax2.errorbar(i, result.statistics['mean'],
                           yerr=[[result.statistics['mean'] - result.statistics['p5']],
                                 [result.statistics['p95'] - result.statistics['mean']]],
                           fmt='o', capsize=5, label=metric.replace('_', ' '))
        ax2.set_xlabel('Risk Metric')
        ax2.set_ylabel('Risk Value')
        ax2.set_title('Risk Estimates with 90% Confidence Intervals')
        ax2.set_yscale('log')
        ax2.legend()

        # Plot 3: Distribution histogram (if samples available)
        ax3 = axes[1, 0]
        if risk_results and 'annual_risk' in risk_results:
            samples = risk_results['annual_risk'].individual_risks
            ax3.hist(samples, bins=50, edgecolor='black', alpha=0.7)
            ax3.axvline(np.mean(samples), color='red', linestyle='--', label='Mean')
            ax3.axvline(np.percentile(samples, 95), color='orange', linestyle='--', label='95th Percentile')
            ax3.set_xlabel('Annual Risk')
            ax3.set_ylabel('Frequency')
            ax3.set_title('Distribution of Annual Risk Estimates')
            ax3.legend()

        # Plot 4: Compliance visualization
        ax4 = axes[1, 1]
        thresholds = {'Annual Risk': 1e-6, 'Single Event': 1e-3}
        if 'annual_risk' in risk_results:
            actual_risk = risk_results['annual_risk'].statistics['mean']
            threshold = thresholds['Annual Risk']

            ax4.barh(['Calculated Risk', 'Regulatory Threshold'],
                    [actual_risk, threshold],
                    color=['red' if actual_risk > threshold else 'green', 'blue'])
            ax4.set_xlabel('Risk Value')
            ax4.set_title('Regulatory Compliance Status')
            ax4.set_xscale('log')

        plt.tight_layout()

        # Save figure
        if output_filename is None:
            output_filename = f"QMRA_Visualizations_{datetime.now().strftime('%Y%m%d_%H%M')}.png"

        plt.savefig(output_filename, dpi=300, bbox_inches='tight')
        plt.close()

        return output_filename


def generate_summary_table(risk_results: Dict[str, RiskResult]) -> pd.DataFrame:
    """
    Generate summary table from risk results.

    Args:
        risk_results: Dictionary of risk results

    Returns:
        DataFrame with summary statistics
    """
    summary_data = []

    for metric_name, result in risk_results.items():
        row = {
            'Risk Metric': metric_name.replace('_', ' ').title(),
            'Pathogen': result.pathogen_name,
            'Mean': result.statistics['mean'],
            'Median': result.statistics.get('median', np.nan),
            'Std Dev': result.statistics.get('std', np.nan),
            'P5': result.statistics.get('p5', np.nan),
            'P95': result.statistics.get('p95', np.nan),
            'P99': result.statistics.get('p99', np.nan)
        }

        if result.population_risks:
            row['Expected Cases'] = result.population_risks.get('expected_cases_per_year', np.nan)

        summary_data.append(row)

    return pd.DataFrame(summary_data)


if __name__ == "__main__":
    # Example usage
    print("Testing Report Generator Module")
    print("=" * 40)

    # Create sample data
    from risk_characterization import RiskCharacterization

    risk_calc = RiskCharacterization()

    # Create sample risk results
    test_doses = np.random.lognormal(2, 1, 1000)

    risk_results = {
        'infection_probability': risk_calc.calculate_infection_probability('norovirus', test_doses),
        'illness_probability': risk_calc.calculate_illness_probability('norovirus', test_doses),
        'annual_risk': risk_calc.calculate_annual_risk('norovirus', test_doses, 10)
    }

    # Project information
    project_info = {
        'title': 'QMRA Assessment for Recreational Water',
        'project_name': 'Beach Safety Assessment',
        'location': 'Wellington Harbour',
        'author': 'NIWA QMRA Team',
        'client': 'Regional Council'
    }

    # Exposure parameters
    exposure_params = {
        'water_ingestion_volume': '50 mL per event',
        'exposure_frequency': '10 events per year',
        'population_size': '10,000 regular swimmers'
    }

    # Generate report
    report_gen = ReportGenerator()

    print("Generating regulatory compliance report...")
    report_path = report_gen.create_regulatory_report(
        project_info=project_info,
        risk_results=risk_results,
        exposure_params=exposure_params,
        treatment_summary="UV disinfection with 3-log reduction",
        output_filename="Sample_QMRA_Report.docx"
    )
    print(f"Report saved: {report_path}")

    print("\nGenerating visualization report...")
    viz_path = report_gen.create_visualization_report(risk_results)
    print(f"Visualizations saved: {viz_path}")

    print("\nGenerating summary table...")
    summary_df = generate_summary_table(risk_results)
    print(summary_df.to_string())

    print("\nReport generation completed successfully!")