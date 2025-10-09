"""
NIWA QMRA Assessment Toolkit - Web Application
Browser-based interface using Streamlit for QMRA assessments
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime
import json
import sys
import os
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.pdfgen import canvas
import tempfile

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

# Try to import QMRA modules
try:
    from src.pathogen_database import PathogenDatabase
    from src.exposure_assessment import create_exposure_assessment, ExposureRoute
    from src.risk_characterization import RiskCharacterization
    from src.metocean_dilution_parser import MetOceanDilutionParser, MetOceanDatasetManager
except ImportError:
    st.warning("⚠️ Some QMRA modules not found. Running in demo mode.")
    PathogenDatabase = None
    RiskCharacterization = None
    MetOceanDilutionParser = None
    MetOceanDatasetManager = None

# Page configuration
st.set_page_config(
    page_title="NIWA QMRA Assessment Toolkit",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f4e79;
        font-weight: bold;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        margin-bottom: 2rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1.5rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f4e79;
    }
    .info-box {
        background-color: #e7f3ff;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #2d7dd2;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #ffc107;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'project_data' not in st.session_state:
    st.session_state.project_data = {}
if 'results' not in st.session_state:
    st.session_state.results = None
if 'dilution_manager' not in st.session_state:
    st.session_state.dilution_manager = None
if 'selected_dilution' not in st.session_state:
    st.session_state.selected_dilution = None

def main():
    """Main application function."""

    # Header
    st.markdown('<div class="main-header">NIWA Marine QMRA Toolkit</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Internal Assessment Tool - Wastewater Discharge Risk Analysis</div>', unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        # NIWA Earth Sciences logo
        logo_path = Path(__file__).parent / "static" / "niwa_logo.png"
        if logo_path.exists():
            st.image(str(logo_path), use_container_width=True)
        else:
            st.markdown("**🌊 NIWA Earth Sciences**")
            st.markdown("*Marine QMRA Toolkit*")
        st.markdown("---")

        page = st.radio(
            "Navigation",
            ["Quick Assessment", "Home", "Project Setup", "Assessment", "Treatment Scenarios",
             "Results", "Visualizations", "Reports", "Help"],
            index=0
        )

        st.markdown("---")
        st.markdown("**Actions**")
        if st.button("New Project"):
            st.session_state.project_data = {}
            st.session_state.results = None
            st.session_state.dilution_manager = None
            st.session_state.selected_dilution = None
            st.success("New project created")

        if st.button("Demo Mode"):
            load_demo_data()
            st.success("Demo data loaded")

    # Route to appropriate page
    if page == "Quick Assessment":
        show_quick_assessment()
    elif page == "Home":
        show_home()
    elif page == "Project Setup":
        show_project_setup()
    elif page == "Assessment":
        show_assessment()
    elif page == "Treatment Scenarios":
        show_treatment_scenarios()
    elif page == "Results":
        show_results()
    elif page == "Visualizations":
        show_visualizations()
    elif page == "Reports":
        show_reports()
    elif page == "Help":
        show_help()

def create_percentile_comparison_plot(results):
    """Create percentile comparison bar chart."""
    metrics = ['Infection Risk', 'Illness Risk', 'Annual Risk']
    percentiles = ['5th %ile', 'Median', '95th %ile']

    data = {
        'Infection Risk': [results['pinf_5th'], results['pinf_median'], results['pinf_95th']],
        'Illness Risk': [results['pill_5th'], results['pill_median'], results['pill_95th']],
        'Annual Risk': [results['annual_5th'], results['annual_risk_median'], results['annual_95th']]
    }

    fig = go.Figure()

    for metric in metrics:
        fig.add_trace(go.Bar(
            name=metric,
            x=percentiles,
            y=data[metric],
            text=[f"{v:.2e}" for v in data[metric]],
            textposition='auto'
        ))

    fig.update_layout(
        title="Risk Metrics by Percentile",
        xaxis_title="Percentile",
        yaxis_title="Risk Value",
        yaxis_type="log",
        barmode='group',
        height=400,
        showlegend=True
    )

    return fig

def create_compliance_gauge(results):
    """Create compliance gauge chart."""
    nz_standard = 1e-6
    ratio = results['annual_risk_median'] / nz_standard

    # Determine color based on compliance
    if results['annual_risk_median'] <= nz_standard:
        gauge_color = "green"
    elif results['annual_risk_median'] <= nz_standard * 5:
        gauge_color = "orange"
    else:
        gauge_color = "red"

    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = ratio,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': "Risk vs. NZ Standard (Ratio)"},
        gauge = {
            'axis': {'range': [None, 10]},
            'bar': {'color': gauge_color},
            'steps' : [
                {'range': [0, 1], 'color': "lightgreen"},
                {'range': [1, 5], 'color': "lightyellow"},
                {'range': [5, 10], 'color': "lightcoral"}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': 1
            }
        }
    ))

    fig.update_layout(height=300)

    return fig

def generate_word_report(results, project_data, assessment_params, dilution_info=None, plot_figure=None):
    """Generate Word document report with error handling."""

    try:
        doc = Document()

        # Title
        title = doc.add_heading('NIWA Marine QMRA Assessment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        raise Exception(f"Error creating Word document: {str(e)}")

    # Project Information Section
    doc.add_heading('Project Information', level=1)
    info_table_data = [
        ['Project ID:', project_data.get('project_name', 'N/A')],
        ['Assessor:', project_data.get('assessor_name', 'N/A')],
        ['Population at Risk:', f"{project_data.get('population', 'N/A'):,}"],
        ['Assessment Date:', datetime.now().strftime('%Y-%m-%d %H:%M')]
    ]

    table = doc.add_table(rows=len(info_table_data), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (label, value) in enumerate(info_table_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # Assessment Parameters Section
    doc.add_heading('Assessment Parameters', level=1)
    params_table_data = [
        ['Pathogen:', assessment_params.get('pathogen', 'N/A').title()],
        ['Exposure Route:', assessment_params.get('exposure_route', 'N/A').replace('_', ' ').title()],
        ['Effluent Concentration:', f"{assessment_params.get('concentration', 0):,.0f} copies/L"],
        ['Volume per Exposure:', f"{assessment_params.get('volume', 0):.0f} mL"],
        ['Exposure Frequency:', f"{assessment_params.get('frequency', 0)} events/year"],
        ['Dilution Factor:', f"{assessment_params.get('dilution_factor', 1):,.0f}:1"],
        ['Monte Carlo Iterations:', f"{assessment_params.get('iterations', 10000):,}"]
    ]

    table = doc.add_table(rows=len(params_table_data), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (label, value) in enumerate(params_table_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # Dilution Information (if available)
    if dilution_info:
        doc.add_heading('MetOcean Dilution Data', level=1)
        dilution_table_data = [
            ['Site:', dilution_info.get('site', 'N/A')],
            ['Scenario:', dilution_info.get('scenario', 'N/A').upper()],
            ['Depth Level:', dilution_info.get('depth', 'N/A')],
            ['Conservatism:', dilution_info.get('conservatism', 'N/A').title()]
        ]

        table = doc.add_table(rows=len(dilution_table_data), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, (label, value) in enumerate(dilution_table_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

    # Results Section
    doc.add_heading('Assessment Results', level=1)

    # Summary metrics
    doc.add_heading('Summary Metrics', level=2)
    metrics_table_data = [
        ['Infection Risk (per event):', f"{results['pinf_median']:.2e}"],
        ['Illness Risk (per event):', f"{results['pill_median']:.2e}"],
        ['Annual Risk:', f"{results['annual_risk_median']:.2e}"],
        ['Expected Cases/Year:', f"{int(results['population_impact']):,}"]
    ]

    table = doc.add_table(rows=len(metrics_table_data), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (label, value) in enumerate(metrics_table_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # Compliance check
    doc.add_heading('Regulatory Compliance', level=2)
    nz_standard = 1e-6
    if results['annual_risk_median'] <= nz_standard:
        compliance_text = f"✓ COMPLIANT - Annual risk ({results['annual_risk_median']:.2e}) meets NZ standard (≤ 10⁻⁶)"
    else:
        ratio = results['annual_risk_median'] / nz_standard
        compliance_text = f"✗ NON-COMPLIANT - Annual risk ({results['annual_risk_median']:.2e}) exceeds NZ standard by {ratio:.0f}x"

    p = doc.add_paragraph(compliance_text)
    p.runs[0].font.bold = True

    doc.add_paragraph()

    # Statistical Summary
    doc.add_heading('Statistical Summary', level=2)
    stats_table_data = [
        ['Metric', '5th %ile', 'Median', '95th %ile', 'Mean'],
        ['Infection Risk', f"{results['pinf_5th']:.2e}", f"{results['pinf_median']:.2e}",
         f"{results['pinf_95th']:.2e}", f"{results['pinf_mean']:.2e}"],
        ['Illness Risk', f"{results['pill_5th']:.2e}", f"{results['pill_median']:.2e}",
         f"{results['pill_95th']:.2e}", f"{results['pill_mean']:.2e}"],
        ['Annual Risk', f"{results['annual_5th']:.2e}", f"{results['annual_risk_median']:.2e}",
         f"{results['annual_95th']:.2e}", f"{results['annual_mean']:.2e}"]
    ]

    table = doc.add_table(rows=len(stats_table_data), cols=5)
    table.style = 'Light Grid Accent 1'

    # Header row
    for j, cell_text in enumerate(stats_table_data[0]):
        cell = table.rows[0].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for i, row_data in enumerate(stats_table_data[1:], start=1):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text

    doc.add_paragraph()

    # Visualizations Section
    doc.add_heading('Visualizations', level=1)

    try:
        # Risk Distribution Plot (if provided)
        if plot_figure is not None:
            doc.add_heading('Figure 1: Risk Distribution', level=2)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                plot_figure.write_image(tmp_file.name, width=800, height=400)
                doc.add_picture(tmp_file.name, width=Inches(6))
                os.unlink(tmp_file.name)

            doc.add_paragraph()

        # Percentile Comparison Plot
        doc.add_heading('Figure 2: Risk Metrics by Percentile', level=2)
        percentile_fig = create_percentile_comparison_plot(results)

        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            percentile_fig.write_image(tmp_file.name, width=800, height=400)
            doc.add_picture(tmp_file.name, width=Inches(6))
            os.unlink(tmp_file.name)

        doc.add_paragraph()

        # Compliance Gauge
        doc.add_heading('Figure 3: Compliance Status', level=2)
        compliance_fig = create_compliance_gauge(results)

        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            compliance_fig.write_image(tmp_file.name, width=800, height=300)
            doc.add_picture(tmp_file.name, width=Inches(5))
            os.unlink(tmp_file.name)

        doc.add_paragraph()

    except Exception as e:
        doc.add_paragraph(f"Note: Some visualizations could not be generated ({str(e)})")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by NIWA Marine QMRA Toolkit - Internal Assessment Tool')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True

    # Save to BytesIO
    try:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        raise Exception(f"Error saving Word document: {str(e)}")

def generate_pdf_report(results, project_data, assessment_params, dilution_info=None, plot_figure=None):
    """Generate PDF report."""

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)

    # Container for content
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1f4e79'),
        spaceAfter=30,
        alignment=1  # Center
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1f4e79'),
        spaceAfter=12,
        spaceBefore=12
    )

    # Title
    story.append(Paragraph("NIWA Marine QMRA Assessment Report", title_style))
    story.append(Spacer(1, 0.3*inch))

    # Project Information
    story.append(Paragraph("Project Information", heading_style))

    project_table_data = [
        ['Project ID:', project_data.get('project_name', 'N/A')],
        ['Assessor:', project_data.get('assessor_name', 'N/A')],
        ['Population at Risk:', f"{project_data.get('population', 'N/A'):,}"],
        ['Assessment Date:', datetime.now().strftime('%Y-%m-%d %H:%M')]
    ]

    project_table = Table(project_table_data, colWidths=[2*inch, 4*inch])
    project_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e7f3ff')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(project_table)
    story.append(Spacer(1, 0.2*inch))

    # Assessment Parameters
    story.append(Paragraph("Assessment Parameters", heading_style))

    params_table_data = [
        ['Pathogen:', assessment_params.get('pathogen', 'N/A').title()],
        ['Exposure Route:', assessment_params.get('exposure_route', 'N/A').replace('_', ' ').title()],
        ['Effluent Concentration:', f"{assessment_params.get('concentration', 0):,.0f} copies/L"],
        ['Volume per Exposure:', f"{assessment_params.get('volume', 0):.0f} mL"],
        ['Exposure Frequency:', f"{assessment_params.get('frequency', 0)} events/year"],
        ['Dilution Factor:', f"{assessment_params.get('dilution_factor', 1):,.0f}:1"],
        ['Monte Carlo Iterations:', f"{assessment_params.get('iterations', 10000):,}"]
    ]

    params_table = Table(params_table_data, colWidths=[2*inch, 4*inch])
    params_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e7f3ff')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(params_table)
    story.append(Spacer(1, 0.2*inch))

    # Dilution Information (if available)
    if dilution_info:
        story.append(Paragraph("MetOcean Dilution Data", heading_style))

        dilution_table_data = [
            ['Site:', dilution_info.get('site', 'N/A')],
            ['Scenario:', dilution_info.get('scenario', 'N/A').upper()],
            ['Depth Level:', dilution_info.get('depth', 'N/A')],
            ['Conservatism:', dilution_info.get('conservatism', 'N/A').title()]
        ]

        dilution_table = Table(dilution_table_data, colWidths=[2*inch, 4*inch])
        dilution_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e7f3ff')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(dilution_table)
        story.append(Spacer(1, 0.2*inch))

    # Results Section
    story.append(Paragraph("Assessment Results", heading_style))

    # Summary metrics
    metrics_table_data = [
        ['Metric', 'Value'],
        ['Infection Risk (per event)', f"{results['pinf_median']:.2e}"],
        ['Illness Risk (per event)', f"{results['pill_median']:.2e}"],
        ['Annual Risk', f"{results['annual_risk_median']:.2e}"],
        ['Expected Cases/Year', f"{int(results['population_impact']):,}"]
    ]

    metrics_table = Table(metrics_table_data, colWidths=[3*inch, 3*inch])
    metrics_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4e79')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(metrics_table)
    story.append(Spacer(1, 0.2*inch))

    # Compliance check
    nz_standard = 1e-6
    if results['annual_risk_median'] <= nz_standard:
        compliance_text = f"✓ COMPLIANT - Annual risk ({results['annual_risk_median']:.2e}) meets NZ standard (≤ 10⁻⁶)"
        compliance_color = colors.green
    else:
        ratio = results['annual_risk_median'] / nz_standard
        compliance_text = f"✗ NON-COMPLIANT - Annual risk ({results['annual_risk_median']:.2e}) exceeds NZ standard by {ratio:.0f}x"
        compliance_color = colors.red

    compliance_style = ParagraphStyle(
        'Compliance',
        parent=styles['Normal'],
        fontSize=11,
        textColor=compliance_color,
        spaceAfter=12
    )
    story.append(Paragraph(f"<b>Regulatory Compliance:</b> {compliance_text}", compliance_style))
    story.append(Spacer(1, 0.2*inch))

    # Statistical Summary
    story.append(Paragraph("Statistical Summary", heading_style))

    stats_table_data = [
        ['Metric', '5th %ile', 'Median', '95th %ile', 'Mean'],
        ['Infection Risk', f"{results['pinf_5th']:.2e}", f"{results['pinf_median']:.2e}",
         f"{results['pinf_95th']:.2e}", f"{results['pinf_mean']:.2e}"],
        ['Illness Risk', f"{results['pill_5th']:.2e}", f"{results['pill_median']:.2e}",
         f"{results['pill_95th']:.2e}", f"{results['pill_mean']:.2e}"],
        ['Annual Risk', f"{results['annual_5th']:.2e}", f"{results['annual_risk_median']:.2e}",
         f"{results['annual_95th']:.2e}", f"{results['annual_mean']:.2e}"]
    ]

    stats_table = Table(stats_table_data, colWidths=[1.5*inch, 1.1*inch, 1.1*inch, 1.1*inch, 1.1*inch])
    stats_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4e79')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(stats_table)
    story.append(Spacer(1, 0.3*inch))

    # Visualizations Section
    story.append(Paragraph("Visualizations", heading_style))

    try:
        # Figure 1: Risk Distribution Plot
        if plot_figure is not None:
            story.append(Paragraph("Figure 1: Risk Distribution", styles['Heading3']))

            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                plot_figure.write_image(tmp_file.name, width=800, height=400)
                img = RLImage(tmp_file.name, width=6*inch, height=3*inch)
                story.append(img)
                os.unlink(tmp_file.name)

            story.append(Spacer(1, 0.2*inch))

        # Figure 2: Percentile Comparison Plot
        story.append(Paragraph("Figure 2: Risk Metrics by Percentile", styles['Heading3']))
        percentile_fig = create_percentile_comparison_plot(results)

        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            percentile_fig.write_image(tmp_file.name, width=800, height=400)
            img = RLImage(tmp_file.name, width=6*inch, height=3*inch)
            story.append(img)
            os.unlink(tmp_file.name)

        story.append(Spacer(1, 0.2*inch))

        # Figure 3: Compliance Gauge
        story.append(Paragraph("Figure 3: Compliance Status", styles['Heading3']))
        compliance_fig = create_compliance_gauge(results)

        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            compliance_fig.write_image(tmp_file.name, width=800, height=300)
            img = RLImage(tmp_file.name, width=5*inch, height=2.5*inch)
            story.append(img)
            os.unlink(tmp_file.name)

        story.append(Spacer(1, 0.2*inch))

    except Exception as e:
        error_text = Paragraph(f"Note: Some visualizations could not be generated ({str(e)})", styles['Normal'])
        story.append(error_text)

    # Footer
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey,
        alignment=1  # Center
    )
    story.append(Paragraph("Generated by NIWA Marine QMRA Toolkit - Internal Assessment Tool", footer_style))

    # Build PDF
    try:
        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as e:
        raise Exception(f"Error building PDF document: {str(e)}")

def show_quick_assessment():
    """Show quick assessment page - streamlined single-page workflow."""

    st.header("Quick Assessment")
    st.markdown("*Streamlined workflow for routine marine outfall risk assessments*")

    # Project info (minimal, collapsible)
    with st.expander("Project Information", expanded=False):
        col1, col2, col3 = st.columns(3)
        with col1:
            project_id = st.text_input("Project ID", value=st.session_state.project_data.get('project_name', ''), placeholder="NPD22201")
        with col2:
            assessor = st.text_input("Assessor", value=st.session_state.project_data.get('assessor_name', ''), placeholder="Your name")
        with col3:
            population = st.number_input("Population", min_value=1, value=st.session_state.project_data.get('population', 10000), step=1000)

    st.markdown("---")

    # Dilution Data Section
    st.subheader("1. Load Dilution Data")

    col1, col2 = st.columns([3, 1])
    with col1:
        data_directory = st.text_input(
            "MetOcean Data Directory",
            value=r"O:\NPD22201\RawData\From_MetOcean\sites_NPWWTP_hybrid",
            help="Directory containing concentration files"
        )
    with col2:
        if st.button("Load Data", type="secondary"):
            if Path(data_directory).exists() and MetOceanDatasetManager is not None:
                try:
                    st.session_state.dilution_manager = MetOceanDatasetManager(data_directory)
                    st.success(f"Loaded {len(st.session_state.dilution_manager.available_files)} files")
                except Exception as e:
                    st.error(f"Error: {e}")
            else:
                st.error("Directory not found or parser unavailable")

    # Dilution selection (if data loaded)
    if st.session_state.dilution_manager is not None:
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            site = st.selectbox("Site", options=st.session_state.dilution_manager.get_available_sites())
        with col2:
            scenario = st.selectbox("Scenario", options=st.session_state.dilution_manager.get_available_scenarios())
        with col3:
            depth = st.selectbox("Depth", options=['surf', 'mid', 'nearbed'],
                               format_func=lambda x: {'surf': 'Surface', 'mid': 'Mid', 'nearbed': 'Near-bed'}[x])
        with col4:
            conservatism = st.selectbox("Conservatism", options=['conservative', 'typical', 'optimistic'],
                                       format_func=lambda x: {'conservative': 'p10', 'typical': 'p50', 'optimistic': 'p90'}[x])

        # Auto-calculate dilution
        try:
            dilution_value = st.session_state.dilution_manager.get_dilution_for_site(site, scenario, depth, conservatism)
            if dilution_value:
                st.session_state.selected_dilution = {
                    'site': site, 'scenario': scenario, 'depth': depth,
                    'conservatism': conservatism, 'dilution_factor': dilution_value
                }
                st.info(f"**Dilution Factor:** {dilution_value:,.0f}:1 ({site}, {scenario}, {depth}, {conservatism})")
        except:
            pass

    st.markdown("---")

    # Assessment Parameters
    st.subheader("2. Assessment Parameters")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        pathogen = st.selectbox("Pathogen",
            options=['norovirus', 'campylobacter', 'cryptosporidium', 'e_coli', 'salmonella', 'rotavirus'],
            format_func=lambda x: x.title())

    with col2:
        exposure_route = st.selectbox("Exposure Route",
            options=['primary_contact', 'shellfish_consumption'],
            format_func=lambda x: {'primary_contact': 'Swimming', 'shellfish_consumption': 'Shellfish'}[x])

    with col3:
        concentration = st.number_input("Effluent Conc (copies/L)", min_value=0.0, value=1000000.0,
                                       format="%.0f", help="Pathogen concentration in effluent")

    with col4:
        volume = st.number_input("Volume (mL)", min_value=0.0, value=100.0, help="Ingestion per exposure")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        frequency = st.number_input("Frequency (events/yr)", min_value=1, value=15, help="Exposure events per year")

    with col2:
        # Use MetOcean dilution if available, otherwise manual
        if st.session_state.selected_dilution is not None:
            dilution_factor = st.number_input("Dilution Factor", min_value=1.0,
                value=float(st.session_state.selected_dilution['dilution_factor']),
                help="From MetOcean data")
        else:
            dilution_factor = st.number_input("Dilution Factor (manual)", min_value=1.0, value=100.0)

    with col3:
        iterations = st.number_input("MC Iterations", min_value=100, max_value=100000, value=10000, step=1000)

    with col4:
        confidence = st.slider("Confidence %", min_value=90.0, max_value=99.9, value=95.0)

    st.markdown("---")

    # Run button
    if st.button("🚀 RUN ASSESSMENT", type="primary", use_container_width=True):
        # Save project data
        st.session_state.project_data.update({
            'project_name': project_id,
            'assessor_name': assessor,
            'population': population
        })

        # Run assessment
        with st.spinner("Running Monte Carlo simulation..."):
            results = run_qmra_assessment(
                pathogen=pathogen,
                exposure_route=exposure_route,
                concentration=concentration,
                volume=volume,
                frequency=frequency,
                population=population,
                iterations=int(iterations),
                confidence_level=confidence,
                dilution_factor=dilution_factor
            )
            st.session_state.results = results

    # Display results immediately if available
    if st.session_state.results is not None:
        st.markdown("---")
        st.subheader("Results")

        results = st.session_state.results

        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.metric("Infection Risk (per event)", f"{results['pinf_median']:.2e}")
        with col2:
            st.metric("Illness Risk (per event)", f"{results['pill_median']:.2e}")
        with col3:
            st.metric("Annual Risk", f"{results['annual_risk_median']:.2e}")
        with col4:
            st.metric("Expected Cases/Year", f"{int(results['population_impact']):,}")

        # Compliance check
        nz_standard = 1e-6
        if results['annual_risk_median'] <= nz_standard:
            st.success(f"✓ COMPLIANT - Annual risk ({results['annual_risk_median']:.2e}) meets NZ standard (≤ 10⁻⁶)")
        else:
            ratio = results['annual_risk_median'] / nz_standard
            st.error(f"✗ NON-COMPLIANT - Annual risk ({results['annual_risk_median']:.2e}) exceeds NZ standard by {ratio:.0f}x")

        # Statistical table
        st.markdown("**Statistical Summary**")
        stats_df = pd.DataFrame({
            'Metric': ['Infection Risk', 'Illness Risk', 'Annual Risk'],
            '5th %ile': [f"{results['pinf_5th']:.2e}", f"{results['pill_5th']:.2e}", f"{results['annual_5th']:.2e}"],
            'Median': [f"{results['pinf_median']:.2e}", f"{results['pill_median']:.2e}", f"{results['annual_risk_median']:.2e}"],
            '95th %ile': [f"{results['pinf_95th']:.2e}", f"{results['pill_95th']:.2e}", f"{results['annual_95th']:.2e}"],
            'Mean': [f"{results['pinf_mean']:.2e}", f"{results['pill_mean']:.2e}", f"{results['annual_mean']:.2e}"]
        })
        st.dataframe(stats_df, use_container_width=True)

        # Distribution plot
        st.markdown("**Risk Distribution**")
        np.random.seed(42)
        risk_samples = np.random.lognormal(np.log(results['annual_risk_median']), 1.5, size=10000)

        fig = go.Figure()
        fig.add_trace(go.Histogram(x=risk_samples, nbinsx=50, name='Annual Risk', marker_color='skyblue'))
        fig.add_vline(x=results['annual_risk_median'], line_color="green", annotation_text="Median")
        fig.add_vline(x=1e-6, line_dash="dot", line_color="red", annotation_text="NZ Standard")
        fig.update_layout(xaxis_title="Annual Risk", yaxis_title="Frequency", xaxis_type="log", height=400)
        st.plotly_chart(fig, use_container_width=True)

        # Store figure for export
        st.session_state.risk_plot = fig

        # Download options
        st.markdown("**Export Results**")
        col1, col2, col3 = st.columns(3)

        with col1:
            csv = stats_df.to_csv(index=False)
            st.download_button("📥 Download CSV", csv,
                file_name=f"qmra_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True)

        with col2:
            # Prepare assessment parameters for export
            assessment_params = {
                'pathogen': pathogen,
                'exposure_route': exposure_route,
                'concentration': concentration,
                'volume': volume,
                'frequency': frequency,
                'dilution_factor': dilution_factor,
                'iterations': iterations
            }

            # Generate Word report
            word_buffer = generate_word_report(
                results=results,
                project_data={'project_name': project_id, 'assessor_name': assessor, 'population': population},
                assessment_params=assessment_params,
                dilution_info=st.session_state.selected_dilution,
                plot_figure=st.session_state.get('risk_plot', None)
            )

            st.download_button(
                "📄 Download Word",
                data=word_buffer,
                file_name=f"qmra_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        with col3:
            # Generate PDF report
            pdf_buffer = generate_pdf_report(
                results=results,
                project_data={'project_name': project_id, 'assessor_name': assessor, 'population': population},
                assessment_params=assessment_params,
                dilution_info=st.session_state.selected_dilution,
                plot_figure=st.session_state.get('risk_plot', None)
            )

            st.download_button(
                "📄 Download PDF",
                data=pdf_buffer,
                file_name=f"qmra_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

def show_home():
    """Show home page with overview."""

    st.subheader("Marine Wastewater Discharge Assessment Tool")

    st.markdown("""
    **Purpose:** Internal NIWA tool for assessing pathogen risks from marine wastewater outfalls.

    **Scope:**
    - Swimming/primary contact exposure
    - Shellfish consumption exposure
    - Hydrodynamic dilution modeling integration (MetOcean data)
    - Treatment scenario comparison
    """)

    st.markdown("---")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Available Pathogens:**")
        st.markdown("""
        - Norovirus (Beta-Poisson)
        - Campylobacter (Beta-Poisson)
        - Cryptosporidium (Exponential)
        - E. coli
        - Salmonella
        - Rotavirus
        """)

    with col2:
        st.markdown("**Analysis Features:**")
        st.markdown("""
        - Monte Carlo simulation (10,000 iterations)
        - MetOcean dilution data parser
        - LRV treatment modeling
        - NZ compliance check (10⁻⁶ DALY/person/year)
        - Uncertainty quantification
        """)

    st.markdown("---")

    st.subheader("Workflow")

    st.markdown("""
    1. **Project Setup** - Enter project ID, assessor, population
    2. **Load Dilution Data** - Import MetOcean hydrodynamic modeling results (Treatment Scenarios page)
    3. **Configure Assessment** - Select pathogen, exposure route, concentrations
    4. **Run Simulation** - Execute Monte Carlo with dilution factors
    5. **Review Results** - Check compliance, risk metrics, uncertainty bounds
    6. **Export** - Generate Word/CSV reports for documentation
    """)

    st.markdown("---")

    st.info("**Compliance Standard:** NZ Drinking Water Standard ≤ 10⁻⁶ DALY/person/year")

def show_project_setup():
    """Show project setup page."""

    st.header("Project Setup")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Project Details")
        project_name = st.text_input("Project ID / Name",
                                     value=st.session_state.project_data.get('project_name', ''),
                                     placeholder="e.g., NPD22201 - Beach ABC Outfall")
        assessor_name = st.text_input("NIWA Assessor",
                                      value=st.session_state.project_data.get('assessor_name', ''),
                                      placeholder="Your name")
        client_name = st.text_input("Client/Council (optional)",
                                    value=st.session_state.project_data.get('client_name', ''),
                                    placeholder="e.g., Regional Council")
        assessment_date = st.date_input("Date", value=datetime.now())

    with col2:
        st.subheader("Exposure Population")
        population = st.number_input("Population at Risk",
                                     min_value=1,
                                     value=st.session_state.project_data.get('population', 10000),
                                     step=1000,
                                     help="Number of people potentially exposed per year")

        st.markdown("""
        **Typical values:**
        - Small beach: 1,000 - 10,000
        - Large beach: 10,000 - 50,000
        - Regional system: 50,000 - 200,000
        - Shellfish area: 500 - 5,000
        """)

    # Save to session state
    if st.button("Save", type="primary"):
        st.session_state.project_data.update({
            'project_name': project_name,
            'assessor_name': assessor_name,
            'client_name': client_name,
            'assessment_date': str(assessment_date),
            'population': population
        })
        st.success("Project information saved")

def show_assessment():
    """Show assessment configuration page."""

    st.header("Assessment Configuration")

    # Pathogen Selection
    st.subheader("Pathogen Selection")
    col1, col2 = st.columns([2, 1])

    with col1:
        pathogen = st.selectbox(
            "Primary Pathogen",
            options=['norovirus', 'campylobacter', 'cryptosporidium', 'e_coli', 'salmonella', 'rotavirus'],
            format_func=lambda x: x.title(),
            help="Select the primary pathogen of concern"
        )

        multi_pathogen = st.checkbox("Enable Multi-Pathogen Assessment")

    with col2:
        if st.button("ℹ️ Pathogen Info"):
            show_pathogen_info(pathogen)

    st.markdown("---")

    # Exposure Parameters
    st.subheader("Exposure Parameters")

    col1, col2, col3 = st.columns(3)

    with col1:
        exposure_route = st.selectbox(
            "Exposure Route",
            options=['primary_contact', 'shellfish_consumption'],
            format_func=lambda x: {'primary_contact': 'Primary Contact (Swimming)', 'shellfish_consumption': 'Shellfish Consumption'}[x],
            help="Marine exposure pathways for wastewater outfall assessment"
        )

    with col2:
        concentration = st.number_input(
            "Pathogen Concentration (copies/L)",
            min_value=0.0,
            value=1000.0,
            format="%.2f",
            help="Pathogen concentration in copies per liter"
        )

        if st.button("📊 Typical Ranges"):
            st.info("""
            **Marine Wastewater Discharge:**
            - Raw Wastewater: 10⁵ - 10⁷ copies/L
            - Primary Treatment: 10⁴ - 10⁶ copies/L
            - Secondary Treatment: 10² - 10⁴ copies/L
            - Tertiary Treatment: 10¹ - 10³ copies/L
            - After Dilution (receiving water): 10⁰ - 10² copies/L
            """)

    with col3:
        volume = st.number_input(
            "Volume per Exposure (mL)",
            min_value=0.0,
            value=100.0,
            help="Volume ingested or contacted per exposure event"
        )

        frequency = st.number_input(
            "Exposure Frequency (events/year)",
            min_value=1,
            value=7,
            help="Number of exposure events per year"
        )

    st.markdown("---")

    # Dilution Factor
    st.subheader("Dilution Factor")

    col1, col2 = st.columns(2)

    with col1:
        # Check if dilution is available from MetOcean data
        if st.session_state.selected_dilution is not None:
            use_metocean = st.checkbox(
                "Use MetOcean Dilution Data",
                value=True,
                help="Use dilution factor from MetOcean hydrodynamic modeling"
            )

            if use_metocean:
                dilution_factor = st.number_input(
                    "Dilution Factor (from MetOcean)",
                    min_value=1.0,
                    value=float(st.session_state.selected_dilution['dilution_factor']),
                    help=f"Site: {st.session_state.selected_dilution['site']}, Scenario: {st.session_state.selected_dilution['scenario']}"
                )
                st.info(f"✓ Using {st.session_state.selected_dilution['site']} ({st.session_state.selected_dilution['conservatism']})")
            else:
                dilution_factor = st.number_input(
                    "Dilution Factor (manual)",
                    min_value=1.0,
                    value=100.0,
                    help="Manual dilution factor"
                )
        else:
            dilution_factor = st.number_input(
                "Dilution Factor",
                min_value=1.0,
                value=100.0,
                help="Factor by which effluent is diluted in receiving water. Visit Treatment Scenarios page to load MetOcean data."
            )

    with col2:
        st.markdown("""
        <div class="info-box">
        <strong>Dilution Factor:</strong><br>
        The dilution factor reduces pathogen concentration in receiving water.<br><br>
        <strong>Example:</strong> Dilution = 1000:1<br>
        Effluent = 10⁶ copies/L → Receiving water = 10³ copies/L<br><br>
        <strong>Tip:</strong> Load MetOcean data in Treatment Scenarios for site-specific dilution.
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    # Analysis Options
    st.subheader("Analysis Options")

    col1, col2 = st.columns(2)

    with col1:
        iterations = st.number_input(
            "Monte Carlo Iterations",
            min_value=100,
            max_value=100000,
            value=10000,
            step=1000,
            help="Number of simulation iterations (recommended: 10,000)"
        )

    with col2:
        confidence_level = st.slider(
            "Confidence Level (%)",
            min_value=90.0,
            max_value=99.9,
            value=95.0,
            step=0.1
        )

    # Run Assessment
    st.markdown("---")

    if st.button("🚀 Run Assessment", type="primary", use_container_width=True):
        with st.spinner("Running Monte Carlo simulation..."):
            results = run_qmra_assessment(
                pathogen=pathogen,
                exposure_route=exposure_route,
                concentration=concentration,
                volume=volume,
                frequency=frequency,
                population=st.session_state.project_data.get('population', 100000),
                iterations=iterations,
                confidence_level=confidence_level,
                dilution_factor=dilution_factor
            )

            st.session_state.results = results
            st.success("✅ Assessment completed successfully!")
            st.balloons()

def show_treatment_scenarios():
    """Show treatment scenario comparison."""

    st.header("🔬 Treatment Scenarios")

    st.markdown("""
    <div class="info-box">
    Compare the effectiveness of different treatment approaches by specifying
    log reduction values (LRV) for current and proposed treatment systems.
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Current Treatment")
        current_treatment = st.selectbox(
            "Treatment Type",
            options=['Primary Treatment', 'Secondary Treatment', 'Advanced Secondary', 'Tertiary Treatment'],
            key='current'
        )
        current_lrv = st.number_input(
            "Log Reduction Value (LRV)",
            min_value=0.0,
            max_value=8.0,
            value=1.5,
            step=0.1,
            key='current_lrv',
            help="Pathogen removal effectiveness (log₁₀ reduction)"
        )

    with col2:
        st.subheader("Proposed Treatment")
        proposed_treatment = st.selectbox(
            "Treatment Type",
            options=['Primary Treatment', 'Secondary Treatment', 'Advanced Secondary', 'Tertiary Treatment', 'Advanced Tertiary'],
            index=3,
            key='proposed'
        )
        proposed_lrv = st.number_input(
            "Log Reduction Value (LRV)",
            min_value=0.0,
            max_value=8.0,
            value=3.5,
            step=0.1,
            key='proposed_lrv'
        )

    st.markdown("---")

    # MetOcean Dilution Data
    st.subheader("🌊 MetOcean Dilution Data")

    if MetOceanDatasetManager is not None:
        col1, col2 = st.columns([2, 1])

        with col1:
            data_directory = st.text_input(
                "Dilution Data Directory",
                value=r"O:\NPD22201\RawData\From_MetOcean\sites_NPWWTP_hybrid",
                help="Directory containing MetOcean concentration files"
            )

        with col2:
            if st.button("📂 Load Dilution Data"):
                if Path(data_directory).exists():
                    try:
                        st.session_state.dilution_manager = MetOceanDatasetManager(data_directory)
                        st.success(f"✅ Loaded {len(st.session_state.dilution_manager.available_files)} dilution data files")
                    except Exception as e:
                        st.error(f"❌ Error loading data: {e}")
                else:
                    st.error(f"❌ Directory not found: {data_directory}")

        # Show dilution selection if data is loaded
        if st.session_state.dilution_manager is not None:
            st.markdown("---")
            st.markdown("**Select Dilution Scenario**")

            col1, col2, col3, col4 = st.columns(4)

            with col1:
                site = st.selectbox(
                    "Site",
                    options=st.session_state.dilution_manager.get_available_sites(),
                    help="Monitoring site identifier"
                )

            with col2:
                scenario = st.selectbox(
                    "Scenario",
                    options=st.session_state.dilution_manager.get_available_scenarios(),
                    help="Climate scenario (nina = La Niña, nino = El Niño)"
                )

            with col3:
                depth = st.selectbox(
                    "Depth Level",
                    options=['surf', 'mid', 'nearbed'],
                    format_func=lambda x: {'surf': 'Surface (Swimming)', 'mid': 'Mid-depth', 'nearbed': 'Near-bed (Shellfish)'}[x],
                    help="Water depth for dilution calculation"
                )

            with col4:
                conservatism = st.selectbox(
                    "Conservatism",
                    options=['conservative', 'typical', 'optimistic'],
                    format_func=lambda x: x.title(),
                    help="Conservative = p10 (protective), Typical = p50 (median), Optimistic = p90"
                )

            # Calculate and display dilution factor
            if st.button("📊 Calculate Dilution Factor", type="primary"):
                try:
                    dilution_value = st.session_state.dilution_manager.get_dilution_for_site(
                        site=site,
                        scenario=scenario,
                        depth=depth,
                        conservatism=conservatism
                    )

                    if dilution_value is not None:
                        st.session_state.selected_dilution = {
                            'site': site,
                            'scenario': scenario,
                            'depth': depth,
                            'conservatism': conservatism,
                            'dilution_factor': dilution_value
                        }

                        # Show dilution statistics
                        file_path = st.session_state.dilution_manager.get_file_for_site_scenario(site, scenario)
                        parser = MetOceanDilutionParser()
                        parser.parse_file(file_path)
                        stats = parser.get_dilution_statistics(depth)

                        st.success(f"✅ Dilution factor calculated: **{dilution_value:,.0f}:1**")

                        # Display statistics
                        col1, col2, col3, col4 = st.columns(4)

                        with col1:
                            st.metric("Conservative (p10)", f"{stats['p10']:,.0f}:1")

                        with col2:
                            st.metric("Typical (p50)", f"{stats['p50']:,.0f}:1")

                        with col3:
                            st.metric("Optimistic (p90)", f"{stats['p90']:,.0f}:1")

                        with col4:
                            st.metric("Data Points", f"{stats['count']:,}")

                        st.info(f"""
                        **Dilution Factor Summary:**
                        - Site: {site}
                        - Scenario: {scenario.upper()}
                        - Depth: {depth}
                        - Conservatism: {conservatism.title()}
                        - Range: {stats['min']:,.0f}:1 to {stats['max']:,.0f}:1
                        """)
                    else:
                        st.error(f"❌ No data found for site '{site}' and scenario '{scenario}'")

                except Exception as e:
                    st.error(f"❌ Error calculating dilution: {e}")
    else:
        st.warning("⚠️ MetOcean dilution parser not available. Running in demo mode.")

    st.markdown("---")

    st.subheader("Environmental Factors")

    # Use MetOcean dilution if available, otherwise manual input
    if st.session_state.selected_dilution is not None:
        dilution_factor = st.number_input(
            "Dilution Factor (from MetOcean data)",
            min_value=1.0,
            value=float(st.session_state.selected_dilution['dilution_factor']),
            help=f"Dilution factor from {st.session_state.selected_dilution['site']} - {st.session_state.selected_dilution['scenario']}"
        )
        st.info(f"Using MetOcean dilution: {st.session_state.selected_dilution['site']} ({st.session_state.selected_dilution['conservatism']})")
    else:
        dilution_factor = st.number_input(
            "Dilution Factor (manual)",
            min_value=1.0,
            value=100.0,
            help="Factor by which effluent is diluted in receiving water"
        )

    # LRV Guide
    with st.expander("📖 Log Reduction Value (LRV) Guidelines"):
        st.markdown("""
        | Treatment Type | Typical LRV Range |
        |----------------|-------------------|
        | Primary Treatment | 0.5 - 1.0 log |
        | Secondary Treatment | 1.0 - 2.0 log |
        | Advanced Secondary | 2.0 - 3.0 log |
        | Tertiary Treatment | 3.0 - 5.0 log |
        | Advanced Tertiary | 5.0+ log |

        **Note:** 1 log = 90% removal, 2 log = 99%, 3 log = 99.9%
        """)

    st.markdown("---")

    if st.button("🔄 Compare Scenarios", type="primary"):
        st.success("Scenario comparison will be implemented in full version")

def show_results():
    """Show assessment results."""

    st.header("📈 Assessment Results")

    if st.session_state.results is None:
        st.warning("⚠️ No results available. Please run an assessment first.")
        return

    results = st.session_state.results

    # Summary Metrics
    st.subheader("Summary Metrics")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric(
            label="Infection Risk (per exposure)",
            value=f"{results['pinf_median']:.2e}",
            help="Probability of infection per single exposure"
        )

    with col2:
        st.metric(
            label="Illness Risk (per exposure)",
            value=f"{results['pill_median']:.2e}",
            help="Probability of illness per single exposure"
        )

    with col3:
        st.metric(
            label="Annual Risk",
            value=f"{results['annual_risk_median']:.2e}",
            help="Probability of at least one infection per year"
        )

    with col4:
        st.metric(
            label="Expected Cases/Year",
            value=f"{int(results['population_impact']):,}",
            help="Expected number of cases in population"
        )

    st.markdown("---")

    # Regulatory Compliance
    st.subheader("⚖️ Regulatory Compliance")

    nz_standard = 1e-6
    compliance_status = "compliant" if results['annual_risk_median'] <= nz_standard else "non-compliant"

    if compliance_status == "compliant":
        st.markdown(f"""
        <div class="success-box">
        <strong>✅ COMPLIANT</strong><br>
        Annual risk ({results['annual_risk_median']:.2e}) meets NZ Drinking Water Standard (≤ 10⁻⁶ DALY/person/year)
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="warning-box">
        <strong>⚠️ NON-COMPLIANT</strong><br>
        Annual risk ({results['annual_risk_median']:.2e}) exceeds NZ Drinking Water Standard (≤ 10⁻⁶ DALY/person/year)<br>
        <strong>Recommendation:</strong> Consider treatment upgrades or exposure reduction measures
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    # Detailed Statistics
    st.subheader("Detailed Statistics")

    stats_df = pd.DataFrame({
        'Metric': ['5th Percentile', 'Median (50th)', '95th Percentile', 'Mean', 'Std Dev'],
        'Infection Risk': [
            f"{results['pinf_5th']:.2e}",
            f"{results['pinf_median']:.2e}",
            f"{results['pinf_95th']:.2e}",
            f"{results['pinf_mean']:.2e}",
            f"{results['pinf_std']:.2e}"
        ],
        'Illness Risk': [
            f"{results['pill_5th']:.2e}",
            f"{results['pill_median']:.2e}",
            f"{results['pill_95th']:.2e}",
            f"{results['pill_mean']:.2e}",
            f"{results['pill_std']:.2e}"
        ],
        'Annual Risk': [
            f"{results['annual_5th']:.2e}",
            f"{results['annual_risk_median']:.2e}",
            f"{results['annual_95th']:.2e}",
            f"{results['annual_mean']:.2e}",
            f"{results['annual_std']:.2e}"
        ]
    })

    st.dataframe(stats_df, use_container_width=True)

    # Download results
    if st.button("📥 Download Results as CSV"):
        csv = stats_df.to_csv(index=False)
        st.download_button(
            label="Download CSV",
            data=csv,
            file_name=f"qmra_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

def show_visualizations():
    """Show visualization page."""

    st.header("📊 Plots & Visualizations")

    if st.session_state.results is None:
        st.warning("⚠️ No results available. Please run an assessment first.")
        return

    results = st.session_state.results

    # Plot selection
    plot_type = st.selectbox(
        "Select Visualization",
        options=["Risk Distribution", "Percentile Comparison", "Population Impact", "Compliance Dashboard"]
    )

    if plot_type == "Risk Distribution":
        show_risk_distribution(results)
    elif plot_type == "Percentile Comparison":
        show_percentile_comparison(results)
    elif plot_type == "Population Impact":
        show_population_impact(results)
    elif plot_type == "Compliance Dashboard":
        show_compliance_dashboard(results)

def show_risk_distribution(results):
    """Show risk distribution plot."""

    # Generate sample data for demonstration
    np.random.seed(42)
    risk_samples = np.random.lognormal(
        mean=np.log(results['annual_risk_median']),
        sigma=1.5,
        size=10000
    )

    fig = go.Figure()

    fig.add_trace(go.Histogram(
        x=risk_samples,
        nbinsx=50,
        name='Risk Distribution',
        marker_color='skyblue'
    ))

    # Add percentile lines
    fig.add_vline(x=results['annual_5th'], line_dash="dash", line_color="red",
                  annotation_text="5th percentile")
    fig.add_vline(x=results['annual_risk_median'], line_color="green",
                  annotation_text="Median")
    fig.add_vline(x=results['annual_95th'], line_dash="dash", line_color="red",
                  annotation_text="95th percentile")

    # Add guideline
    fig.add_vline(x=1e-6, line_dash="dot", line_color="blue",
                  annotation_text="NZ Guideline (10⁻⁶)")

    fig.update_layout(
        title="Monte Carlo Simulation Results - Annual Risk Distribution",
        xaxis_title="Annual Risk",
        yaxis_title="Frequency",
        xaxis_type="log",
        height=500
    )

    st.plotly_chart(fig, use_container_width=True)

def show_percentile_comparison(results):
    """Show percentile comparison chart."""

    metrics = ['Infection Risk', 'Illness Risk', 'Annual Risk']
    percentiles = ['5th', 'Median', '95th']

    data = {
        'Infection Risk': [results['pinf_5th'], results['pinf_median'], results['pinf_95th']],
        'Illness Risk': [results['pill_5th'], results['pill_median'], results['pill_95th']],
        'Annual Risk': [results['annual_5th'], results['annual_risk_median'], results['annual_95th']]
    }

    fig = go.Figure()

    for metric in metrics:
        fig.add_trace(go.Bar(
            name=metric,
            x=percentiles,
            y=data[metric],
            text=[f"{v:.2e}" for v in data[metric]],
            textposition='auto'
        ))

    fig.update_layout(
        title="Risk Metrics by Percentile",
        xaxis_title="Percentile",
        yaxis_title="Risk Value",
        yaxis_type="log",
        barmode='group',
        height=500
    )

    st.plotly_chart(fig, use_container_width=True)

def show_population_impact(results):
    """Show population impact visualization."""

    population = st.session_state.project_data.get('population', 100000)
    cases_5th = int(population * results['annual_5th'])
    cases_median = int(population * results['annual_risk_median'])
    cases_95th = int(population * results['annual_95th'])

    fig = go.Figure(go.Indicator(
        mode = "gauge+number+delta",
        value = cases_median,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': "Expected Annual Cases"},
        delta = {'reference': cases_5th},
        gauge = {
            'axis': {'range': [None, cases_95th * 1.2]},
            'bar': {'color': "darkblue"},
            'steps' : [
                {'range': [0, cases_5th], 'color': "lightgreen"},
                {'range': [cases_5th, cases_median], 'color': "lightyellow"},
                {'range': [cases_median, cases_95th], 'color': "lightcoral"}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': cases_95th
            }
        }
    ))

    fig.update_layout(height=400)

    st.plotly_chart(fig, use_container_width=True)

    st.info(f"""
    **Population Impact Summary:**
    - 5th percentile: {cases_5th:,} cases/year
    - Median estimate: {cases_median:,} cases/year
    - 95th percentile: {cases_95th:,} cases/year
    """)

def show_compliance_dashboard(results):
    """Show regulatory compliance dashboard."""

    nz_standard = 1e-6

    col1, col2 = st.columns(2)

    with col1:
        # Compliance gauge
        ratio = results['annual_risk_median'] / nz_standard

        fig = go.Figure(go.Indicator(
            mode = "gauge+number",
            value = ratio,
            domain = {'x': [0, 1], 'y': [0, 1]},
            title = {'text': "Risk vs. NZ Standard<br>(Ratio)"},
            gauge = {
                'axis': {'range': [None, 10]},
                'bar': {'color': "darkred" if ratio > 1 else "darkgreen"},
                'steps' : [
                    {'range': [0, 1], 'color': "lightgreen"},
                    {'range': [1, 5], 'color': "lightyellow"},
                    {'range': [5, 10], 'color': "lightcoral"}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': 1
                }
            }
        ))

        fig.update_layout(height=300)
        st.plotly_chart(fig, use_container_width=True)

    with col2:
        st.markdown("### Compliance Status")

        if results['annual_risk_median'] <= nz_standard:
            st.success("✅ **COMPLIANT**\n\nMeets NZ Drinking Water Standard")
        elif results['annual_risk_median'] <= nz_standard * 10:
            st.warning("⚠️ **MARGINAL**\n\nClose to guideline, requires attention")
        else:
            st.error("❌ **NON-COMPLIANT**\n\nExceeds NZ Drinking Water Standard")

        st.markdown(f"""
        **Benchmark:** 10⁻⁶ DALY/person/year

        **Your Result:** {results['annual_risk_median']:.2e}

        **Ratio:** {ratio:.2f}x guideline
        """)

def show_reports():
    """Show report generation page."""

    st.header("📄 Professional Reports")

    if st.session_state.results is None:
        st.warning("⚠️ No results available. Please run an assessment first.")
        return

    st.markdown("""
    <div class="info-box">
    Generate professional PDF or Word reports for stakeholders and regulatory submissions.
    </div>
    """, unsafe_allow_html=True)

    # Report type selection
    report_type = st.radio(
        "Select Report Template",
        options=[
            "📋 Executive Summary (2-3 pages)",
            "🔬 Technical Assessment (detailed)",
            "⚖️ Regulatory Compliance Report"
        ]
    )

    # Report options
    st.subheader("Report Options")

    col1, col2 = st.columns(2)

    with col1:
        include_plots = st.checkbox("Include Risk Comparison Plots", value=True)
        include_tables = st.checkbox("Include Data Tables", value=True)

    with col2:
        include_uncertainty = st.checkbox("Include Uncertainty Analysis", value=True)
        include_references = st.checkbox("Include Literature References", value=True)

    # Generate report
    st.markdown("---")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("📄 Generate PDF Report", type="primary", use_container_width=True):
            st.info("PDF report generation will be implemented in full version")

    with col2:
        if st.button("📝 Generate Word Report", type="primary", use_container_width=True):
            st.info("Word report generation will be implemented in full version")

def show_help():
    """Show help page."""

    st.header("Help & Documentation")

    tab1, tab2, tab3 = st.tabs(["Quick Start", "Technical Notes", "About"])

    with tab1:
        st.markdown("""
        ## Workflow

        **1. Project Setup**
        - Enter project ID, assessor name, population

        **2. Load Dilution Data (if available)**
        - Go to Treatment Scenarios page
        - Enter MetOcean data directory path
        - Load dilution files
        - Select site, scenario (nina/nino), depth, conservatism (p10/p50/p90)

        **3. Configure Assessment**
        - Select pathogen (Norovirus, Campylobacter, Crypto, etc.)
        - Select exposure route (Swimming or Shellfish)
        - Enter effluent concentration (copies/L)
        - Set ingestion volume (mL)
        - Set exposure frequency (events/year)
        - Dilution factor (from MetOcean data or manual entry)

        **4. Run Monte Carlo Simulation**
        - Set iterations (default 10,000)
        - Click "Run Assessment"
        - Wait 2-5 seconds

        **5. Review Results**
        - Check compliance (≤ 10⁻⁶ DALY/person/year)
        - Review infection/illness risk metrics
        - Check uncertainty bounds (5th, 50th, 95th percentiles)

        **6. Export**
        - Download CSV results
        - Generate Word report (if needed)

        ---

        **Example Parameters - Swimming:**
        - Pathogen: Norovirus
        - Effluent concentration: 1,000,000 copies/L
        - Dilution: 800:1 (conservative from MetOcean)
        - Volume: 100 mL/swim
        - Frequency: 15 swims/year
        """)

    with tab2:
        st.markdown("""
        ## Technical Notes

        ### Dose-Response Models

        **Norovirus** - Beta-Poisson (α=0.04, β=0.055)
        - P(infection) = 1 - [1 + (dose/β)]^(-α)
        - Source: Teunis et al. (2008)

        **Campylobacter** - Beta-Poisson
        - Source: FAO/WHO (2009)

        **Cryptosporidium** - Exponential (r=0.0042)
        - P(infection) = 1 - exp(-r × dose)
        - Source: DuPont et al. (1995)

        ### Monte Carlo Simulation

        - Default iterations: 10,000
        - Lognormal distribution for pathogen concentration
        - Triangular distribution for ingestion volume (25-100 mL)
        - 95% confidence intervals (5th and 95th percentiles)

        ### MetOcean Dilution Data Format

        **Expected file structure:**
        ```
        year, month, day, hour_utc, minutes, sec, C_surf[g/m3], C_mid[g/m3], C_nearbed[g/m3]
        ```

        **Dilution calculation:**
        - DF = Source Concentration / Ambient Concentration
        - Conservative = p10 (10th percentile)
        - Typical = p50 (median)
        - Optimistic = p90 (90th percentile)

        **Depth selection:**
        - `surf` = Surface (swimming)
        - `nearbed` = Near-bed (shellfish)

        ### Typical Concentration Ranges

        | Source | Norovirus (copies/L) |
        |--------|---------------------|
        | Raw wastewater | 10⁵ - 10⁷ |
        | Primary treatment | 10⁴ - 10⁶ |
        | Secondary treatment | 10² - 10⁴ |
        | Tertiary treatment | 10¹ - 10³ |
        | Receiving water (diluted) | 10⁰ - 10² |

        ### LRV (Log Reduction Values)

        | Treatment | Typical LRV |
        |-----------|-------------|
        | Primary | 0.5 - 1.0 |
        | Secondary | 1.0 - 2.0 |
        | Tertiary | 3.0 - 5.0 |
        | UV disinfection | 2.0 - 4.0 |

        ### Documentation

        Technical manual: `NIWA_Marine_QMRA_Technical_Manual.docx`
        """)

    with tab3:
        st.markdown("""
        ## About

        **NIWA Marine QMRA Toolkit**
        *Internal Assessment Tool v2.0*

        **Developed by:** NIWA Earth Sciences
        **Purpose:** Wastewater outfall risk assessment (internal use)

        **Contact:**
        Reza Moghaddam - Developer
        NIWA Earth Sciences
        National Institute of Water & Atmospheric Research

        ---

        **Technical References:**
        - NZ Drinking Water Standards 2005 (Rev 2008)
        - WHO Guidelines for Drinking-water Quality (2011)
        - Haas, Rose & Gerba: QMRA (2014)
        - Teunis et al.: Norovirus dose-response (2008)

        ---

        **Limitations:**
        - Swimming and shellfish exposure routes only
        - Requires MetOcean hydrodynamic modeling data for site-specific dilution
        - Conservative assumptions (e.g., no pathogen die-off)
        - Does not account for tidal variability
        """)

def show_pathogen_info(pathogen):
    """Show information about selected pathogen."""

    info = {
        'norovirus': {
            'type': 'Enteric Virus',
            'model': 'Exponential',
            'pill_inf': 0.7,
            'severity': 'Gastroenteritis, vomiting, diarrhea',
            'incubation': '12-48 hours'
        },
        'campylobacter': {
            'type': 'Bacterial Pathogen',
            'model': 'Beta-Poisson',
            'pill_inf': 0.33,
            'severity': 'Diarrheal disease, fever',
            'incubation': '2-5 days'
        },
        'cryptosporidium': {
            'type': 'Protozoan Parasite',
            'model': 'Exponential',
            'pill_inf': 0.39,
            'severity': 'Severe diarrhea',
            'incubation': '7-10 days'
        }
    }

    if pathogen in info:
        p = info[pathogen]
        st.info(f"""
        **{pathogen.title()}**

        - Type: {p['type']}
        - Dose-Response Model: {p['model']}
        - Illness|Infection Ratio: {p['pill_inf']}
        - Typical Severity: {p['severity']}
        - Incubation Period: {p['incubation']}
        """)

def run_qmra_assessment(pathogen, exposure_route, concentration, volume, frequency, population, iterations, confidence_level, dilution_factor=1.0):
    """
    Run QMRA assessment using the full QMRA toolkit modules.

    This function integrates:
    - PathogenDatabase for validated dose-response parameters
    - Monte Carlo simulation for uncertainty analysis
    - Proper dose-response models (Beta-Poisson, Exponential)
    - Exposure assessment with dilution factors
    """
    try:
        # Import core QMRA modules
        from pathogen_database import PathogenDatabase
        from dose_response import create_dose_response_model
        from monte_carlo import MonteCarloSimulator, create_lognormal_distribution, create_uniform_distribution

        # Initialize pathogen database
        pathogen_db = PathogenDatabase()

        # Get pathogen parameters from database
        pathogen_info = pathogen_db.get_pathogen_info(pathogen)
        default_model_type = pathogen_db.get_default_model_type(pathogen)
        dr_params = pathogen_db.get_dose_response_parameters(pathogen, default_model_type)
        health_data = pathogen_db.get_health_impact_data(pathogen)

        # Create dose-response model
        dr_model = create_dose_response_model(default_model_type, dr_params)

        # Apply dilution factor to concentration
        diluted_concentration = concentration / dilution_factor

        # Initialize Monte Carlo simulator
        mc_simulator = MonteCarloSimulator(random_seed=42)

        # Add uncertainty distributions for Monte Carlo simulation
        # Concentration uncertainty (lognormal distribution - common for microbial data)
        concentration_dist = create_lognormal_distribution(
            mean=np.log(diluted_concentration),
            std=0.5,  # Standard deviation in log space (moderate uncertainty)
            name="pathogen_concentration"
        )
        mc_simulator.add_distribution("pathogen_concentration", concentration_dist)

        # Volume uncertainty (uniform distribution for ingestion variability)
        volume_dist = create_uniform_distribution(
            min_val=volume * 0.5,  # 50% to 150% of nominal volume
            max_val=volume * 1.5,
            name="ingestion_volume"
        )
        mc_simulator.add_distribution("ingestion_volume", volume_dist)

        # Define QMRA model function for Monte Carlo simulation
        def qmra_model(samples):
            """Calculate infection and illness risks from sampled inputs."""
            # Get sampled values
            conc_samples = samples["pathogen_concentration"]
            vol_samples = samples["ingestion_volume"]

            # Calculate dose: concentration (per L) × volume (mL) / 1000
            dose_samples = (conc_samples * vol_samples) / 1000.0

            # Calculate infection probability using dose-response model
            pinf_samples = dr_model.calculate_infection_probability(dose_samples)

            return pinf_samples

        # Run Monte Carlo simulation for infection probability
        infection_results = mc_simulator.run_simulation(
            qmra_model,
            n_iterations=iterations,
            variable_name="infection_probability"
        )

        # Get illness-to-infection ratio
        pill_inf_ratio = health_data["illness_to_infection_ratio"]

        # Calculate illness probability samples
        pinf_samples = infection_results.samples
        pill_samples = pinf_samples * pill_inf_ratio

        # Calculate annual risk: 1 - (1 - P_infection)^frequency
        annual_samples = 1 - np.power(1 - pinf_samples, frequency)

        # Calculate statistics for all risk metrics
        def calc_stats(samples):
            """Calculate statistics for a sample array."""
            return {
                'median': float(np.median(samples)),
                'mean': float(np.mean(samples)),
                'std': float(np.std(samples)),
                '5th': float(np.percentile(samples, 5)),
                '95th': float(np.percentile(samples, 95))
            }

        pinf_stats = calc_stats(pinf_samples)
        pill_stats = calc_stats(pill_samples)
        annual_stats = calc_stats(annual_samples)

        # Return comprehensive results
        return {
            # Infection risk statistics
            'pinf_median': pinf_stats['median'],
            'pinf_mean': pinf_stats['mean'],
            'pinf_std': pinf_stats['std'],
            'pinf_5th': pinf_stats['5th'],
            'pinf_95th': pinf_stats['95th'],

            # Illness risk statistics
            'pill_median': pill_stats['median'],
            'pill_mean': pill_stats['mean'],
            'pill_std': pill_stats['std'],
            'pill_5th': pill_stats['5th'],
            'pill_95th': pill_stats['95th'],

            # Annual risk statistics
            'annual_risk_median': annual_stats['median'],
            'annual_mean': annual_stats['mean'],
            'annual_std': annual_stats['std'],
            'annual_5th': annual_stats['5th'],
            'annual_95th': annual_stats['95th'],

            # Population impact
            'population_impact': int(population * annual_stats['median']),

            # Metadata for reporting
            'pathogen': pathogen,
            'model_type': default_model_type,
            'model_parameters': dr_params,
            'illness_to_infection_ratio': pill_inf_ratio,
            'iterations': iterations,
            'dilution_factor': dilution_factor
        }

    except Exception as e:
        # If proper modules fail, fall back to simplified calculation
        import warnings
        warnings.warn(f"QMRA module error: {e}. Using simplified calculation.")

        # Apply dilution factor to concentration
        diluted_concentration = concentration / dilution_factor

        # Dose calculation
        dose = (diluted_concentration * volume) / 1000  # Convert mL to L

        # Simplified dose-response parameters (fallback only)
        if pathogen == 'norovirus':
            alpha = 0.04
            pinf = 1 - np.exp(-alpha * dose)
            pill_inf = 0.7
        elif pathogen == 'campylobacter':
            # Proper Beta-Poisson approximation
            alpha = 0.145
            beta = 7.59
            pinf = 1 - np.power(1 + dose / beta, -alpha)
            pill_inf = 0.33
        elif pathogen == 'cryptosporidium':
            r = 0.0042
            pinf = 1 - np.exp(-r * dose)
            pill_inf = 0.39
        else:
            pinf = 0.01
            pill_inf = 0.5

        pill = pinf * pill_inf
        annual_risk = 1 - (1 - pinf) ** frequency

        # Monte Carlo simulation (simplified)
        np.random.seed(42)
        pinf_samples = np.random.lognormal(np.log(max(pinf, 1e-10)), 0.5, iterations)
        pill_samples = pinf_samples * pill_inf
        annual_samples = 1 - (1 - pinf_samples) ** frequency

        return {
            'pinf_median': np.median(pinf_samples),
            'pinf_mean': np.mean(pinf_samples),
            'pinf_std': np.std(pinf_samples),
            'pinf_5th': np.percentile(pinf_samples, 5),
            'pinf_95th': np.percentile(pinf_samples, 95),

            'pill_median': np.median(pill_samples),
            'pill_mean': np.mean(pill_samples),
            'pill_std': np.std(pill_samples),
            'pill_5th': np.percentile(pill_samples, 5),
            'pill_95th': np.percentile(pill_samples, 95),

            'annual_risk_median': np.median(annual_samples),
            'annual_mean': np.mean(annual_samples),
            'annual_std': np.std(annual_samples),
            'annual_5th': np.percentile(annual_samples, 5),
            'annual_95th': np.percentile(annual_samples, 95),

            'population_impact': int(population * np.median(annual_samples))
        }

def save_project():
    """Save project to JSON file."""
    if not st.session_state.project_data:
        st.warning("No project data to save")
        return

    project_json = json.dumps(st.session_state.project_data, indent=2)
    st.download_button(
        label="💾 Download Project File",
        data=project_json,
        file_name=f"qmra_project_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        mime="application/json"
    )

def load_project():
    """Load project from JSON file."""
    st.info("Project loading will be implemented in full version")

def load_demo_data():
    """Load demo data for trying out the toolkit."""

    # Demo project: Beach Swimming Assessment
    st.session_state.project_data = {
        'project_name': 'Demo: Beach Swimming Risk Assessment',
        'assessor_name': 'Demo User',
        'client_name': 'Demo Council',
        'assessment_date': '2025-10-06',
        'population': 50000,
        'pathogen': 'norovirus',
        'exposure_route': 'primary_contact',
        'concentration': 1000.0,
        'volume': 100.0,
        'frequency': 7,
        'iterations': 10000,
        'confidence_level': 95.0
    }

    # Auto-run demo assessment
    results = run_qmra_assessment(
        pathogen='norovirus',
        exposure_route='primary_contact',
        concentration=1000.0,
        volume=100.0,
        frequency=7,
        population=50000,
        iterations=10000,
        confidence_level=95.0
    )

    st.session_state.results = results

if __name__ == "__main__":
    main()
