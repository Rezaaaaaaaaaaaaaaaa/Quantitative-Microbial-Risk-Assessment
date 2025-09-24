#!/usr/bin/env python3
"""
Generate Updated QMRA SIP Application Document
Incorporating RW (Richard Woods) feedback and comments from PDF review
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_updated_sip_document():
    """Create an updated SIP application document addressing RW comments from PDF review."""

    # Create new document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "QMRA Workflow Engine - Updated SIP Application (RW Comments Addressed)"
    doc.core_properties.author = "Reza Moghaddam, David Wood - Earth Sciences New Zealand"
    doc.core_properties.subject = "Strategic Investment Proposal - QMRA Workflow Engine"

    # Title
    title = doc.add_heading('Structured internal project application 2025-2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Main SIP Table - matching original format exactly
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Set column widths to match original
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)

    # Helper function to add table rows
    def add_table_row(label, content, make_bold=False):
        row = table.add_row()
        label_cell = row.cells[0]
        content_cell = row.cells[1]

        label_para = label_cell.paragraphs[0]
        label_para.text = label
        if make_bold:
            label_para.runs[0].bold = True

        content_cell.text = content

    # Project Overview section
    add_table_row("Project Overview", "", True)
    add_table_row("Project name: (Short title)", "Development of QMRA Assessment Toolkit")
    add_table_row("Staff: (who will be completing the work?)", "Reza Moghaddam (Lead Developer - 200 hrs), David Wood (Model Review & Support - 40 hrs)")
    add_table_row("Project Manager: (usually a Group Manager)", "Andrew Hughes")
    add_table_row("Region:", "Hamilton")
    add_table_row("Centre:", "Freshwater")
    add_table_row("Type: (science, operations activity, or other - explain)", "Science (Applied Research & Development)")
    add_table_row("Budget: (attach costing prepared by your project coordinator)", "")

    # Project objective - RW47: Define acronym
    add_table_row("Project objective: (30 words max)",
                  "Develop a Python-based Quantitative Microbial Risk Assessment (QMRA) assessment toolkit to standardise processes, improve reproducibility and auditability, and reduce manual work for regulatory compliance assessments.")

    # Project outline - Updated with RW feedback addressing RW48, RW49, RW50, RW51
    project_outline = (
        "Earth Sciences New Zealand currently undertakes Quantitative Microbial Risk Assessment (QMRA) projects "
        "using @Risk Excel add-in, which is commercial licensed software that has proven problematic and costly. "
        "@Risk relies on error-prone Excel-based workflows with extensive manual processes that could be automated, "
        "leading to inconsistencies and quality control issues. Recent projects have lost up to 80 hours due to "
        "security system conflicts within NIWA's firewall environment, requiring client extensions and budget overruns. "
        "Based on our project experience, typical QMRA projects involve 40-60 hours of manual work including "
        "dose-response model setup, exposure assessment, dilution modelling integration, simulation configuration, "
        "and report generation using @Risk's cumbersome Excel interface.\n\n"

        "This project will develop a Python-based QMRA assessment toolkit following a Minimum Viable Product (MVP) "
        "approach, focusing specifically on norovirus exposure scenarios for primary contact and shellfish consumption. "
        "The toolkit will replace @Risk dependency, eliminate manual Excel-based processes through automation, "
        "incorporate Earth Sciences New Zealand's dilution modelling capabilities (our key differentiator), and work "
        "with engineer-provided log reduction values rather than attempting complex treatment calculations. This focused "
        "approach avoids the 'do everything' trap while delivering immediate value through automated, reproducible workflows.\n\n"

        "The Python implementation will develop technical capabilities in QMRA methodology while maintaining our "
        "competitive position in the QMRA market. The toolkit will integrate dilution modelling inputs, automate "
        "routine calculations that are currently manual and error-prone in Excel, and generate standardised outputs. "
        "While Charlotte Jones-Todd has developed an R package, this Python approach provides greater integration "
        "with Earth Sciences New Zealand systems and builds internal technical depth essential for our $40-70K QMRA "
        "projects that generate additional consenting work opportunities. The project includes collaboration with "
        "NIWA's communications team to ensure effective promotion and market awareness of the new toolkit capabilities."
    )
    add_table_row("Project outline: (150-300 words max)", project_outline)

    # Project outputs
    project_outputs = (
        "QMRA Assessment Toolkit (Python MVP replacing @Risk dependency)\n"
        "Norovirus exposure models for primary contact and shellfish consumption\n"
        "Dilution modelling integration module (NIWA's key differentiator)\n"
        "Validated dose-response database with engineer-provided LRV inputs\n"
        "Standardised reporting templates for regulatory compliance"
    )
    add_table_row("Project outputs: (e.g., a journal paper or an App, or a safe operating procedure or guidance document for operations activities)", project_outputs)

    # Project impact
    add_table_row("Project impact: (choose an SCI impact area that the project aligns with, see graphic below)", "Protecting our diversity\nImproved environmental health")

    # Alignment
    alignment_text = (
        "This project aligns with Earth Sciences New Zealand's analytical capabilities development "
        "and supports regulatory compliance services. It develops our technical capacity for water quality "
        "risk assessment and supports our role in environmental protection. The improved reproducibility "
        "and auditability will strengthen our credibility with regulatory bodies."
    )
    add_table_row("Alignment: (with a programme and/or National Centre outcomes or KPIs)", alignment_text)

    # Outcomes for Māori
    maori_outcomes = (
        "Supporting improved water quality assessment capabilities that contribute to protecting "
        "water bodies important for cultural values and mahinga kai. The developed QMRA capabilities "
        "will support decision-making that considers cultural significance of water resources and "
        "traditional food gathering practices."
    )
    add_table_row("Outcomes for Māori: (may include partnerships, resourcing, alignment with aspirations)", maori_outcomes)

    # Operations alignment
    add_table_row("Operations alignment: (for non-science projects, how does this work contribute to inputs or enablers from the graphic below)", "Not applicable")

    doc.add_paragraph()

    # Add the actual QMRA schematic diagram
    doc.add_paragraph()

    # Try to add the generated schematic image
    try:
        # First, ensure the schematic is generated
        import subprocess
        subprocess.run(['python', 'create_qmra_schematic.py'],
                      cwd='C:\\Users\\moghaddamr\\OneDrive - NIWA\\Quantitative Microbial Risk Assessment\\src',
                      capture_output=True)

        # Add the image
        doc.add_picture('C:\\Users\\moghaddamr\\OneDrive - NIWA\\Quantitative Microbial Risk Assessment\\src\\qmra_assessment_toolkit_schematic.png',
                       width=Inches(7))

        # Add figure caption
        caption_para = doc.add_paragraph()
        caption_para.add_run('Figure 1: ').bold = True
        caption_para.add_run('QMRA Assessment Toolkit - System Architecture & Workflow')
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    except Exception as e:
        # Fallback: Add a text description if image generation fails
        caption_para = doc.add_paragraph()
        caption_para.add_run('Figure 1: ').bold = True
        caption_para.add_run('QMRA Assessment Toolkit Architecture')
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add detailed textual description
        fallback_desc = doc.add_paragraph(
            'The QMRA Assessment Toolkit follows a structured workflow: Input data (water quality measurements, '
            'pathogen concentrations, log reduction values, population demographics, exposure scenarios, and '
            'treatment parameters) flows into a Python processing core containing four key modules: Norovirus '
            'Database, Monte Carlo Simulation, Dose-Response Models, and Exposure Assessment. The processing '
            'pipeline includes data validation, risk calculation, uncertainty quantification, and report generation '
            'stages. The system produces comprehensive outputs including risk assessment results, compliance reports, '
            'health risk estimates, regulatory documentation, uncertainty analysis, and decision support tools. '
            'The modular design enables easy updates and integration with existing NIWA systems.'
        )
        print(f"Note: Image generation failed ({e}), using text description instead")

    doc.add_paragraph()

    # Add detailed description of the assessment toolkit
    engine_desc = doc.add_paragraph(
        'The QMRA Assessment Toolkit processes multiple data inputs through a modular Python-based system. '
        'Input data includes water quality measurements, pathogen concentrations, user-defined log reduction values, '
        'population demographics, and exposure scenarios. The core processing engine integrates a norovirus pathogen '
        'database as the initial proof-of-concept, applies Monte Carlo simulation for uncertainty analysis, and '
        'implements validated dose-response models. The system generates comprehensive outputs including risk assessment '
        'results, regulatory compliance reports, health risk estimates, and decision support documentation for '
        'environmental health protection, while integrating Earth Sciences New Zealand\'s dilution modelling capabilities.'
    )

    doc.add_paragraph()

    # Work Programme and Timeline - in original format with RW52 addressed
    doc.add_heading('WORK PROGRAMME AND TIMELINE', 1)
    doc.add_paragraph('Outline the tasks to be done, who will do what and by when. Be as specific as possible.')

    # Work programme table
    work_table = doc.add_table(rows=1, cols=3)
    work_table.style = 'Table Grid'
    header_cells = work_table.rows[0].cells
    header_cells[0].text = 'Task'
    header_cells[1].text = 'Specific activity (who, what)'
    header_cells[2].text = 'Hours'

    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # RW52: More explicit detail on when reviews will take place - after each step
    work_items = [
        ('QMRA Literature Review & Charlotte\'s R Package Assessment',
         'Review Charlotte Jones-Todd\'s R package, assess existing QMRA tools, analyse current models and methodologies to avoid reinventing wheel (Reza)',
         '25'),
        ('Review Meeting 1',
         'David to review literature findings and provide feedback on methodology approach (David & Reza)',
         '2'),
        ('MVP Requirements & Design',
         'Define minimum viable product scope (norovirus primary contact + shellfish), system architecture, @Risk replacement strategy (Reza)',
         '35'),
        ('Review Meeting 2',
         'David to review system design and architecture decisions (David & Reza)',
         '3'),
        ('Core Development (Norovirus MVP)',
         'Develop norovirus exposure models for primary contact and shellfish consumption, Python framework replacing @Risk (Reza)',
         '45'),
        ('Review Meeting 3',
         'David to review core implementation and model accuracy (David & Reza)',
         '5'),
        ('Dilution Modelling Integration',
         'Implement Earth Sciences New Zealand\'s dilution modelling capabilities integration, our key differentiator in NZ QMRA market (Reza)',
         '25'),
        ('Review Meeting 4',
         'David to review dilution modelling integration (David & Reza)',
         '3'),
        ('Monte Carlo & Uncertainty Analysis',
         'Implement Monte Carlo simulation replacing @Risk functionality, uncertainty quantification for decision support (Reza)',
         '20'),
        ('Testing & Validation',
         'Validate against known benchmarks, test @Risk replacement functionality, ensure regulatory compliance outputs (Reza)',
         '20'),
        ('Review Meeting 5',
         'David to review testing results and validation approach (David & Reza)',
         '2'),
        ('QMRA Model Review & Validation',
         'Technical review of implemented models, validation of dose-response relationships and dilution integration (David)',
         '25'),
        ('Documentation & Training',
         'Technical documentation, user guides, training materials for @Risk transition, regulatory templates (David)',
         '15'),
        ('Deployment & @Risk Transition',
         'System deployment, staff training on new toolkit, transition away from @Risk dependency (Reza/David)',
         '10'),
        ('Marketing & Communications Strategy',
         'Collaborate with NIWA communications team to develop product promotion strategy, create website content, develop client outreach materials to ensure market awareness of new QMRA toolkit capabilities (Reza)',
         '30'),
        ('Final Review Meeting',
         'Project completion review and handover (David, Reza & Andrew)',
         '3')
    ]

    for task, activity, hours in work_items:
        row_cells = work_table.add_row().cells
        row_cells[0].text = task
        row_cells[1].text = activity
        row_cells[2].text = hours

    # Add table caption - updated total hours
    table_caption = doc.add_paragraph()
    table_caption.add_run('Table 1: ').bold = True
    table_caption.add_run('Work Programme and Timeline for QMRA Assessment Toolkit Development with Regular Review Meetings (Total: 258 hours)')
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Project Management and Oversight section (addressing Richard's broader comment)
    doc.add_heading('PROJECT MANAGEMENT AND OVERSIGHT', 1)
    project_mgmt_text = (
        'To ensure project success, a structured collaboration approach has been established between Reza Moghaddam '
        'and David Wood with regular oversight from Andrew Hughes as Project Manager.\n\n'

        'Weekly Review Meetings: David and Reza will conduct weekly review meetings after each major development phase '
        'as outlined in the work programme. These meetings will cover technical progress, methodology validation, '
        'and any challenges encountered. David will provide technical guidance on QMRA methodologies and model validation, '
        'while Reza will lead the development implementation.\n\n'

        'Problem Resolution: Any technical issues or roadblocks will be immediately flagged during weekly reviews. '
        'David will provide expert guidance on QMRA-specific challenges, while Andrew Hughes will be consulted '
        'for any project scope or resource issues. This ensures rapid problem identification and resolution.\n\n'

        'Quality Assurance: David will conduct formal technical reviews at key milestones (literature review, '
        'system design, core implementation, dilution integration, and final validation). This staged approach '
        'ensures quality control throughout development rather than only at project completion.\n\n'

        'Documentation and Knowledge Transfer: All review meetings will be documented with action items and '
        'decisions recorded. David will contribute to technical documentation to ensure knowledge transfer '
        'and future maintainability of the system.'
    )
    doc.add_paragraph(project_mgmt_text)

    doc.add_paragraph()

    # Collaboration section
    doc.add_heading('EMERGING COLLABORATION OPPORTUNITIES', 1)
    collab_text = (
        'Earth Sciences New Zealand has completed seven QMRA projects over the past three years, with individual project values '
        'of $40-70K. The real business value extends beyond the QMRA work itself, as these projects provide '
        'entry into the consenting process and generate additional consulting opportunities, as demonstrated '
        'by the Beachlands QMRA follow-on work.\n\n'

        'The New Zealand Institute for Public Health and Forensic Science (PHF) has approached Earth Sciences New Zealand to develop '
        'QMRA guidance specifically for shellfish safety assessment, providing immediate application for the '
        'norovirus shellfish consumption models planned in this toolkit. This collaboration demonstrates '
        'market demand and offers real-world validation opportunities.\n\n'

        'The @Risk replacement toolkit addresses multiple critical operational issues. Recent projects have lost up to '
        '80 hours due to @Risk security conflicts within NIWA\'s firewall environment, requiring client extensions and budget overruns. '
        'Additionally, @Risk\'s Excel-based workflows are inherently error-prone with extensive manual processes that introduce '
        'inconsistencies and quality control challenges. Moving to our own Python-based solution eliminates this dependency, '
        'automates manual processes, improves reproducibility and auditability, while building internal technical capability '
        'essential for maintaining competitive advantage in the changing regulatory environment.'
    )
    doc.add_paragraph(collab_text)

    doc.add_paragraph()

    # Chief Scientist Support section - matching original format
    doc.add_heading('CHIEF SCIENTIST SUPPORT', 1)

    chief_scientist_para = doc.add_paragraph()
    chief_scientist_para.add_run('Chief Scientist comment: ').bold = True
    chief_scientist_para.add_run('(For example - If agreement that project required, indicate why SIP mechanism versus Centre Funds; What is/are the key output(s) and how will NIWA/National Centre/programme/individual benefit from that; note that there must be an output at the end of the project)')

    doc.add_paragraph()
    doc.add_paragraph()

    signature_para = doc.add_paragraph()
    signature_para.add_run('Signature: ').bold = True

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f'Updated SIP document addressing RW comments - Generated on {datetime.now().strftime("%d %B %Y")}').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    filename = f'QMRA_SIP_Final_RW_Addressed_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    doc.save(filename)
    print(f'Updated SIP document (RW comments addressed) created: {filename}')
    return filename

if __name__ == '__main__':
    create_updated_sip_document()