#!/usr/bin/env python3
"""
Extract content and style from SIP_example.docx to understand NIWA's preferred tone and format
"""

import docx
from docx import Document
import re
import os

def extract_document_content(docx_path):
    """Extract all text content from the Word document."""
    try:
        doc = Document(docx_path)

        content = {
            'paragraphs': [],
            'headings': [],
            'tables': [],
            'style_analysis': {
                'paragraph_count': 0,
                'heading_levels': {},
                'avg_paragraph_length': 0,
                'sentence_patterns': [],
                'tone_indicators': []
            }
        }

        total_chars = 0

        # Extract paragraphs and analyze style
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content['paragraphs'].append({
                    'text': text,
                    'style': para.style.name if para.style else 'Normal',
                    'level': getattr(para, 'level', 0) if hasattr(para, 'level') else 0
                })

                # Check if it's a heading
                if any(keyword in para.style.name.lower() for keyword in ['heading', 'title']):
                    content['headings'].append({
                        'text': text,
                        'level': para.style.name,
                        'style': para.style.name
                    })

                    # Count heading levels
                    level = para.style.name
                    content['style_analysis']['heading_levels'][level] = \
                        content['style_analysis']['heading_levels'].get(level, 0) + 1

                content['style_analysis']['paragraph_count'] += 1
                total_chars += len(text)

                # Analyze sentence patterns
                sentences = re.split(r'[.!?]+', text)
                for sentence in sentences:
                    sentence = sentence.strip()
                    if sentence:
                        content['style_analysis']['sentence_patterns'].append({
                            'length': len(sentence.split()),
                            'starts_with': sentence.split()[0] if sentence.split() else '',
                            'contains_numbers': bool(re.search(r'\d', sentence)),
                            'contains_technical': bool(re.search(r'(research|development|analysis|assessment|methodology|framework|approach|system|process)', sentence.lower()))
                        })

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            content['tables'].append(table_data)

        # Calculate average paragraph length
        if content['style_analysis']['paragraph_count'] > 0:
            content['style_analysis']['avg_paragraph_length'] = \
                total_chars / content['style_analysis']['paragraph_count']

        return content

    except Exception as e:
        print(f"Error reading document: {e}")
        return None

def analyze_writing_style(content):
    """Analyze the writing style and tone from extracted content."""
    if not content:
        return None

    analysis = {
        'tone_characteristics': [],
        'sentence_structure': {},
        'vocabulary_patterns': {},
        'formatting_style': {},
        'typical_phrases': []
    }

    # Analyze sentence patterns
    sentence_lengths = [s['length'] for s in content['style_analysis']['sentence_patterns']]
    if sentence_lengths:
        analysis['sentence_structure'] = {
            'avg_length': sum(sentence_lengths) / len(sentence_lengths),
            'range': (min(sentence_lengths), max(sentence_lengths)),
            'technical_ratio': sum(1 for s in content['style_analysis']['sentence_patterns'] if s['contains_technical']) / len(sentence_lengths)
        }

    # Analyze vocabulary and phrases
    all_text = ' '.join([p['text'] for p in content['paragraphs']])

    # Common NIWA/scientific phrases
    technical_phrases = [
        'research and development', 'scientific understanding', 'knowledge gaps',
        'research outcomes', 'strategic investment', 'capability development',
        'collaborative research', 'industry engagement', 'research excellence',
        'innovation', 'methodology', 'framework', 'approach', 'assessment'
    ]

    found_phrases = []
    for phrase in technical_phrases:
        if phrase.lower() in all_text.lower():
            found_phrases.append(phrase)

    analysis['typical_phrases'] = found_phrases

    # Tone indicators
    formal_indicators = ['will', 'shall', 'proposed', 'objectives', 'outcomes', 'deliverables']
    collaborative_indicators = ['partnership', 'collaboration', 'engagement', 'stakeholder']
    technical_indicators = ['methodology', 'framework', 'analysis', 'assessment', 'research']

    analysis['tone_characteristics'] = {
        'formal': sum(1 for indicator in formal_indicators if indicator.lower() in all_text.lower()),
        'collaborative': sum(1 for indicator in collaborative_indicators if indicator.lower() in all_text.lower()),
        'technical': sum(1 for indicator in technical_indicators if indicator.lower() in all_text.lower())
    }

    return analysis

def generate_style_report(content, analysis):
    """Generate a comprehensive style analysis report."""
    if not content or not analysis:
        return "Unable to analyze document style."

    report = f"""
NIWA SIP PROPOSAL STYLE ANALYSIS
=================================

DOCUMENT STRUCTURE:
- Total paragraphs: {content['style_analysis']['paragraph_count']}
- Heading levels found: {list(content['style_analysis']['heading_levels'].keys())}
- Average paragraph length: {content['style_analysis']['avg_paragraph_length']:.1f} characters
- Number of tables: {len(content['tables'])}

WRITING STYLE CHARACTERISTICS:
- Average sentence length: {analysis['sentence_structure'].get('avg_length', 0):.1f} words
- Technical content ratio: {analysis['sentence_structure'].get('technical_ratio', 0)*100:.1f}%
- Sentence length range: {analysis['sentence_structure'].get('range', (0,0))}

TONE ANALYSIS:
- Formal indicators: {analysis['tone_characteristics']['formal']} occurrences
- Collaborative indicators: {analysis['tone_characteristics']['collaborative']} occurrences
- Technical indicators: {analysis['tone_characteristics']['technical']} occurrences

TYPICAL PHRASES FOUND:
{chr(10).join([f"- {phrase}" for phrase in analysis['typical_phrases']])}

SAMPLE HEADINGS:
{chr(10).join([f"- {h['text']}" for h in content['headings'][:5]])}

SAMPLE CONTENT EXCERPTS:
{chr(10).join([f"- {p['text'][:100]}..." for p in content['paragraphs'][:3] if len(p['text']) > 50])}
"""

    return report

def main():
    """Main function to extract and analyze SIP example style."""
    docx_path = "SIP_example.docx"

    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found in current directory")
        return

    print("Extracting content from SIP_example.docx...")
    content = extract_document_content(docx_path)

    if content:
        print("Analyzing writing style and tone...")
        analysis = analyze_writing_style(content)

        print("Generating style analysis report...")
        report = generate_style_report(content, analysis)

        # Save analysis to file
        with open('sip_style_analysis.txt', 'w', encoding='utf-8') as f:
            f.write(report)

        print(report)
        print(f"\nStyle analysis saved to: sip_style_analysis.txt")

        # Save raw content for reference
        with open('sip_content_extract.txt', 'w', encoding='utf-8') as f:
            f.write("EXTRACTED CONTENT FROM SIP_EXAMPLE.DOCX\n")
            f.write("="*50 + "\n\n")

            f.write("HEADINGS:\n")
            for heading in content['headings']:
                f.write(f"{heading['level']}: {heading['text']}\n")

            f.write("\n\nFULL CONTENT:\n")
            for i, para in enumerate(content['paragraphs'], 1):
                f.write(f"\n[{i}] {para['style']}: {para['text']}\n")

        print("Raw content saved to: sip_content_extract.txt")

    else:
        print("Failed to extract content from the document.")

if __name__ == "__main__":
    main()