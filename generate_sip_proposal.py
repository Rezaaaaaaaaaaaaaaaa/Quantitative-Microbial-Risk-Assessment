#!/usr/bin/env python3
"""
SIP Proposal Generator for QMRA Workflow Engine
Creates a comprehensive SIP funding proposal document
"""

import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import re
import os

class SIPProposalGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document_styles()

    def setup_document_styles(self):
        """Setup document styles for professional formatting."""
        # Title style
        title_style = self.doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(16)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        # Heading styles
        heading1_style = self.doc.styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_font = heading1_style.font
        heading1_font.name = 'Arial'
        heading1_font.size = Pt(14)
        heading1_font.bold = True
        heading1_style.paragraph_format.space_before = Pt(12)
        heading1_style.paragraph_format.space_after = Pt(6)

        heading2_style = self.doc.styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_font = heading2_style.font
        heading2_font.name = 'Arial'
        heading2_font.size = Pt(12)
        heading2_font.bold = True
        heading2_style.paragraph_format.space_before = Pt(6)
        heading2_style.paragraph_format.space_after = Pt(3)

        # Normal text style
        normal_style = self.doc.styles.add_style('Custom Normal', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15

    def add_title_page(self):
        """Add the title page with project information."""
        # NIWA Logo placeholder and title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("NIWA Strategic Investment Programme (SIP)\nProject Proposal")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title.space_after = Pt(24)

        # Project title
        project_title = self.doc.add_paragraph()
        project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_title_run = project_title.add_run("Development of a Comprehensive Quantitative Microbial Risk Assessment (QMRA) Workflow Engine")
        project_title_run.font.size = Pt(16)
        project_title_run.font.bold = True
        project_title.space_after = Pt(24)

        # Project details table
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'

        details = [
            ("Principal Investigator:", "Reza Moghaddam"),
            ("Co-Investigator:", "David Wood"),
            ("Supporting Team:", "Andrew Hughes, Michael Bruce"),
            ("Programme:", "Strategic Investment Programme"),
            ("Requested Duration:", "6 months"),
            ("Requested Budget:", "$25,000 (150 hours + David Wood support)"),
            ("Start Date:", "February 2025"),
            ("Submission Date:", datetime.now().strftime("%B %d, %Y"))
        ]

        for i, (key, value) in enumerate(details):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
            # Bold the keys
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

        self.doc.add_page_break()

    def add_executive_summary(self):
        """Add executive summary section."""
        self.doc.add_heading('Executive Summary', level=1).style = 'Custom Heading 1'

        summary_text = """
        NIWA proposes to develop a comprehensive Quantitative Microbial Risk Assessment (QMRA) workflow engine to capitalize on New Zealand's evolving regulatory landscape and strengthen our position as the leading QMRA consultancy provider. This Python-based tool will integrate our existing research expertise into a standardized, efficient workflow system that addresses critical market timing and competitive challenges.

        The regulatory environment is rapidly changing with national wastewater environmental performance standards taking effect in August 2025, requiring standardized risk assessment methods. Sixty percent of public wastewater treatment plants will need consent renewals in the next decade, and Taumata Arowai's expanded role demands more rigorous, standardized risk assessments. This positions NIWA's QMRA expertise at the center of industry needs.

        Currently, NIWA rebuilds models from scratch for each consulting project, which has resulted in lost bids to competitors who can deliver faster turnaround times. The proposed workflow engine will reduce QMRA project delivery time by 60-70% while ensuring consistent methodology across all team members.

        Key deliverables include: standardized pathogen dose-response libraries with uncertainty distributions, treatment efficacy modules for common water treatment processes, automated Monte Carlo simulation engine, and customizable reporting templates that meet regulatory requirements. The tool will leverage existing R code components through Python interfaces while building a scalable, maintainable system integrated with NIWA's data infrastructure.

        This investment of $25,000 over 6 months will transform our consulting efficiency, strengthen our competitive position, and ensure NIWA remains the go-to organization for QMRA expertise in New Zealand's changing regulatory environment.
        """

        para = self.doc.add_paragraph(summary_text.strip())
        para.style = 'Custom Normal'

    def add_project_background(self):
        """Add project background and rationale."""
        self.doc.add_heading('1. Project Background and Rationale', level=1).style = 'Custom Heading 1'

        self.doc.add_heading('1.1 Regulatory Context and Market Timing', level=2).style = 'Custom Heading 2'

        background_text = """
        The New Zealand water sector is experiencing unprecedented regulatory transformation that creates both urgent need and significant opportunity for NIWA's QMRA capabilities:

        • National wastewater environmental performance standards become mandatory in August 2025
        • 60% of public wastewater treatment plants require consent renewals within the next decade
        • Taumata Arowai's expanded regulatory mandate requires more rigorous, standardized risk assessment methodologies
        • Regional councils are increasingly demanding evidence-based risk assessments for water management decisions

        This regulatory shift positions quantitative microbial risk assessment as an essential tool for compliance, making NIWA's expertise critical to the water industry's adaptation to new national standards.
        """

        para = self.doc.add_paragraph(background_text.strip())
        para.style = 'Custom Normal'

        self.doc.add_heading('1.2 Current Challenges and Competitive Position', level=2).style = 'Custom Heading 2'

        challenges_text = """
        Despite NIWA's recognized expertise in QMRA, we face several operational challenges that impact our competitive position:

        • Project delivery inefficiency: Each QMRA project requires rebuilding models from scratch, leading to longer delivery times
        • Lost competitive bids: Recent project losses to competitors who offer faster turnaround times
        • Inconsistent methodology: Manual processes lead to variation in approach across different team members
        • Knowledge retention risk: Institutional QMRA expertise not captured in reusable systems
        • Scaling limitations: Current manual processes prevent efficient handling of large multi-site assessments

        Competitors with standardized tools are gaining market advantage through faster delivery and lower costs, threatening NIWA's market leadership in this critical area.
        """

        para = self.doc.add_paragraph(challenges_text.strip())
        para.style = 'Custom Normal'

        self.doc.add_heading('1.3 Scientific Foundation', level=2).style = 'Custom Heading 2'

        scientific_text = """
        QMRA is a scientifically rigorous methodology that quantifies pathogen-related health risks through four key components: hazard identification, exposure assessment, dose-response modeling, and risk characterization (Ramos et al., 2021). Recent advances in QMRA applications demonstrate its expanding utility across diverse contexts including:

        • Dairy product safety assessment (Ramos et al., 2021)
        • Respiratory pathogen transmission in policy development (Tang et al., 2024)
        • Occupational health risk assessment in wastewater treatment (Dada & Gyawali, 2021)
        • Food safety risk profiling in international trade contexts (Ng et al., 2022)

        New Zealand-specific research highlights the critical importance of localized QMRA applications, particularly for drinking water contamination at outdoor recreation facilities (Phiri et al., 2021) and novel pathogen identification in freshwater systems (Bloomfield et al., 2020). These studies demonstrate both the scientific validity and practical necessity of standardized QMRA approaches for New Zealand's unique environmental and regulatory context.
        """

        para = self.doc.add_paragraph(scientific_text.strip())
        para.style = 'Custom Normal'

    def add_project_objectives(self):
        """Add project objectives and expected outcomes."""
        self.doc.add_heading('2. Project Objectives and Expected Outcomes', level=1).style = 'Custom Heading 1'

        self.doc.add_heading('2.1 Primary Objective', level=2).style = 'Custom Heading 2'

        primary_obj = """
        Develop a comprehensive, Python-based QMRA workflow engine that integrates NIWA's institutional knowledge into standardized, reusable components, reducing project delivery time by 60-70% while ensuring consistent, high-quality methodology across all team members.
        """

        para = self.doc.add_paragraph(primary_obj.strip())
        para.style = 'Custom Normal'

        self.doc.add_heading('2.2 Specific Technical Objectives', level=2).style = 'Custom Heading 2'

        objectives = [
            "Create standardized pathogen dose-response libraries with comprehensive uncertainty distributions",
            "Develop treatment efficacy modules for common water and wastewater treatment processes",
            "Build automated Monte Carlo simulation engine with sensitivity analysis capabilities",
            "Design customizable reporting templates that meet New Zealand regulatory requirements",
            "Integrate existing R code components through Python interfaces (rpy2)",
            "Establish version control and documentation standards for long-term maintainability",
            "Create user-friendly interfaces for both technical and non-technical staff",
            "Develop validation protocols using established datasets"
        ]

        for obj in objectives:
            para = self.doc.add_paragraph(f"• {obj}", style='Custom Normal')

        self.doc.add_heading('2.3 Expected Outcomes and Impact', level=2).style = 'Custom Heading 2'

        outcomes_text = """
        The successful completion of this project will deliver transformational benefits:

        Operational Efficiency:
        • 60-70% reduction in QMRA project delivery time
        • Standardized methodology ensuring consistent quality across all projects
        • Improved capacity to handle multiple concurrent projects
        • Enhanced ability to respond rapidly to urgent regulatory requirements

        Competitive Advantage:
        • Faster turnaround times compared to current market competitors
        • Standardized approach reducing project costs
        • Enhanced credibility through consistent, validated methodologies
        • Capability to bid on larger, multi-site assessments

        Strategic Positioning:
        • Strengthened position as New Zealand's leading QMRA consultancy
        • Enhanced capability to support regulatory agencies with standardized tools
        • Improved knowledge retention and transfer within NIWA
        • Foundation for future tool development and commercialization opportunities
        """

        para = self.doc.add_paragraph(outcomes_text.strip())
        para.style = 'Custom Normal'

    def add_technical_approach(self):
        """Add technical approach and methodology."""
        self.doc.add_heading('3. Technical Approach and Methodology', level=1).style = 'Custom Heading 1'

        self.doc.add_heading('3.1 Technology Platform Selection', level=2).style = 'Custom Heading 2'

        tech_rationale = """
        The workflow engine will be developed in Python, representing a strategic shift from our current R-based approach. This decision is justified by several technical advantages:

        • Superior large-scale data processing capabilities
        • More active development ecosystem with continuously evolving libraries
        • Seamless integration with existing R packages through rpy2 interfaces
        • Comprehensive scientific computing stack (NumPy, SciPy, pandas)
        • Better long-term maintainability and integration with NIWA's data infrastructure
        • Enhanced deployment options for web-based and cloud computing environments
        """

        para = self.doc.add_paragraph(tech_rationale.strip())
        para.style = 'Custom Normal'

        self.doc.add_heading('3.2 System Architecture', level=2).style = 'Custom Heading 2'

        architecture_text = """
        The QMRA workflow engine will follow a modular architecture design:

        Core Engine Components:
        • Pathogen Database Module: Standardized dose-response relationships with uncertainty distributions
        • Treatment Process Module: Efficacy models for filtration, disinfection, and advanced treatment
        • Simulation Engine: Monte Carlo methods with sensitivity and uncertainty analysis
        • Risk Characterization Module: Health outcome calculation and interpretation frameworks

        Data Management Layer:
        • Input/Output standardization with common data formats (CSV, Excel, databases)
        • Version control integration for reproducible analyses
        • Metadata management for traceability and quality assurance

        User Interface Components:
        • Command-line interface for advanced users and automation
        • Configuration files for project-specific parameters
        • Automated report generation with customizable templates
        • Integration hooks for external data sources and visualization tools
        """

        para = self.doc.add_paragraph(architecture_text.strip())
        para.style = 'Custom Normal'

        self.doc.add_heading('3.3 Integration with Existing Systems', level=2).style = 'Custom Heading 2'

        integration_text = """
        The workflow engine will leverage and enhance existing NIWA capabilities:

        • Integration with current R code libraries through rpy2 interfaces
        • Compatibility with NIWA's data management systems and protocols
        • Alignment with established quality assurance procedures
        • Coordination with ongoing work by Mike Hickford on @Risk software transition
        • Foundation for future integration with NIWA's broader modeling platform vision
        """

        para = self.doc.add_paragraph(integration_text.strip())
        para.style = 'Custom Normal'

    def add_work_programme(self):
        """Add detailed work programme and timeline."""
        self.doc.add_heading('4. Work Programme and Timeline', level=1).style = 'Custom Heading 1'

        # Create timeline table
        self.doc.add_heading('4.1 Project Timeline (6 months)', level=2).style = 'Custom Heading 2'

        table = self.doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'

        # Table headers
        headers = ["Phase", "Duration", "Key Activities", "Deliverables"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True

        # Timeline data
        timeline_data = [
            ("Month 1", "4 weeks", "Requirements analysis, system design, R code audit", "System architecture document, Requirements specification"),
            ("Month 2", "4 weeks", "Core engine development, pathogen database creation", "Pathogen dose-response library, Core simulation framework"),
            ("Month 3", "4 weeks", "Treatment module development, Monte Carlo engine", "Treatment efficacy modules, Simulation engine"),
            ("Month 4", "4 weeks", "Report generation, user interface development", "Automated reporting system, User interface prototypes"),
            ("Month 5", "4 weeks", "Integration testing, validation with existing datasets", "Validated workflow engine, Performance benchmarks"),
            ("Month 6", "4 weeks", "Documentation, training materials, final testing", "Complete documentation, User manuals, Training workshops")
        ]

        for i, (phase, duration, activities, deliverables) in enumerate(timeline_data, 1):
            table.cell(i, 0).text = phase
            table.cell(i, 1).text = duration
            table.cell(i, 2).text = activities
            table.cell(i, 3).text = deliverables

        self.doc.add_heading('4.2 Resource Allocation', level=2).style = 'Custom Heading 2'

        resources_text = """
        Personnel Requirements:
        • Reza Moghaddam (Principal Investigator): 150 hours over 6 months
          - System design and architecture
          - Python development and integration
          - Testing and validation
          - Documentation and training material development

        • David Wood (Co-Investigator): 40 hours over 6 months
          - R code integration and optimization
          - Technical review and quality assurance
          - System testing and validation
          - Knowledge transfer and documentation review

        • Supporting team (Andrew Hughes, Michael Bruce): Consultation and review
          - Requirements validation
          - User acceptance testing
          - Strategic guidance and oversight
        """

        para = self.doc.add_paragraph(resources_text.strip())
        para.style = 'Custom Normal'

    def add_budget_justification(self):
        """Add detailed budget and justification."""
        self.doc.add_heading('5. Budget and Justification', level=1).style = 'Custom Heading 1'

        # Budget table
        table = self.doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'

        # Headers
        headers = ["Item", "Rate ($/hour)", "Hours", "Total ($)"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True

        # Budget items
        budget_data = [
            ("Principal Investigator (Reza Moghaddam)", "125", "150", "18,750"),
            ("Co-Investigator (David Wood)", "125", "40", "5,000"),
            ("Computing resources and software", "-", "-", "750"),
            ("Documentation and training materials", "-", "-", "500")
        ]

        for i, (item, rate, hours, total) in enumerate(budget_data, 1):
            table.cell(i, 0).text = item
            table.cell(i, 1).text = rate
            table.cell(i, 2).text = hours
            table.cell(i, 3).text = total

        # Total row
        total_row = table.add_row()
        total_row.cells[0].text = "TOTAL PROJECT COST"
        total_row.cells[0].paragraphs[0].runs[0].font.bold = True
        total_row.cells[3].text = "$25,000"
        total_row.cells[3].paragraphs[0].runs[0].font.bold = True

        justification_text = """
        Budget Justification:

        The requested budget of $25,000 represents excellent value for investment considering:

        • Market Impact: Addresses immediate competitive disadvantage causing lost project bids
        • Efficiency Gains: 60-70% reduction in project delivery time will rapidly recover investment costs
        • Revenue Protection: Maintains NIWA's market leadership in growing QMRA consultancy sector
        • Strategic Positioning: Positions NIWA to capitalize on August 2025 regulatory changes
        • Knowledge Preservation: Captures and systematizes institutional expertise for long-term value

        Return on Investment Analysis:
        • Typical QMRA consulting projects range from $15,000-$50,000
        • Time savings on just 2-3 projects will recover the full investment
        • Enhanced competitiveness expected to increase project win rate by 25-30%
        • Potential revenue uplift of $100,000-$200,000 annually from improved efficiency and competitiveness
        """

        para = self.doc.add_paragraph(justification_text.strip())
        para.style = 'Custom Normal'

    def add_risk_management(self):
        """Add risk assessment and management strategies."""
        self.doc.add_heading('6. Risk Assessment and Management', level=1).style = 'Custom Heading 1'

        # Risk table
        table = self.doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'

        headers = ["Risk", "Probability", "Impact", "Mitigation Strategy"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True

        risks_data = [
            ("Technical integration challenges with existing R code", "Medium", "Medium", "Early prototyping, David Wood's expertise, rpy2 testing"),
            ("Staff availability conflicts", "Low", "High", "Flexible timeline, distributed workload, backup resources"),
            ("Scope creep beyond core functionality", "Medium", "Medium", "Clear requirements specification, regular reviews"),
            ("User adoption challenges", "Low", "Medium", "User involvement in design, comprehensive training"),
            ("Regulatory requirements change during development", "Low", "Low", "Modular design allows rapid adaptation")
        ]

        for i, (risk, prob, impact, mitigation) in enumerate(risks_data, 1):
            table.cell(i, 0).text = risk
            table.cell(i, 1).text = prob
            table.cell(i, 2).text = impact
            table.cell(i, 3).text = mitigation

    def add_references(self):
        """Add references section with QMRA citations."""
        self.doc.add_heading('7. References', level=1).style = 'Custom Heading 1'

        references = [
            "Bloomfield, S., Wilkinson, D., Rogers, L., Biggs, P., French, N., Mohan, V., Savoian, M., Venter, P., & Midwinter, A. (2020). Campylobacter novaezeelandiae sp. nov., isolated from birds and water in New Zealand. International Journal of Systematic and Evolutionary Microbiology, 70(6), 3775-3784.",

            "Dada, A.C., & Gyawali, P. (2021). Quantitative microbial risk assessment (QMRA) of occupational exposure to SARS-CoV-2 in wastewater treatment plants. Science of the Total Environment, 763, 142989.",

            "Ng, S., Shao, S., & Ling, N. (2022). Food safety risk-assessment systems utilized by China, Australia/New Zealand, Canada, and the United States. Journal of Food Science, 87(11), 4780-4795.",

            "Phiri, B.J., French, N.P., Biggs, P.J., Stevenson, M.A., Reynolds, A.D., Garcia, R.J., & Hayman, D.T.S. (2021). Microbial contamination in drinking water at public outdoor recreation facilities in New Zealand. Journal of Applied Microbiology, 130(1), 302-312.",

            "Ramos, G.L.P.A., Nascimento, J.S., Margalho, L.P., Duarte, M.C.K.H., Esmerino, E.A., Freitas, M.Q., Cruz, A.G., & Sant'Ana, A.S. (2021). Quantitative microbiological risk assessment in dairy products: Concepts and applications. Trends in Food Science & Technology, 111, 610-616.",

            "Tang, L., Rhoads, W.J., Eichelberg, A., Hamilton, K.A., & Julian, T.R. (2024). Applications of quantitative microbial risk assessment to respiratory pathogens and implications for uptake in policy: A state-of-the-science review. Environmental Health Perspectives, 132(5), 56001."
        ]

        for ref in references:
            para = self.doc.add_paragraph(ref, style='Custom Normal')
            para.paragraph_format.hanging_indent = Inches(0.5)

    def add_appendices(self):
        """Add appendices with supporting information."""
        self.doc.add_heading('Appendix A: Technical Specifications', level=1).style = 'Custom Heading 1'

        tech_specs = """
        System Requirements:
        • Python 3.8+ with scientific computing libraries (NumPy, SciPy, pandas)
        • R 4.0+ with rpy2 interface for legacy code integration
        • Database support (SQLite/PostgreSQL) for pathogen and treatment data
        • Web framework capability (Flask/Django) for future GUI development
        • Version control integration (Git) with NIWA repositories
        • Documentation generation (Sphinx) for automated technical documentation

        Data Standards:
        • Input/output formats: CSV, Excel, JSON, database connections
        • Metadata standards for traceability and reproducibility
        • Quality assurance protocols aligned with NIWA procedures
        • Backup and version control for all project data and configurations

        Performance Targets:
        • Simulation speed: >1000 iterations per second for standard QMRA models
        • Memory efficiency: Handle datasets with >100,000 exposure scenarios
        • Scalability: Support concurrent analysis of multiple sites/scenarios
        • Reliability: >99% uptime for production use
        """

        para = self.doc.add_paragraph(tech_specs.strip())
        para.style = 'Custom Normal'

        self.doc.add_heading('Appendix B: Stakeholder Engagement Plan', level=1).style = 'Custom Heading 1'

        stakeholder_plan = """
        Internal Stakeholders:
        • QMRA team members: Monthly progress reviews and feedback sessions
        • IT services: Coordination on system deployment and maintenance
        • Quality assurance: Integration with NIWA's QA procedures and standards
        • Management: Quarterly progress reports and strategic alignment reviews

        External Stakeholders:
        • Regulatory agencies (Taumata Arowai, regional councils): Consultation on reporting requirements
        • Industry partners: Feedback on practical implementation needs
        • Academic collaborators: Peer review of technical approaches and validation

        Communication Strategy:
        • Monthly technical newsletters highlighting development progress
        • Webinar series on QMRA applications and tool capabilities
        • Conference presentations at water industry and scientific meetings
        • Publication of technical reports and peer-reviewed articles
        """

        para = self.doc.add_paragraph(stakeholder_plan.strip())
        para.style = 'Custom Normal'

    def generate_proposal(self, filename="SIP_QMRA_Proposal.docx"):
        """Generate the complete SIP proposal document."""
        print("Generating SIP Proposal for QMRA Workflow Engine...")

        # Add all sections
        self.add_title_page()
        self.add_executive_summary()
        self.add_project_background()
        self.add_project_objectives()
        self.add_technical_approach()
        self.add_work_programme()
        self.add_budget_justification()
        self.add_risk_management()
        self.add_references()
        self.add_appendices()

        # Save document
        filepath = os.path.join(os.getcwd(), filename)
        self.doc.save(filepath)
        print(f"SIP Proposal saved as: {filepath}")

        return filepath

def main():
    """Main function to generate the SIP proposal."""
    generator = SIPProposalGenerator()
    proposal_path = generator.generate_proposal()

    print("\n" + "="*60)
    print("SIP PROPOSAL GENERATION COMPLETE")
    print("="*60)
    print(f"Document saved as: {proposal_path}")
    print("\nProposal Summary:")
    print("• Project: QMRA Workflow Engine Development")
    print("• Budget: $25,000")
    print("• Duration: 6 months")
    print("• Expected impact: 60-70% reduction in project delivery time")
    print("• Strategic positioning for August 2025 regulatory changes")
    print("\nNext steps:")
    print("1. Review the generated document")
    print("2. Customize specific budget details if needed")
    print("3. Add any additional institutional requirements")
    print("4. Submit through NIWA's SIP application process")

if __name__ == "__main__":
    main()