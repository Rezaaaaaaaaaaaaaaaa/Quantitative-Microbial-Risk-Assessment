#!/usr/bin/env python3
"""
NIWA SIP Proposal - QMRA Workflow Engine (Revised)
Addresses Andrew Hughes' comments and follows exact template format
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import FancyBboxPatch, Rectangle
from io import BytesIO

class RevisedQMRAProposal:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Setup document styles."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)

    def create_concise_workflow_diagram(self):
        """Create a concise, professional workflow diagram."""
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.axis('off')

        # Title
        ax.text(5, 5.5, 'QMRA Workflow Engine', fontsize=16, fontweight='bold', ha='center')

        # Define clean colors
        input_color = '#E3F2FD'    # Light blue
        process_color = '#BBDEFB'   # Medium blue
        output_color = '#E8F5E8'    # Light green

        # INPUT LAYER - Simplified
        inputs = [
            {'x': 0.5, 'y': 4, 'w': 2, 'h': 0.6, 'text': 'Water Quality\nData'},
            {'x': 3, 'y': 4, 'w': 2, 'h': 0.6, 'text': 'Treatment\nPerformance'},
            {'x': 5.5, 'y': 4, 'w': 2, 'h': 0.6, 'text': 'Pathogen\nData'},
            {'x': 8, 'y': 4, 'w': 1.5, 'h': 0.6, 'text': 'Standards'}
        ]

        for inp in inputs:
            rect = FancyBboxPatch((inp['x'], inp['y']), inp['w'], inp['h'],
                                 boxstyle="round,pad=0.05",
                                 facecolor=input_color, edgecolor='#1976D2', linewidth=1)
            ax.add_patch(rect)
            ax.text(inp['x'] + inp['w']/2, inp['y'] + inp['h']/2, inp['text'],
                   ha='center', va='center', fontsize=9, fontweight='bold')

        # PROCESSING ENGINE - Single unified box
        engine_rect = FancyBboxPatch((1, 2.5), 8, 0.8,
                                   boxstyle="round,pad=0.05",
                                   facecolor=process_color, edgecolor='#1565C0', linewidth=2)
        ax.add_patch(engine_rect)
        ax.text(5, 2.9, 'PYTHON QMRA ENGINE', ha='center', va='center',
               fontsize=12, fontweight='bold')
        ax.text(5, 2.6, 'Dose-Response Models • Monte Carlo Simulation • Risk Calculation',
               ha='center', va='center', fontsize=9)

        # OUTPUT LAYER - Simplified
        outputs = [
            {'x': 1, 'y': 1, 'w': 2.5, 'h': 0.6, 'text': 'Risk Assessment\nResults'},
            {'x': 3.75, 'y': 1, 'w': 2.5, 'h': 0.6, 'text': 'Compliance\nReports'},
            {'x': 6.5, 'y': 1, 'w': 2, 'h': 0.6, 'text': 'Documentation'}
        ]

        for out in outputs:
            rect = FancyBboxPatch((out['x'], out['y']), out['w'], out['h'],
                                 boxstyle="round,pad=0.05",
                                 facecolor=output_color, edgecolor='#388E3C', linewidth=1)
            ax.add_patch(rect)
            ax.text(out['x'] + out['w']/2, out['y'] + out['h']/2, out['text'],
                   ha='center', va='center', fontsize=9, fontweight='bold')

        # Add clean arrows
        arrow_props = dict(arrowstyle='->', lw=2, color='#424242')

        # Inputs to engine
        for inp in inputs:
            start_x = inp['x'] + inp['w']/2
            ax.annotate('', xy=(start_x, 3.3), xytext=(start_x, inp['y']),
                       arrowprops=arrow_props)

        # Engine to outputs
        for out in outputs:
            end_x = out['x'] + out['w']/2
            ax.annotate('', xy=(end_x, out['y'] + out['h']), xytext=(end_x, 2.5),
                       arrowprops=arrow_props)

        # Add workflow labels
        ax.text(5, 3.8, 'INPUT', ha='center', va='center', fontsize=10,
               fontweight='bold', color='#1976D2')
        ax.text(5, 0.5, 'OUTPUT', ha='center', va='center', fontsize=10,
               fontweight='bold', color='#388E3C')

        # Save to buffer
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight',
                   facecolor='white', edgecolor='none', pad_inches=0.2)
        img_buffer.seek(0)
        plt.close()

        return img_buffer

    def add_header(self):
        """Add NIWA header."""
        header_para = self.doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run("NIWA")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_para.space_after = Pt(12)

    def add_title_section(self):
        """Add title section."""
        title = self.doc.add_paragraph("Structured internal project application 2025-2026")
        title.space_after = Pt(12)

        overview = self.doc.add_paragraph("Project Overview")
        overview.space_after = Pt(12)

    def add_project_overview_table(self):
        """Add the main project overview table following exact template format."""
        table = self.doc.add_table(rows=12, cols=2)
        table.style = 'Table Grid'

        # Set column widths to match template
        for row in table.rows:
            row.cells[0].width = Inches(2.8)
            row.cells[1].width = Inches(4.2)

        # Project name
        table.cell(0, 0).text = "Project name: (Short title)"
        table.cell(0, 1).text = "Development of QMRA Workflow Engine"

        # Staff
        table.cell(1, 0).text = "Staff: (who will be completing the work?)"
        table.cell(1, 1).text = "Reza Moghaddam (Lead Developer - 150 hrs), David Wood (Model Review & Support - 40 hrs)"

        # Project Manager
        table.cell(2, 0).text = "Project Manager: (usually a Group Manager)"
        table.cell(2, 1).text = "Andrew Hughes"

        # Region
        table.cell(3, 0).text = "Region:"
        table.cell(3, 1).text = "Hamilton"

        # Centre
        table.cell(4, 0).text = "Centre:"
        table.cell(4, 1).text = "Freshwater"

        # Type
        table.cell(5, 0).text = "Type: (science, operations activity, or other - explain)"
        table.cell(5, 1).text = "Science (Applied Research & Development)"

        # Budget - Left empty as per Andrew's feedback
        table.cell(6, 0).text = "Budget: (attach costing prepared by your project coordinator)"
        table.cell(6, 1).text = ""

        # Project objective
        table.cell(7, 0).text = "Project objective: (30 words max)"
        table.cell(7, 1).text = "Develop a Python-based QMRA workflow engine to standardise processes, reduce manual work, and improve efficiency of regulatory compliance assessments."

        # Project outline - Updated to address Andrew's comments
        table.cell(8, 0).text = "Project outline: (150-300 words max)"
        outline_text = """NIWA currently undertakes QMRA projects that require significant manual effort for each assessment. Based on our recent project experience, typical QMRA projects involve 40-60 hours of manual work including dose-response model setup, treatment calculations, simulation configuration, and report generation.

This project will develop a Python-based QMRA workflow engine to standardise these processes. The system will automate routine calculations, provide validated dose-response relationships for common pathogens, and generate standardised reporting templates. This approach will reduce manual effort and improve consistency across projects.

The technical implementation will use Python's scientific libraries (NumPy, SciPy, pandas) for computational tasks and create reusable modules for pathogen databases, treatment assessment, and risk simulation. The modular design will allow for easy updates and extensions as new requirements emerge.

NIWA has established capabilities in QMRA through previous projects for regulatory clients. The workflow engine will enhance our ability to deliver timely, consistent assessments while maintaining technical rigour. The system will be designed to support regulatory compliance requirements and provide clear documentation for decision-making."""

        table.cell(8, 1).text = outline_text

        # Project outputs
        table.cell(9, 0).text = "Project outputs: (e.g., a journal paper or an App, or a safe operating procedure or guidance document for operations activities)"
        table.cell(9, 1).text = "• QMRA Workflow Engine (Python application)\n• Technical documentation and user guide\n• Standardised pathogen database with dose-response relationships\n• Template reporting system for regulatory compliance"

        # Project impact
        table.cell(10, 0).text = "Project impact: (choose an SCI impact area that the project aligns with, see graphic below)"
        table.cell(10, 1).text = "Protecting our diversity\nImproved environmental health"

        # Alignment
        table.cell(11, 0).text = "Alignment: (with a programme and/or National Centre outcomes or KPIs)"
        table.cell(11, 1).text = "This project aligns with the Freshwater Centre's analytical capabilities development and supports regulatory compliance services. It enhances NIWA's technical capacity for water quality risk assessment and supports our role in environmental protection."

        # Make labels bold
        for i in range(12):
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True

    def add_workflow_diagram(self):
        """Add the concise workflow diagram."""
        # Add some space before diagram
        self.doc.add_paragraph()

        # Create centered paragraph for the diagram
        diagram_para = self.doc.add_paragraph()
        diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generate and insert the diagram
        img_buffer = self.create_concise_workflow_diagram()
        run = diagram_para.add_run()
        run.add_picture(img_buffer, width=Inches(6))

        # Add figure caption
        caption_para = self.doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run("Figure 1: QMRA Workflow Engine - Streamlined architecture for efficient risk assessment")
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        caption_para.space_after = Pt(12)

    def add_additional_sections(self):
        """Add the remaining template sections."""
        # Outcomes for Māori
        maori_table = self.doc.add_table(rows=1, cols=2)
        maori_table.style = 'Table Grid'
        maori_table.cell(0, 0).width = Inches(2.8)
        maori_table.cell(0, 1).width = Inches(4.2)
        maori_table.cell(0, 0).text = "Outcomes for Māori: (may include partnerships, resourcing, alignment with aspirations)"
        maori_table.cell(0, 1).text = "Supporting improved water quality assessment capabilities that contribute to protecting water bodies important for cultural values and mahinga kai."
        maori_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True

        self.doc.add_paragraph()

        # Operations alignment
        ops_table = self.doc.add_table(rows=1, cols=2)
        ops_table.style = 'Table Grid'
        ops_table.cell(0, 0).width = Inches(2.8)
        ops_table.cell(0, 1).width = Inches(4.2)
        ops_table.cell(0, 0).text = "Operations alignment: (for non-science projects, how does this work contribute to inputs or enablers from the graphic below)"
        ops_table.cell(0, 1).text = "Not applicable"
        ops_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True

    def add_work_programme(self):
        """Add work programme section following template format."""
        self.doc.add_page_break()

        # Work programme heading
        work_heading = self.doc.add_paragraph("WORK PROGRAMME AND TIMELINE")
        work_heading.space_after = Pt(12)

        # Instructions text
        instruction_para = self.doc.add_paragraph("Outline the tasks to be done, who will do what and by when. Be as specific as possible.")
        instruction_para.space_after = Pt(12)

        # Work programme table
        table = self.doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'

        # Headers
        headers = ["Task", "Specific activity (who, what)", "By when", "Hours"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Work data - Realistic timeline and hours
        work_data = [
            ("Requirements & Design", "System architecture definition, QMRA methodology analysis, stakeholder consultation (Reza)", "31/01/2025", "30"),
            ("Core Development", "Pathogen database creation, dose-response model implementation, Python framework development (Reza)", "28/02/2025", "60"),
            ("Advanced Features", "Monte Carlo simulation engine, statistical modelling implementation (Reza)", "31/03/2025", "35"),
            ("Testing & Validation", "Performance testing, validation against known benchmarks, quality assurance (Reza)", "30/04/2025", "25"),
            ("Model Review & Validation", "Technical review of QMRA models, validation of dose-response relationships (David)", "30/04/2025", "25"),
            ("Documentation", "Technical documentation, user guides, training materials (David)", "31/05/2025", "15"),
            ("Deployment & Transfer", "System deployment, staff training, knowledge transfer protocols (Reza/David)", "30/06/2025", "10")
        ]

        for i, (task, activity, when, hours) in enumerate(work_data, 1):
            table.cell(i, 0).text = task
            table.cell(i, 1).text = activity
            table.cell(i, 2).text = when
            table.cell(i, 3).text = hours

    def add_chief_scientist_section(self):
        """Add Chief Scientist support section."""
        self.doc.add_paragraph()

        # Chief Scientist heading
        cs_heading = self.doc.add_paragraph("CHIEF SCIENTIST SUPPORT")
        cs_heading.space_after = Pt(12)

        # Comment section
        comment_para = self.doc.add_paragraph()
        comment_para.add_run("Chief Scientist comment: ").font.bold = True
        comment_para.add_run("(For example - If agreement that project required, indicate why SIP mechanism versus Centre Funds; What is/are the key output(s) and how will NIWA/National Centre/programme/individual benefit from that; note that there must be an output at the end of the project)")
        comment_para.space_after = Pt(12)

        # Response area
        response_para = self.doc.add_paragraph()
        response_para.space_after = Pt(24)

        # Signature
        sig_para = self.doc.add_paragraph("Signature")
        sig_para.space_after = Pt(24)

        # Reference note
        ref_para = self.doc.add_paragraph("For reference: NIWA Impact Strategy, also see Statement of Corporate Intent")
        ref_para.space_after = Pt(12)

    def add_collaboration_section(self):
        """Add collaboration opportunities section - addressing Andrew's comments."""
        collab_heading = self.doc.add_paragraph("EMERGING COLLABORATION OPPORTUNITIES")
        collab_heading.space_after = Pt(12)

        # Narrative format as requested by Andrew, explaining PHF
        collab_text = """Recent developments have strengthened the business case for this QMRA workflow engine. Primary Health Foundation (PHF), which provides health services to rural communities, has approached NIWA to develop QMRA guidance specifically for shellfish safety assessment. PHF has confirmed their interest through direct communication with Taumata Arowai about collaborating with NIWA on this initiative.

This emerging opportunity demonstrates immediate market validation for our QMRA capabilities and provides direct application potential for the workflow engine in shellfish safety assessment. The collaboration creates strategic partnership opportunities with regulatory bodies and demonstrates that there is demand for NIWA's enhanced QMRA services.

The shellfish QMRA guidance project would serve as an ideal pilot application for our workflow engine, providing real-world validation while generating project revenue. This collaboration would allow concurrent testing and refinement of the system with actual regulatory requirements, potentially offsetting some development costs through direct application to a paying project."""

        self.doc.add_paragraph(collab_text)

    def generate_proposal(self, filename="NIWA_QMRA_SIP_Final_with_Diagram.docx"):
        """Generate the complete revised proposal."""
        print("Generating revised NIWA SIP proposal addressing all comments...")

        # Add all sections in template order
        self.add_header()
        self.add_title_section()
        self.add_project_overview_table()
        self.add_workflow_diagram()
        self.add_additional_sections()
        self.add_work_programme()
        self.add_collaboration_section()
        self.add_chief_scientist_section()

        # Save document
        filepath = os.path.join(os.getcwd(), filename)
        self.doc.save(filepath)
        print(f"Revised NIWA SIP proposal saved as: {filepath}")

        return filepath

def main():
    """Generate the revised SIP proposal."""
    generator = RevisedQMRAProposal()
    proposal_path = generator.generate_proposal()

    print("\n" + "="*80)
    print("REVISED NIWA SIP PROPOSAL COMPLETE")
    print("="*80)
    print(f"Document: {proposal_path}")
    print("\nAddressed Comments:")
    print("• Removed speculative numbers and hyperbolic language")
    print("• Explained PHF (Primary Health Foundation)")
    print("• Converted bullet points to narrative format")
    print("• Followed exact template structure from example.pdf")
    print("• Used realistic project timelines and hours")
    print("• Focused on factual benefits and capabilities")
    print("• Added improved concise workflow diagram")
    print("• Added proper template sections (Māori outcomes, operations alignment)")

if __name__ == "__main__":
    main()