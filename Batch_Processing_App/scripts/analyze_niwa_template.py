#!/usr/bin/env python3
"""
Analyze NIWA Report Template Structure
"""

from docx import Document
from pathlib import Path

def analyze_template():
    """Analyze the NIWA report template structure."""

    template_path = Path("../../NIWA QMRA report.docx")

    if not template_path.exists():
        print(f"Template not found: {template_path}")
        return

    print(f"Analyzing: {template_path.absolute()}\n")
    print("="*80)

    doc = Document(str(template_path))

    # Analyze document properties
    print("\nDOCUMENT PROPERTIES:")
    print("-"*80)
    core_props = doc.core_properties
    print(f"Title: {core_props.title}")
    print(f"Subject: {core_props.subject}")
    print(f"Author: {core_props.author}")
    print(f"Keywords: {core_props.keywords}")
    print(f"Comments: {core_props.comments}")

    # Analyze sections
    print("\nSECTIONS:")
    print("-"*80)
    for i, section in enumerate(doc.sections, 1):
        print(f"\nSection {i}:")
        print(f"  Page height: {section.page_height}")
        print(f"  Page width: {section.page_width}")
        print(f"  Left margin: {section.left_margin}")
        print(f"  Right margin: {section.right_margin}")

    # Analyze paragraphs
    print("\nPARAGRAPHS STRUCTURE:")
    print("-"*80)
    print(f"Total paragraphs: {len(doc.paragraphs)}\n")

    for i, para in enumerate(doc.paragraphs[:50], 1):  # First 50 paragraphs
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            print(f"{i:3d}. [{style:20s}] {para.text[:80]}")

    if len(doc.paragraphs) > 50:
        print(f"\n... and {len(doc.paragraphs) - 50} more paragraphs")

    # Analyze tables
    print("\nTABLES:")
    print("-"*80)
    print(f"Total tables: {len(doc.tables)}\n")

    for i, table in enumerate(doc.tables, 1):
        print(f"Table {i}: {len(table.rows)} rows × {len(table.columns)} columns")
        if table.rows:
            print(f"  First row: {[cell.text[:30] for cell in table.rows[0].cells]}")

    # Analyze styles
    print("\nAVAILABLE STYLES:")
    print("-"*80)
    styles_used = set()
    for para in doc.paragraphs:
        if para.style:
            styles_used.add(para.style.name)

    for style in sorted(styles_used):
        print(f"  - {style}")

    print("\n" + "="*80)
    print("Analysis complete!")

if __name__ == '__main__':
    analyze_template()
