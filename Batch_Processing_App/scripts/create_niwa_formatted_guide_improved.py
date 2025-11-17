#!/usr/bin/env python3
"""
Create QMRA User Guide Using NIWA Template Format - IMPROVED VERSION
=====================================================================

Professional NIWA-formatted document with complete sections.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
import shutil

def create_niwa_user_guide():
    """Create comprehensive user guide using NIWA template."""

    # Copy template to new document
    template_path = Path("../../NIWA QMRA report.docx")
    output_path = Path("../../QMRA_User_Guide_NIWA_Format.docx")

    print(f"\nCreating NIWA-formatted QMRA User Guide...")
    print(f"Template: {template_path}")
    print(f"Output: {output_path}\n")

    # Load the template
    doc = Document(str(template_path))

    print(f"Loaded template with {len(doc.paragraphs)} paragraphs")
    print(f"Sections: {len(doc.sections)}")
    print(f"Tables: {len(doc.tables)}\n")

    # COVER PAGE (if template doesn't have one, add it)
    print("Adding cover page...")
    para = doc.add_paragraph()
    para.style = 'z_Report Title'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('QMRA Web Application\nUser Guide and Technical Documentation')

    para = doc.add_paragraph()
    para.style = 'z_Report Subtitle'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Quantitative Microbial Risk Assessment Tool\nfor Waterborne Pathogen Exposure')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(f'\n\nPrepared for:\nNIWA (National Institute of Water and Atmospheric Research)\n\n')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(f'Version 1.0\n{datetime.now().strftime("%B %Y")}')

    doc.add_page_break()

    # TABLE OF CONTENTS
    print("Adding table of contents...")
    para = doc.add_paragraph()
    para.style = 'Heading 1'
    para.add_run('Table of Contents')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('[Table of Contents will be auto-generated in Microsoft Word]\n')
    para.add_run('To update: Right-click → Update Field → Update Entire Table')

    doc.add_page_break()

    # EXECUTIVE SUMMARY
    print("Adding Executive Summary...")
    para = doc.add_paragraph()
    para.style = 'z_Exec Summary heading'
    para.add_run('Executive Summary')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Purpose: ')
    run.bold = True
    para.add_run(
        'This document provides comprehensive guidance for using the QMRA Web Application, '
        'a professional tool for quantitative microbial risk assessment (QMRA) of waterborne '
        'pathogens in New Zealand water environments.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Scope: ')
    run.bold = True
    para.add_run(
        'The application evaluates public health risks from pathogen exposure through '
        'shellfish consumption, recreational water contact, and drinking water using validated '
        'dose-response models and Monte Carlo simulation.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Production Mode (Default): ')
    run.bold = True
    para.add_run(
        'This application implements Production Mode restricting pathogen selection to norovirus '
        'with the validated Beta-Binomial dose-response model (α=0.04, β=0.055 from Teunis et al. 2008). '
        'This ensures all risk assessments meet scientific standards and align with project deliverables.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Key Features:')
    run.bold = True

    # Use proper NIWA bullet style
    bullets = [
        'Five specialized assessment workflows for different risk scenarios',
        'Validated Beta-Binomial dose-response model (exact formulation, not Beta-Poisson approximation)',
        'Monte Carlo simulation with 10,000 iterations for uncertainty quantification',
        'WHO guideline compliance checking (10⁻⁴ annual infection risk target)',
        'Interactive visualizations and comprehensive statistical summaries',
        'Professional PDF reports and CSV data exports for regulatory submissions'
    ]

    for bullet in bullets:
        para = doc.add_paragraph(bullet, style='Bullet Level 1')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Document Structure: ')
    run.bold = True
    para.add_run(
        'This guide covers system requirements, installation, all assessment workflows with '
        'step-by-step procedures, results interpretation, technical background on the Beta-Binomial '
        'model, troubleshooting, and limitations.'
    )

    doc.add_page_break()

    # 1. INTRODUCTION
    print("Adding Introduction...")
    heading = doc.add_heading('Introduction to QMRA', level=1)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'Quantitative Microbial Risk Assessment (QMRA) is a scientific framework for estimating '
        'the probability of infection or illness from exposure to pathogenic microorganisms in '
        'water, food, or environmental media. Unlike traditional water quality monitoring that '
        'relies solely on indicator organisms (e.g., E. coli), QMRA directly estimates health '
        'risks by modeling the entire exposure pathway from source to receptor.'
    )

    # 1.1 QMRA Framework
    heading = doc.add_heading('The QMRA Framework', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('The QMRA process integrates four key components:')

    components = [
        ('Hazard Identification', 'Identifying pathogenic microorganisms of concern (e.g., norovirus, Cryptosporidium)'),
        ('Exposure Assessment', 'Quantifying pathogen concentrations in source water and calculating human exposure doses'),
        ('Dose-Response Modeling', 'Relating pathogen dose to probability of infection using validated mathematical models'),
        ('Risk Characterization', 'Calculating annual infection/illness probabilities and comparing to health-based guidelines')
    ]

    for title, desc in components:
        para = doc.add_paragraph()
        para.style = 'Body Text'
        run = para.add_run(f'{title}: ')
        run.bold = True
        para.add_run(desc)

    # 1.2 Application Purpose
    heading = doc.add_heading('Application Purpose and Scope', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'This web-based application enables water quality professionals to conduct comprehensive '
        'QMRA studies without requiring specialized programming skills or advanced statistical software. '
        'Primary use cases include:'
    )

    use_cases = [
        'Evaluating health risks from wastewater discharge to shellfish harvesting areas',
        'Assessing recreational water quality at swimming beaches and coastal locations',
        'Comparing treatment technology effectiveness for pathogen reduction',
        'Analyzing temporal risk trends from long-term monitoring programs',
        'Supporting regulatory compliance and water safety plan development',
        'Shellfish harvesting area classification (approved, conditionally approved, restricted)'
    ]

    for case in use_cases:
        para = doc.add_paragraph(case, style='Bullet Level 1')

    doc.add_page_break()

    # 2. GETTING STARTED
    print("Adding Getting Started section...")
    heading = doc.add_heading('Getting Started', level=1)

    # 2.1 System Requirements
    heading2 = doc.add_heading('System Requirements', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('The QMRA application requires the following:')

    requirements = [
        'Operating System: Windows 10/11, macOS 10.14+, or Linux',
        'Python: Version 3.8 or higher',
        'RAM: Minimum 4 GB (8 GB recommended for large datasets)',
        'Browser: Chrome, Firefox, Safari, or Edge (latest versions)',
        'Internet Connection: Required for initial installation only',
        'Disk Space: Approximately 500 MB for application and dependencies'
    ]

    for req in requirements:
        para = doc.add_paragraph(req, style='Bullet Level 1')

    # 2.2 Installation
    heading2 = doc.add_heading('Installation and Setup', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Step 1: Install Python Dependencies')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Open a terminal/command prompt in the application directory and run:')

    para = doc.add_paragraph('pip install -r requirements.txt', style='Normal')
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Step 2: Verify Installation')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Test the installation by running:')

    para = doc.add_paragraph('python -c "import streamlit, pandas, numpy; print(\'Success!\')"', style='Normal')
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # 2.3 Launching
    heading2 = doc.add_heading('Launching the Application', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Method 1: Using the Batch File (Windows)')
    run.bold = True

    steps = [
        'Navigate to: Batch_Processing_App/app/',
        'Double-click: launch_web_gui.bat',
        'Your default web browser will automatically open to http://localhost:8501'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Numbered para - Level 1')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Method 2: Using Command Line (All Platforms)')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Open a terminal and run:')

    para = doc.add_paragraph('cd Batch_Processing_App/app', style='Normal')
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    para = doc.add_paragraph('streamlit run web_app.py', style='Normal')
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # 2.4 Production Mode
    heading2 = doc.add_heading('Production Mode vs Research Mode', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('CRITICAL: Production Mode (Default Setting)')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'Production Mode is enabled by default and restricts pathogen selection to norovirus only. '
        'This ensures all risk assessments use the validated Beta-Binomial dose-response model and '
        'produce scientifically defensible results for regulatory submissions.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Production Mode provides:')

    features = [
        'Norovirus-only pathogen selection (highly infectious waterborne virus)',
        'Validated Beta-Binomial dose-response model (α=0.04, β=0.055)',
        'Complete verification against reference Excel calculations (David Wood, NIWA)',
        'Alignment with project contract deliverables and scope',
        'Scientifically defensible results for regulatory compliance'
    ]

    for feature in features:
        para = doc.add_paragraph(feature, style='Bullet Level 1')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Note: ')
    run.bold = True
    para.add_run(
        'Research Mode (unchecking Production Mode checkbox) allows selection of additional '
        'pathogens for exploratory analysis, but results should not be used for regulatory submissions '
        'without independent validation.'
    )

    doc.add_page_break()

    # 3. BETA-BINOMIAL MODEL (CRITICAL SECTION)
    print("Adding Beta-Binomial model section...")
    heading = doc.add_heading('Dose-Response Model: Beta-Binomial (Norovirus)', level=1)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('THE CRITICAL DISTINCTION FOR NOROVIRUS')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(200, 0, 0)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'This application uses the EXACT Beta-Binomial dose-response model for norovirus, which is the '
        'mathematically correct formulation. This is NOT the same as the Beta-Poisson approximation.'
    )

    heading2 = doc.add_heading('Beta-Binomial Formula (CORRECT)', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('The exact mathematical formula implemented in this application:')

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run(
        'P(infection) = 1 - exp[ln Γ(β+dose) + ln Γ(α+β) - ln Γ(α+β+dose) - ln Γ(β)]'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 100, 0)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Parameters (Teunis et al. 2008):')
    run.bold = True

    params = [
        'α (alpha) = 0.04',
        'β (beta) = 0.055',
        'Source: Teunis et al. (2008) "Norwalk virus: How infectious is it?" J. Med. Virol. 80:1468-1476'
    ]

    for param in params:
        para = doc.add_paragraph(param, style='Bullet Level 1')

    heading2 = doc.add_heading('Why Beta-Poisson is INVALID for Norovirus', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('CRITICAL WARNING: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    para.add_run(
        'The Beta-Poisson approximation is mathematically INVALID for norovirus and produces '
        'incorrect risk estimates.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Validity requirement: β >> 1 (beta must be much greater than 1)')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('For norovirus: β = 0.055 << 1 (violates validity requirement)')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Consequences of using Beta-Poisson (incorrect):')
    run.bold = True

    consequences = [
        'Underestimates risk by 2-4× at low doses (1-100 virions)',
        'Produces scientifically indefensible results',
        'Does not match validated reference calculations',
        'Not acceptable for regulatory submissions'
    ]

    for consequence in consequences:
        para = doc.add_paragraph(consequence, style='Bullet Level 1')

    heading2 = doc.add_heading('Validation and Verification', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'This application\'s Beta-Binomial implementation has been rigorously validated against '
        'reference Excel spreadsheets (David Wood, NIWA) and produces exact agreement for all test doses.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Example comparison (1 virion dose):')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('Beta-Binomial (CORRECT): 42.1% infection probability')
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('Beta-Poisson (WRONG): 11.1% infection probability')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('Underestimation factor: 3.8× (unacceptable error)')
    run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_page_break()

    # 4. ASSESSMENT WORKFLOWS
    print("Adding Assessment Workflows section...")
    heading = doc.add_heading('Assessment Workflows', level=1)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'The application provides five specialized workflows, each designed for specific '
        'risk assessment scenarios. This section provides detailed step-by-step instructions for each workflow.'
    )

    # 4.1 Batch Scenarios
    heading2 = doc.add_heading('Batch Scenarios Assessment', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Purpose: ')
    run.bold = True
    para.add_run('Run multiple pre-configured scenarios simultaneously (15+ scenarios)')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Step-by-step procedure:')

    steps = [
        'Select "Batch Scenarios" from Assessment Mode dropdown',
        'Check "Use example data" to load 15 pre-configured scenarios',
        'Review pathogen_data.csv, scenarios.csv, and master_scenarios.csv in preview tabs',
        'Click "Run Batch Assessment"',
        'Review results table showing all scenarios with risk metrics',
        'Download batch_scenarios_results.csv for further analysis'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Numbered para - Level 1')

    # 4.2 Spatial Assessment
    heading2 = doc.add_heading('Spatial Risk Assessment', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Purpose: ')
    run.bold = True
    para.add_run('Evaluate risk across multiple sites with different dilution factors')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Step-by-step procedure:')

    steps = [
        'Select "Spatial Assessment" from Assessment Mode dropdown',
        'Configure: Pathogen (norovirus), Exposure Route (Shellfish Consumption or Primary Contact)',
        'Upload dilution data CSV or check "Use example data" (6 sites, dilution 100-100,000)',
        'Set pathogen concentration (e.g., 1×10⁶ organisms/L for raw wastewater)',
        'Set treatment LRV (Log Reduction Value, typically 2-4 for secondary treatment)',
        'Set exposure frequency (e.g., 12 events/year for monthly shellfish consumption)',
        'Click "Run Spatial Assessment"',
        'Review site-by-site risk table, WHO compliance status, risk gradient visualization',
        'Download results as CSV and plots as PNG'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Numbered para - Level 1')

    # 4.3 Temporal Assessment
    heading2 = doc.add_heading('Temporal Risk Assessment', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Purpose: ')
    run.bold = True
    para.add_run('Analyze risk trends over time from monitoring data')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Step-by-step procedure:')

    steps = [
        'Select "Temporal Assessment" from Assessment Mode dropdown',
        'Choose pathogen (norovirus) and exposure route',
        'Upload time-series CSV with columns: date, pathogen_concentration, dilution_factor',
        'Set treatment LRV and exposure parameters',
        'Click "Run Temporal Assessment"',
        'Review time-series plot showing risk variation over time',
        'Identify seasonal patterns or exceedance events',
        'Download temporal_results.csv'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Numbered para - Level 1')

    # 4.4 Treatment Comparison
    heading2 = doc.add_heading('Treatment Technology Comparison', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Purpose: ')
    run.bold = True
    para.add_run('Compare multiple treatment technologies side-by-side')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Step-by-step procedure:')

    steps = [
        'Select "Treatment Comparison" from Assessment Mode dropdown',
        'Choose pathogen and exposure route',
        'Configure multiple treatment scenarios (e.g., "No Treatment", "Secondary", "UV Disinfection")',
        'Set LRV for each treatment (e.g., 0, 2, 4)',
        'Set common exposure parameters',
        'Click "Run Treatment Comparison"',
        'Review comparative bar charts showing risk reduction',
        'Identify most effective treatment for compliance',
        'Download comparison results'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Numbered para - Level 1')

    # 4.5 Multi-Pathogen Assessment
    heading2 = doc.add_heading('Multi-Pathogen Assessment', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Purpose: ')
    run.bold = True
    para.add_run('Compare risks from multiple pathogens (Production Mode: norovirus only)')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Note: ')
    run.bold = True
    para.add_run(
        'In Production Mode (default), only norovirus is available. Uncheck Production Mode '
        'to access additional pathogens for research purposes.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Step-by-step procedure:')

    steps = [
        'Select "Multi-Pathogen Assessment" from dropdown',
        'Select pathogen(s) from multi-select box (norovirus in Production Mode)',
        'Choose exposure route',
        'Set common exposure parameters',
        'Click "Run Multi-Pathogen Assessment"',
        'Review comparative risk analysis across pathogens',
        'Download multi-pathogen results'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Numbered para - Level 1')

    doc.add_page_break()

    # 5. RESULTS INTERPRETATION
    print("Adding Results Interpretation section...")
    heading = doc.add_heading('Understanding Results', level=1)

    heading2 = doc.add_heading('Risk Metrics', level=2)

    metrics = [
        ('Mean Dose (organisms per exposure)',
         'Average number of viable pathogen organisms ingested or inhaled during a single exposure event'),
        ('Infection Probability (per exposure)',
         'Probability of infection from a single exposure event, calculated using Beta-Binomial model'),
        ('Annual Infection Risk',
         'Probability of at least one infection over a year: 1 - (1 - P_inf)^n, where n = exposure frequency'),
        ('Illness Probability',
         'Probability of symptomatic illness: P_infection × P(ill|infected). For norovirus: P(ill|infected) = 0.60'),
        ('95% Confidence Interval',
         'Uncertainty range from Monte Carlo simulation (10,000 iterations)')
    ]

    for metric, description in metrics:
        para = doc.add_paragraph()
        para.style = 'Body Text'
        run = para.add_run(f'{metric}: ')
        run.bold = True
        para.add_run(description)

    heading2 = doc.add_heading('WHO Guideline Compliance', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'The World Health Organization (WHO) recommends a health-based target of 10⁻⁴ (0.0001 or 1 in 10,000) '
        'annual infection risk for drinking water and recreational water quality.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Compliance evaluation:')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('PASS (Compliant): Annual infection risk < 10⁻⁴')
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('FAIL (Exceeds Guideline): Annual infection risk ≥ 10⁻⁴')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True

    heading2 = doc.add_heading('Exporting and Reporting Results', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('The application provides multiple export options:')

    export_options = [
        'CSV Downloads: All results tables can be downloaded as CSV files for further analysis in Excel or R',
        'Plot Downloads: All visualizations can be saved as high-resolution PNG images',
        'PDF Reports: Comprehensive PDF reports including all plots, tables, and metadata',
        'Session State: Results persist during browser session for iterative analysis'
    ]

    for option in export_options:
        para = doc.add_paragraph(option, style='Bullet Level 1')

    doc.add_page_break()

    # 6. TROUBLESHOOTING
    print("Adding Troubleshooting section...")
    heading = doc.add_heading('Troubleshooting and Common Issues', level=1)

    heading2 = doc.add_heading('Installation Issues', level=2)

    issues = [
        ('Python version error',
         'Ensure Python 3.8+ is installed. Check with: python --version'),
        ('Missing dependencies',
         'Run: pip install -r requirements.txt from application directory'),
        ('Browser doesn\'t open automatically',
         'Manually navigate to http://localhost:8501 in your browser'),
        ('Port 8501 already in use',
         'Stop other Streamlit apps or specify different port: streamlit run web_app.py --server.port 8502')
    ]

    for issue, solution in issues:
        para = doc.add_paragraph()
        para.style = 'Body Text'
        run = para.add_run(f'{issue}: ')
        run.bold = True
        para.add_run(solution)

    heading2 = doc.add_heading('Data Upload Issues', level=2)

    issues = [
        ('CSV file not recognized',
         'Ensure file is true CSV (comma-separated), not Excel or tab-delimited. Save as CSV in Excel'),
        ('Column name errors',
         'Check CSV has exact column names as specified in examples (case-sensitive)'),
        ('Numerical format errors',
         'Use standard notation: 1000000 or 1e6, not "1,000,000" or text'),
        ('Empty or NULL values',
         'Remove rows with missing data or replace with valid numerical values')
    ]

    for issue, solution in issues:
        para = doc.add_paragraph()
        para.style = 'Body Text'
        run = para.add_run(f'{issue}: ')
        run.bold = True
        para.add_run(solution)

    heading2 = doc.add_heading('Performance Issues', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'Large datasets (>1000 scenarios or time points) may cause slow performance. Recommendations:'
    )

    recommendations = [
        'Reduce Monte Carlo iterations (default 10,000) to 1,000 for preliminary analysis',
        'Split large temporal datasets into smaller time periods',
        'Ensure at least 4 GB RAM available',
        'Close other browser tabs and applications',
        'Use batch processing mode for very large scenario sets'
    ]

    for rec in recommendations:
        para = doc.add_paragraph(rec, style='Bullet Level 1')

    doc.add_page_break()

    # 7. LIMITATIONS AND ASSUMPTIONS
    print("Adding Limitations section...")
    heading = doc.add_heading('Limitations and Assumptions', level=1)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Model Assumptions:')
    run.bold = True

    assumptions = [
        'Pathogen concentrations are uniformly distributed in water (no spatial aggregation)',
        'Dose-response relationships are based on human challenge studies (limited data)',
        'Exposure events are independent (no immunity from previous infections)',
        'Monte Carlo assumes log-normal distribution for pathogen concentrations',
        'Treatment LRV is constant and does not vary with pathogen load',
        'No pathogen regrowth occurs post-treatment'
    ]

    for assumption in assumptions:
        para = doc.add_paragraph(assumption, style='Bullet Level 1')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Limitations:')
    run.bold = True

    limitations = [
        'Production Mode restricted to norovirus only (validated pathogen)',
        'Does not model viral inactivation kinetics (uses static LRV)',
        'Does not account for population immunity or vaccination',
        'Assumes adult population (dose-response may differ for children)',
        'Does not model secondary transmission (household spread)',
        'Requires user-provided pathogen concentration data (no predictive modeling)'
    ]

    for limitation in limitations:
        para = doc.add_paragraph(limitation, style='Bullet Level 1')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Regulatory Use: ')
    run.bold = True
    para.add_run(
        'Results from Production Mode (norovirus, Beta-Binomial model) are validated for regulatory '
        'submissions in New Zealand. Results from Research Mode (other pathogens/models) require '
        'independent verification before regulatory use.'
    )

    doc.add_page_break()

    # 8. REFERENCES
    print("Adding References section...")
    heading = doc.add_heading('References', level=1)

    references = [
        'Teunis, P.F.M., Moe, C.L., Liu, P., Miller, S.E., Lindesmith, L., Baric, R.S., Le Pendu, J., & Calderon, R.L. (2008). Norwalk virus: How infectious is it? Journal of Medical Virology, 80(8), 1468-1476.',

        'McBride, G.B. (2017). Bell Island Wastewater Treatment Plant Discharge QMRA. NIWA Client Report No: 2017350HN. Prepared for Wellington Water Limited.',

        'World Health Organization (2016). Quantitative Microbial Risk Assessment: Application for Water Safety Management. Geneva: WHO Press.',

        'Haas, C.N., Rose, J.B., & Gerba, C.P. (2014). Quantitative Microbial Risk Assessment (2nd ed.). John Wiley & Sons, Inc.',

        'Haas, C.N. (2002). Conditional dose-response relationships for microorganisms: Development and application. Risk Analysis, 22(3), 455-463.',

        'Van Abel, N., Schoen, M.E., Kissel, J.C., & Meschke, J.S. (2017). Comparison of risk predicted by multiple norovirus dose-response models and implications for quantitative microbial risk assessment. Risk Analysis, 37(2), 245-264.',

        'Ministry for the Environment & Ministry of Health (2003). Microbiological Water Quality Guidelines for Marine and Freshwater Recreational Areas. Wellington: Ministry for the Environment.',
    ]

    for i, ref in enumerate(references, 1):
        para = doc.add_paragraph()
        para.style = 'Body Text'
        para.add_run(f'[{i}] {ref}')
        para.paragraph_format.first_line_indent = Inches(-0.3)
        para.paragraph_format.left_indent = Inches(0.5)

    # Save the document
    doc.save(str(output_path))

    print(f"\n{'='*80}")
    print("IMPROVED NIWA-FORMATTED USER GUIDE CREATED SUCCESSFULLY")
    print(f"{'='*80}\n")
    print(f"Output file: {output_path.absolute()}")
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB\n")
    print("Document structure:")
    print("  - Cover Page with title and metadata")
    print("  - Table of Contents (placeholder)")
    print("  - Executive Summary (concise and structured)")
    print("  - Introduction to QMRA")
    print("  - Getting Started (System Requirements, Installation, Launching)")
    print("  - Beta-Binomial Model (CRITICAL SECTION with validation)")
    print("  - Assessment Workflows (ALL 5 workflows with detailed steps)")
    print("  - Results Interpretation (metrics, WHO compliance, exports)")
    print("  - Troubleshooting (installation, data, performance)")
    print("  - Limitations and Assumptions")
    print("  - Complete References")
    print(f"\n{'='*80}\n")

    return output_path


if __name__ == '__main__':
    create_niwa_user_guide()
