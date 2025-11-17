#!/usr/bin/env python3
"""
Create Final QMRA User Guide - Professional Documentation
==========================================================

Creates a comprehensive Word document without screenshots.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

def add_colored_heading(doc, text, level=1, color_rgb=(31, 119, 180)):
    """Add a colored heading."""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_bullet_with_bold(doc, bold_text, normal_text, style='List Bullet'):
    """Add a bullet point with bold prefix."""
    p = doc.add_paragraph(style=style)
    run1 = p.add_run(bold_text)
    run1.bold = True
    p.add_run(normal_text)
    return p

def add_code_block(doc, code_text, description=''):
    """Add a formatted code block."""
    if description:
        p = doc.add_paragraph(description)

    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.5)
    code_p.paragraph_format.space_before = Pt(6)
    code_p.paragraph_format.space_after = Pt(6)

    code_run = code_p.add_run(code_text)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add light gray background effect (approximation)
    return code_p

def create_user_guide():
    """Create comprehensive user guide."""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ====================
    # TITLE PAGE
    # ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    title_run = title.add_run('QMRA Batch Processing\nWeb Application')
    title_run.font.size = Pt(40)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 119, 180)

    doc.add_paragraph()
    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Professional User Guide\n&\nTechnical Documentation')
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(70, 70, 70)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Quantitative Microbial Risk Assessment Toolkit\n\n')
    footer_run.font.size = Pt(13)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

    footer_run2 = footer.add_run(f'Version 2.0 | {datetime.now().strftime("%B %Y")}\n\n')
    footer_run2.font.size = Pt(11)
    footer_run2.font.color.rgb = RGBColor(128, 128, 128)

    footer_run3 = footer.add_run('NIWA - National Institute of Water & Atmospheric Research\n')
    footer_run3.font.size = Pt(11)
    footer_run3.font.color.rgb = RGBColor(128, 128, 128)

    footer_run4 = footer.add_run('Hamilton, New Zealand')
    footer_run4.font.size = Pt(10)
    footer_run4.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # ====================
    # EXECUTIVE SUMMARY
    # ====================
    add_colored_heading(doc, 'Executive Summary', level=1)

    doc.add_paragraph(
        'The QMRA Batch Processing Web Application is a professional tool for quantitative '
        'microbial risk assessment (QMRA) of waterborne pathogens. It evaluates public health '
        'risks from pathogen exposure through shellfish consumption, recreational water contact, '
        'and other routes using validated dose-response models and Monte Carlo simulation.'
    )

    doc.add_paragraph()

    doc.add_paragraph(
        'This application implements Production Mode as the default configuration, restricting '
        'pathogen selection to norovirus only with the validated Beta-Binomial dose-response model. '
        'This ensures all risk assessments meet scientific standards and align with project '
        'deliverables.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Key Features:')
    run1.bold = True

    doc.add_paragraph('Five specialized assessment workflows for different scenarios', style='List Bullet')
    doc.add_paragraph('Validated Beta-Binomial dose-response model for norovirus (α=0.04, β=0.055)', style='List Bullet')
    doc.add_paragraph('Monte Carlo simulation with 10,000 iterations for uncertainty quantification', style='List Bullet')
    doc.add_paragraph('WHO guideline compliance checking (10⁻⁴ annual infection risk)', style='List Bullet')
    doc.add_paragraph('Interactive visualizations and statistical summaries', style='List Bullet')
    doc.add_paragraph('Professional PDF reports and CSV data exports', style='List Bullet')

    doc.add_page_break()

    # ====================
    # TABLE OF CONTENTS
    # ====================
    add_colored_heading(doc, 'Table of Contents', level=1)

    toc_items = [
        ('1', 'Introduction to QMRA'),
        ('', '1.1 What is Quantitative Microbial Risk Assessment?'),
        ('', '1.2 Application Purpose and Scope'),
        ('', '1.3 Target Users'),
        ('2', 'Getting Started'),
        ('', '2.1 System Requirements'),
        ('', '2.2 Launching the Application'),
        ('', '2.3 Understanding the Interface'),
        ('', '2.4 Production Mode vs Research Mode'),
        ('3', 'Assessment Workflows'),
        ('', '3.1 Batch Scenarios Assessment'),
        ('', '3.2 Spatial Risk Assessment'),
        ('', '3.3 Temporal Risk Assessment'),
        ('', '3.4 Treatment Comparison'),
        ('', '3.5 Multi-Pathogen Assessment'),
        ('4', 'Input Data Requirements'),
        ('', '4.1 Pathogen Concentration Data'),
        ('', '4.2 Dilution Factor Data'),
        ('', '4.3 Temporal Monitoring Data'),
        ('', '4.4 Data Formats and Templates'),
        ('5', 'Understanding Results'),
        ('', '5.1 Risk Metrics Explained'),
        ('', '5.2 Statistical Outputs'),
        ('', '5.3 Visualizations'),
        ('', '5.4 WHO Guideline Compliance'),
        ('6', 'Technical Background'),
        ('', '6.1 QMRA Methodology'),
        ('', '6.2 Dose-Response Models'),
        ('', '6.3 Beta-Binomial Model (Norovirus)'),
        ('', '6.4 Monte Carlo Simulation'),
        ('7', 'Troubleshooting and Support'),
        ('8', 'References'),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        if num:
            run1 = p.add_run(f'{num}. ')
            run1.bold = True
            run2 = p.add_run(title)
            run2.bold = True
        else:
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(title)

    doc.add_page_break()

    # ====================
    # 1. INTRODUCTION
    # ====================
    add_colored_heading(doc, '1. Introduction to QMRA', level=1)

    # 1.1
    add_colored_heading(doc, '1.1 What is Quantitative Microbial Risk Assessment?', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'Quantitative Microbial Risk Assessment (QMRA) is a scientific framework for estimating '
        'the probability of infection or illness from exposure to pathogenic microorganisms in '
        'water, food, or environmental media. Unlike traditional water quality monitoring that '
        'relies solely on indicator organisms (e.g., E. coli), QMRA directly estimates health '
        'risks by modeling the entire exposure pathway from source to receptor.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('The QMRA Framework integrates four key components:')
    run1.bold = True

    add_bullet_with_bold(doc, 'Hazard Identification: ', 'Identifying pathogenic microorganisms of concern (e.g., norovirus, Cryptosporidium)')
    add_bullet_with_bold(doc, 'Exposure Assessment: ', 'Quantifying pathogen concentrations and human exposure doses')
    add_bullet_with_bold(doc, 'Dose-Response Modeling: ', 'Relating exposure dose to probability of infection')
    add_bullet_with_bold(doc, 'Risk Characterization: ', 'Calculating annual infection/illness probabilities and comparing to guidelines')

    doc.add_paragraph()

    doc.add_paragraph(
        'QMRA outputs are typically expressed as annual infection risk (probability per person per year), '
        'which can be compared against health-based targets such as the World Health Organization (WHO) '
        'guideline of 10⁻⁴ (1 in 10,000) for drinking water and recreational water quality.'
    )

    # 1.2
    add_colored_heading(doc, '1.2 Application Purpose and Scope', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'This web-based application enables water quality professionals to conduct comprehensive '
        'QMRA studies without requiring specialized programming skills or advanced statistical software. '
        'The application handles complex calculations, uncertainty quantification, and report generation '
        'through an intuitive interface.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Primary Use Cases:')
    run1.bold = True

    doc.add_paragraph('Evaluating health risks from wastewater discharge to shellfish harvesting areas', style='List Bullet')
    doc.add_paragraph('Assessing recreational water quality at swimming beaches', style='List Bullet')
    doc.add_paragraph('Comparing treatment technology effectiveness for pathogen reduction', style='List Bullet')
    doc.add_paragraph('Analyzing temporal risk trends from long-term monitoring programs', style='List Bullet')
    doc.add_paragraph('Supporting regulatory compliance and water safety plan development', style='List Bullet')

    # 1.3
    add_colored_heading(doc, '1.3 Target Users', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph('This application is designed for:')

    doc.add_paragraph('Water quality managers and environmental engineers', style='List Bullet')
    doc.add_paragraph('Public health risk assessors and epidemiologists', style='List Bullet')
    doc.add_paragraph('Regulatory compliance officers and auditors', style='List Bullet')
    doc.add_paragraph('Wastewater treatment plant operators and planners', style='List Bullet')
    doc.add_paragraph('Environmental health researchers and consultants', style='List Bullet')
    doc.add_paragraph('Shellfish industry managers and aquaculture professionals', style='List Bullet')

    doc.add_page_break()

    # ====================
    # 2. GETTING STARTED
    # ====================
    add_colored_heading(doc, '2. Getting Started', level=1)

    # 2.1
    add_colored_heading(doc, '2.1 System Requirements', level=2, color_rgb=(44, 62, 80))

    p = doc.add_paragraph()
    run1 = p.add_run('Software Requirements:')
    run1.bold = True

    doc.add_paragraph('Python 3.8 or higher', style='List Bullet')
    doc.add_paragraph('Streamlit web framework', style='List Bullet')
    doc.add_paragraph('Modern web browser (Chrome, Firefox, Edge, Safari)', style='List Bullet')
    doc.add_paragraph('Required Python packages: numpy, pandas, scipy, matplotlib, plotly', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Hardware Requirements:')
    run1.bold = True

    doc.add_paragraph('Minimum 4GB RAM (8GB recommended for large datasets)', style='List Bullet')
    doc.add_paragraph('Modern CPU (Monte Carlo simulation is computationally intensive)', style='List Bullet')
    doc.add_paragraph('Screen resolution: 1920x1080 or higher recommended', style='List Bullet')

    # 2.2
    add_colored_heading(doc, '2.2 Launching the Application', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'The QMRA application runs as a local web server using Streamlit. Follow these steps to launch:'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Method 1: Using the Batch File (Windows)')
    run1.bold = True

    doc.add_paragraph('Navigate to: Batch_Processing_App/app/', style='List Number')
    doc.add_paragraph('Double-click: launch_web_gui.bat', style='List Number')
    doc.add_paragraph('Your default web browser will automatically open to http://localhost:8501', style='List Number')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Method 2: Using Command Line')
    run1.bold = True

    add_code_block(doc, 'cd Batch_Processing_App/app\nstreamlit run web_app.py')

    doc.add_paragraph()
    doc.add_paragraph(
        'The terminal will display "Local URL: http://localhost:8501" when ready. '
        'The application will remain running until you close the terminal window or press Ctrl+C.'
    )

    # 2.3
    add_colored_heading(doc, '2.3 Understanding the Interface', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph('The application interface consists of two main areas:')

    doc.add_paragraph()

    add_bullet_with_bold(doc, 'Left Sidebar: ',
        'Contains configuration controls, production mode toggle, and assessment type selection. '
        'This panel remains visible as you navigate through different workflows.')

    add_bullet_with_bold(doc, 'Main Panel: ',
        'Displays input forms, data preview tables, results, visualizations, and download options. '
        'Content changes based on the selected assessment mode.')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Navigation Pattern:')
    run1.bold = True

    doc.add_paragraph('Select assessment mode from sidebar dropdown', style='List Number')
    doc.add_paragraph('Configure input parameters and upload data (if required)', style='List Number')
    doc.add_paragraph('Click "Run Assessment" button', style='List Number')
    doc.add_paragraph('Review results, visualizations, and download outputs', style='List Number')

    # 2.4
    add_colored_heading(doc, '2.4 Production Mode vs Research Mode', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'The application operates in two distinct modes, controlled by a checkbox in the sidebar:'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Production Mode (Default - Recommended)')
    run1.bold = True
    run1.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()
    doc.add_paragraph('Production Mode is enabled by default and provides:')

    doc.add_paragraph('Pathogen selection restricted to norovirus only', style='List Bullet')
    doc.add_paragraph('Validated Beta-Binomial dose-response model (α=0.04, β=0.055)', style='List Bullet')
    doc.add_paragraph('Complete verification against reference Excel calculations', style='List Bullet')
    doc.add_paragraph('Alignment with project contract deliverables', style='List Bullet')
    doc.add_paragraph('Ensures scientifically defensible results for regulatory submissions', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('When to use Production Mode:')
    run1.bold = True

    doc.add_paragraph('Official risk assessments for regulatory compliance', style='List Bullet')
    doc.add_paragraph('Reports for public health authorities', style='List Bullet')
    doc.add_paragraph('Water safety plan documentation', style='List Bullet')
    doc.add_paragraph('Shellfish harvesting area classifications', style='List Bullet')
    doc.add_paragraph('Any scenario requiring WHO guideline compliance evaluation', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Research Mode')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 140, 0)

    doc.add_paragraph()
    doc.add_paragraph('Research Mode can be enabled by unchecking the Production Mode checkbox. This provides:')

    doc.add_paragraph('Access to all 6 pathogens (norovirus, Campylobacter, Cryptosporidium, E. coli, rotavirus, Salmonella)', style='List Bullet')
    doc.add_paragraph('Multiple dose-response model options', style='List Bullet')
    doc.add_paragraph('Exploratory analysis capabilities', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('IMPORTANT WARNING: ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run(
        'Non-norovirus pathogens in Research Mode have NOT been validated to the same standard '
        'as norovirus. These should be used for exploratory research purposes only, not for '
        'regulatory or compliance assessments.'
    )

    doc.add_page_break()

    # ====================
    # 3. ASSESSMENT WORKFLOWS
    # ====================
    add_colored_heading(doc, '3. Assessment Workflows', level=1)

    doc.add_paragraph(
        'The application provides five specialized assessment workflows, each designed for '
        'specific risk assessment scenarios. This section provides detailed instructions for '
        'each workflow.'
    )

    doc.add_page_break()

    # 3.1 Batch Scenarios
    add_colored_heading(doc, '3.1 Batch Scenarios Assessment', level=2, color_rgb=(44, 62, 80))

    p = doc.add_paragraph()
    run1 = p.add_run('Overview:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'The Batch Scenarios workflow allows you to evaluate multiple pre-configured scenarios '
        'simultaneously using a "library approach" with three CSV input files. This is the most '
        'comprehensive workflow, suitable for complex risk assessments with multiple combinations '
        'of pathogens, sites, and exposure conditions.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Use Cases:')
    run1.bold = True

    doc.add_paragraph('Master risk assessment covering multiple locations and scenarios', style='List Bullet')
    doc.add_paragraph('Comparing risks across different exposure routes (shellfish vs swimming)', style='List Bullet')
    doc.add_paragraph('Evaluating seasonal variations with different parameter sets', style='List Bullet')
    doc.add_paragraph('Batch processing for annual reporting requirements', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Input Files Required:')
    run1.bold = True

    doc.add_paragraph()

    add_bullet_with_bold(doc, 'pathogen_data.csv: ',
        'Contains pathogen concentration data with columns for pathogen_id, name, concentration '
        '(mean, stdev, distribution type), and source information.')

    add_bullet_with_bold(doc, 'scenarios.csv: ',
        'Defines scenario configurations including scenario_id, name, exposure_route, dilution_factor, '
        'treatment_lrv, exposure_frequency, exposure_volume, and population_size.')

    add_bullet_with_bold(doc, 'master_scenarios.csv: ',
        'Links pathogens to scenarios with master_scenario_id, pathogen_id, and scenario_id columns.')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step-by-Step Procedure:')
    run1.bold = True

    doc.add_paragraph('Select "Batch Scenarios" from the Assessment Mode dropdown in the sidebar', style='List Number')
    doc.add_paragraph('Choose one of two data input options:', style='List Number')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Check "Use example data" to load 15 pre-configured scenarios')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Upload your own CSV files using the file upload widgets')

    doc.add_paragraph('Review the data preview tables to verify correct data loading', style='List Number')
    doc.add_paragraph('Scroll down and click "Run Batch Assessment"', style='List Number')
    doc.add_paragraph('Wait for Monte Carlo simulation to complete (10,000 iterations per scenario)', style='List Number')
    doc.add_paragraph('Review results in the comprehensive summary table showing:', style='List Number')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Mean dose and infection probability per scenario')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Annual infection risk')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('WHO guideline compliance status')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('95% confidence intervals')

    doc.add_paragraph('Download results as CSV or complete assessment package as ZIP', style='List Number')

    doc.add_page_break()

    # 3.2 Spatial Assessment
    add_colored_heading(doc, '3.2 Spatial Risk Assessment', level=2, color_rgb=(44, 62, 80))

    p = doc.add_paragraph()
    run1 = p.add_run('Overview:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'The Spatial Assessment workflow evaluates risk across multiple geographic sites with '
        'different dilution factors. This is particularly useful for mapping risk gradients near '
        'wastewater discharge points or comparing shellfish harvesting zones.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Use Cases:')
    run1.bold = True

    doc.add_paragraph('Shellfish harvesting area classification (approved, conditionally approved, restricted)', style='List Bullet')
    doc.add_paragraph('Beach water quality assessment at multiple sampling locations', style='List Bullet')
    doc.add_paragraph('Wastewater discharge impact assessment with distance from outfall', style='List Bullet')
    doc.add_paragraph('Comparing risk at upstream vs downstream locations', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Input Requirements:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Dilution data CSV with two columns: Site (location name) and Dilution_Factor (dimensionless). '
        'Example data includes 6 sites with dilution factors ranging from 100 (near-field) to 100,000 (far-field).'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step-by-Step Procedure:')
    run1.bold = True

    doc.add_paragraph('Select "Spatial Assessment" from the Assessment Mode dropdown', style='List Number')
    doc.add_paragraph('Configure parameters:', style='List Number')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    run1 = p.add_run('Pathogen: ')
    run1.bold = True
    p.add_run('Select norovirus (Production Mode enforces this)')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    run1 = p.add_run('Exposure Route: ')
    run1.bold = True
    p.add_run('Choose Shellfish Consumption or Primary Contact (Swimming)')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    run1 = p.add_run('Pathogen Concentration: ')
    run1.bold = True
    p.add_run('Enter mean concentration in organisms/L (e.g., 1e6 for raw wastewater)')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    run1 = p.add_run('Treatment LRV: ')
    run1.bold = True
    p.add_run('Log reduction value (0-6, typically 2-4 for secondary treatment)')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    run1 = p.add_run('Exposure Frequency: ')
    run1.bold = True
    p.add_run('Events per year (e.g., 12 for monthly shellfish consumption)')

    doc.add_paragraph('Upload dilution data CSV or use example data (6 sites)', style='List Number')
    doc.add_paragraph('Click "Run Spatial Assessment"', style='List Number')
    doc.add_paragraph('Review results:', style='List Number')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Site-by-site risk table with compliance status')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Risk gradient visualization showing spatial trends')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Recommended classification for each site')

    doc.add_paragraph('Download results and visualizations', style='List Number')

    doc.add_page_break()

    # 3.3 Temporal Assessment
    add_colored_heading(doc, '3.3 Temporal Risk Assessment', level=2, color_rgb=(44, 62, 80))

    p = doc.add_paragraph()
    run1 = p.add_run('Overview:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'The Temporal Assessment workflow analyzes risk trends over time using pathogen '
        'monitoring data. This is essential for understanding seasonal patterns, evaluating '
        'treatment performance over time, and identifying risk peaks.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Use Cases:')
    run1.bold = True

    doc.add_paragraph('Seasonal risk pattern analysis (summer vs winter)', style='List Bullet')
    doc.add_paragraph('Before/after treatment upgrade evaluation', style='List Bullet')
    doc.add_paragraph('Long-term monitoring program assessment', style='List Bullet')
    doc.add_paragraph('Identifying peak risk periods for management interventions', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Input Requirements:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Temporal data CSV with two columns: Date (YYYY-MM-DD format) and Concentration (organisms/L). '
        'The application handles missing data and irregular sampling intervals.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step-by-Step Procedure:')
    run1.bold = True

    doc.add_paragraph('Select "Temporal Assessment" from dropdown', style='List Number')
    doc.add_paragraph('Configure pathogen and exposure route', style='List Number')
    doc.add_paragraph('Upload temporal monitoring data CSV', style='List Number')
    doc.add_paragraph('Set assessment parameters (treatment LRV, dilution, exposure frequency)', style='List Number')
    doc.add_paragraph('Click "Run Temporal Assessment"', style='List Number')
    doc.add_paragraph('Review time-series visualization showing:', style='List Number')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Daily/monthly risk estimates over time')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('WHO guideline reference line')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Trend analysis and seasonal patterns')

    doc.add_paragraph('Examine summary statistics (mean, median, percentiles over time period)', style='List Number')
    doc.add_paragraph('Download results and plots', style='List Number')

    doc.add_page_break()

    # 3.4 Treatment Comparison
    add_colored_heading(doc, '3.4 Treatment Comparison', level=2, color_rgb=(44, 62, 80))

    p = doc.add_paragraph()
    run1 = p.add_run('Overview:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'The Treatment Comparison workflow enables side-by-side evaluation of multiple treatment '
        'technologies or log reduction values (LRVs). This supports decision-making for treatment '
        'upgrades or optimization.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Use Cases:')
    run1.bold = True

    doc.add_paragraph('Comparing UV disinfection vs chlorination effectiveness', style='List Bullet')
    doc.add_paragraph('Evaluating membrane filtration vs conventional treatment', style='List Bullet')
    doc.add_paragraph('Assessing incremental risk reduction from treatment upgrades', style='List Bullet')
    doc.add_paragraph('Cost-benefit analysis of different LRV targets', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step-by-Step Procedure:')
    run1.bold = True

    doc.add_paragraph('Select "Treatment Comparison" from dropdown', style='List Number')
    doc.add_paragraph('Configure base parameters (pathogen, exposure route)', style='List Number')
    doc.add_paragraph('Define 2-4 treatment scenarios by specifying:', style='List Number')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Scenario Name (e.g., "Secondary Treatment", "Secondary + UV")')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Log Reduction Value (LRV) for each scenario')

    doc.add_paragraph('Set shared parameters (concentration, dilution, exposure frequency)', style='List Number')
    doc.add_paragraph('Click "Run Treatment Comparison"', style='List Number')
    doc.add_paragraph('Review comparative results table showing:', style='List Number')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Risk estimates for each treatment scenario')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Compliance status (pass/fail)')

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p.add_run('Risk reduction percentages')

    doc.add_paragraph('Examine side-by-side bar chart comparison', style='List Number')
    doc.add_paragraph('Download results', style='List Number')

    doc.add_page_break()

    # 3.5 Multi-Pathogen Assessment
    add_colored_heading(doc, '3.5 Multi-Pathogen Assessment', level=2, color_rgb=(44, 62, 80))

    p = doc.add_paragraph()
    run1 = p.add_run('Overview:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'The Multi-Pathogen Assessment workflow evaluates multiple pathogens simultaneously to '
        'identify the dominant health risk driver. This is important because different pathogens '
        'have vastly different infectivity and dose-response characteristics.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('IMPORTANT NOTE: ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run(
        'In Production Mode, this workflow is restricted to norovirus only. To assess multiple '
        'pathogens, you must disable Production Mode and enable Research Mode. However, non-norovirus '
        'pathogens have NOT been validated and should only be used for exploratory purposes.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step-by-Step Procedure:')
    run1.bold = True

    doc.add_paragraph('Select "Multi-Pathogen Assessment" from dropdown', style='List Number')
    doc.add_paragraph('In Production Mode: Only norovirus will be available', style='List Number')
    doc.add_paragraph('In Research Mode: Select 2-6 pathogens from multi-select box', style='List Number')
    doc.add_paragraph('Configure exposure route and parameters', style='List Number')
    doc.add_paragraph('Click "Run Multi-Pathogen Assessment"', style='List Number')
    doc.add_paragraph('Review pathogen comparison table showing relative risks', style='List Number')
    doc.add_paragraph('Identify the dominant pathogen driving overall risk', style='List Number')

    doc.add_page_break()

    # Continue with remaining sections...
    # (I'll add the rest in the next part due to length)

    # ====================
    # 4. INPUT DATA REQUIREMENTS
    # ====================
    add_colored_heading(doc, '4. Input Data Requirements', level=1)

    add_colored_heading(doc, '4.1 Pathogen Concentration Data', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'Pathogen concentration is the starting point for all QMRA calculations. Concentration should be '
        'measured or estimated at the contamination source (e.g., raw wastewater, secondary effluent) before '
        'dilution and further fate/transport processes.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Typical Concentration Ranges for Norovirus:')
    run1.bold = True

    doc.add_paragraph('Raw wastewater: 10⁶ - 10⁷ gene copies/L', style='List Bullet')
    doc.add_paragraph('Primary effluent: 10⁵ - 10⁶ gene copies/L', style='List Bullet')
    doc.add_paragraph('Secondary effluent (activated sludge): 10³ - 10⁴ gene copies/L', style='List Bullet')
    doc.add_paragraph('UV-disinfected effluent: 10² - 10³ gene copies/L', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: These are gene copies, not infectious particles. A conversion factor (typically 1:100 to 1:1000) '
        'is often applied to estimate infectious virions from qPCR measurements.'
    )

    add_colored_heading(doc, '4.2 Dilution Factor Data', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'Dilution factor represents the volumetric dilution of wastewater in receiving waters. This can be '
        'estimated using:'
    )

    doc.add_paragraph('Wastewater discharge flow rate and receiving water flow rate', style='List Bullet')
    doc.add_paragraph('Hydrodynamic modeling (e.g., MIKE 3, Delft3D)', style='List Bullet')
    doc.add_paragraph('Tracer studies or dye dilution experiments', style='List Bullet')
    doc.add_paragraph('Conservative chemical ratios (e.g., salinity gradients)', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Example Dilution Factors:')
    run1.bold = True

    doc.add_paragraph('Near-field (< 100m from outfall): 10-100', style='List Bullet')
    doc.add_paragraph('Mid-field (100-500m): 100-1,000', style='List Bullet')
    doc.add_paragraph('Far-field (> 500m): 1,000-100,000', style='List Bullet')

    add_colored_heading(doc, '4.3 Temporal Monitoring Data', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'For temporal assessments, provide time-series pathogen concentration measurements in CSV format:'
    )

    doc.add_paragraph()
    add_code_block(doc, 'Date,Concentration\n2024-01-15,5000\n2024-02-12,8000\n2024-03-10,3000')

    doc.add_paragraph()
    doc.add_paragraph(
        'The application handles irregular sampling intervals and can interpolate missing data points. '
        'Minimum recommended: monthly sampling for one year to capture seasonal patterns.'
    )

    add_colored_heading(doc, '4.4 Data Formats and Templates', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'All CSV files should use comma delimiters with headers in the first row. The application provides '
        'example data files that can be used as templates:'
    )

    doc.add_paragraph('pathogen_data.csv - Example pathogen concentration data', style='List Bullet')
    doc.add_paragraph('scenarios.csv - Example scenario configurations', style='List Bullet')
    doc.add_paragraph('dilution_example.csv - Example spatial dilution data', style='List Bullet')
    doc.add_paragraph('temporal_example.csv - Example temporal monitoring data', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'These templates can be found in the Batch_Processing_App/example_data/ directory.'
    )

    doc.add_page_break()

    # ====================
    # 5. UNDERSTANDING RESULTS
    # ====================
    add_colored_heading(doc, '5. Understanding Results', level=1)

    add_colored_heading(doc, '5.1 Risk Metrics Explained', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph('The application reports several key risk metrics for each assessment:')

    doc.add_paragraph()

    add_bullet_with_bold(doc, 'Mean Dose (organisms per exposure): ',
        'Average number of viable pathogen organisms ingested or inhaled during a single exposure event. '
        'Calculated as: (Concentration × Volume) / (Dilution × 10^LRV)')

    add_bullet_with_bold(doc, 'Infection Probability (per exposure): ',
        'Probability of infection from a single exposure event, calculated using the dose-response model. '
        'Range: 0 to 1 (0% to 100%). For norovirus with Beta-Binomial model: P_inf = f(dose, α=0.04, β=0.055)')

    add_bullet_with_bold(doc, 'Annual Infection Risk: ',
        'Probability of at least one infection over a year of repeated exposures. '
        'Formula: 1 - (1 - P_inf)^n, where n = exposure frequency per year')

    add_bullet_with_bold(doc, 'Illness Probability: ',
        'Probability of symptomatic illness given infection. '
        'Formula: P_illness = P_infection × P(ill|infected). '
        'For norovirus: P(ill|infected) = 0.60 (60% of infected become symptomatic)')

    add_bullet_with_bold(doc, 'Annual Illness Risk: ',
        'Probability of at least one illness episode per year. '
        'Calculated similarly to annual infection risk but using illness probability')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Uncertainty Quantification:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'All risk metrics include 95% confidence intervals representing uncertainty in:'
    )

    doc.add_paragraph('Pathogen concentration variability', style='List Bullet')
    doc.add_paragraph('Dilution factor uncertainty', style='List Bullet')
    doc.add_paragraph('Exposure volume and frequency distributions', style='List Bullet')
    doc.add_paragraph('Dose-response model parameter uncertainty', style='List Bullet')

    add_colored_heading(doc, '5.2 Statistical Outputs', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph('For each risk metric, the application provides:')

    add_bullet_with_bold(doc, 'Mean: ', 'Average value across 10,000 Monte Carlo iterations')
    add_bullet_with_bold(doc, 'Median (50th percentile): ', 'Middle value when sorted; less sensitive to outliers than mean')
    add_bullet_with_bold(doc, '5th percentile: ', 'Lower bound representing best-case scenario')
    add_bullet_with_bold(doc, '95th percentile: ', 'Upper bound representing worst-case scenario (conservative estimate)')
    add_bullet_with_bold(doc, 'Standard deviation: ', 'Measure of variability/uncertainty')

    add_colored_heading(doc, '5.3 Visualizations', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph('The application generates several types of visualizations:')

    doc.add_paragraph()

    add_bullet_with_bold(doc, 'Risk Overview Bar Charts: ',
        'Compare mean annual infection risks across scenarios, sites, or time periods. '
        'Bars are color-coded: green (compliant), red (exceeds guideline)')

    add_bullet_with_bold(doc, 'Uncertainty Distribution Histograms: ',
        'Show the full probability distribution of simulated outcomes (10,000 values). '
        'Useful for understanding risk variability and identifying tail risks')

    add_bullet_with_bold(doc, 'Time Series Line Plots: ',
        'Display temporal trends in risk with WHO guideline reference line (10⁻⁴). '
        'Includes confidence bands showing uncertainty over time')

    add_bullet_with_bold(doc, 'Spatial Gradient Plots: ',
        'Visualize how risk changes with distance from source (dilution factor). '
        'Helps identify safe zones and high-risk areas')

    add_colored_heading(doc, '5.4 WHO Guideline Compliance', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'The World Health Organization (WHO) recommends a health-based target of 10⁻⁴ (0.0001 or 1 in 10,000) '
        'annual infection risk for:'
    )

    doc.add_paragraph('Drinking water (waterborne pathogen risk)', style='List Bullet')
    doc.add_paragraph('Recreational water quality (swimming, surfing)', style='List Bullet')
    doc.add_paragraph('Wastewater reuse applications', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Compliance Evaluation:')
    run1.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('PASS (Compliant): ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run('Annual infection risk < 10⁻⁴')

    p = doc.add_paragraph()
    run1 = p.add_run('FAIL (Exceeds Guideline): ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('Annual infection risk ≥ 10⁻⁴')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Important Notes:')
    run1.bold = True

    doc.add_paragraph('The 10⁻⁴ target is a guideline, not a regulatory standard in all jurisdictions', style='List Bullet')
    doc.add_paragraph('Some countries/regions use different targets (e.g., 10⁻⁵ for stringent protection)', style='List Bullet')
    doc.add_paragraph('Compliance should be evaluated using mean or median risk, not worst-case (95th percentile)', style='List Bullet')
    doc.add_paragraph('Consult local health authorities for applicable standards', style='List Bullet')

    doc.add_page_break()

    # ====================
    # 6. TECHNICAL BACKGROUND
    # ====================
    add_colored_heading(doc, '6. Technical Background', level=1)

    add_colored_heading(doc, '6.1 QMRA Methodology', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'Quantitative Microbial Risk Assessment follows a structured four-step framework established '
        'by the National Research Council (1983) and refined by WHO (2016):'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step 1: Hazard Identification')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Identify pathogens of concern based on occurrence data, environmental persistence, and '
        'infectivity. For wastewater-associated risks, key pathogens include norovirus (most infectious), '
        'Cryptosporidium, Campylobacter, and enteric viruses.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step 2: Exposure Assessment')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Quantify the dose (number of organisms) ingested or inhaled per exposure event. '
        'This involves modeling:'
    )

    doc.add_paragraph('Pathogen concentration in source water (C, organisms/L)', style='List Bullet')
    doc.add_paragraph('Treatment reduction (LRV, log₁₀ reduction)', style='List Bullet')
    doc.add_paragraph('Environmental dilution (D, dimensionless)', style='List Bullet')
    doc.add_paragraph('Exposure volume (V, L per event)', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Dose formula: Dose = (C × V) / (D × 10^LRV)')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step 3: Dose-Response Assessment')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Apply mathematical models relating dose to infection probability. Models are derived from '
        'human feeding studies, outbreak data, or animal experiments with host-adjustment factors.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Step 4: Risk Characterization')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'Calculate per-exposure and annual risks, accounting for repeated exposures and uncertainty. '
        'Compare results against health-based targets and guidelines.'
    )

    add_colored_heading(doc, '6.2 Dose-Response Models', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'Dose-response models are mathematical functions describing the relationship between '
        'pathogen dose and probability of infection. The choice of model depends on:'
    )

    doc.add_paragraph('Pathogen characteristics (virus vs bacteria vs protozoa)', style='List Bullet')
    doc.add_paragraph('Infectivity (highly infectious organisms require exact models)', style='List Bullet')
    doc.add_paragraph('Available dose-response data (human studies, outbreak investigations)', style='List Bullet')
    doc.add_paragraph('Mathematical validity conditions (approximations have constraints)', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Common Dose-Response Models:')
    run1.bold = True

    doc.add_paragraph()

    add_bullet_with_bold(doc, 'Exponential Model: ',
        'P(inf) = 1 - exp(-r × dose). Simple single-parameter model assuming constant infection '
        'probability per organism. Used for highly infectious pathogens.')

    add_bullet_with_bold(doc, 'Beta-Poisson Model: ',
        'P(inf) = 1 - (1 + dose/β)^(-α). Two-parameter model accounting for host variability. '
        'VALID only when β >> 1 (beta much greater than 1).')

    add_bullet_with_bold(doc, 'Beta-Binomial Model (Exact): ',
        'P(inf) = 1 - B(α, β+dose) / B(α, β). Exact conditional model for scenarios where individual '
        'doses are known. Required for highly infectious viruses when β < 1.')

    add_colored_heading(doc, '6.3 Beta-Binomial Model (Norovirus)', level=2, color_rgb=(44, 62, 80))

    p = doc.add_paragraph()
    run1 = p.add_run('THE CRITICAL DISTINCTION FOR NOROVIRUS')
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph()

    doc.add_paragraph(
        'This application uses the EXACT Beta-Binomial dose-response model for norovirus, which is the '
        'mathematically correct formulation. This is NOT the same as the Beta-Poisson approximation.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Beta-Binomial Formula (Used by This Application):')
    run1.bold = True

    doc.add_paragraph()
    add_code_block(doc,
        'P(infection) = 1 - exp[\n'
        '    ln Γ(β + dose) + \n'
        '    ln Γ(α + β) - \n'
        '    ln Γ(α + β + dose) - \n'
        '    ln Γ(β)\n'
        ']'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Parameters (Teunis et al. 2008):')
    run1.bold = True

    doc.add_paragraph('α (alpha) = 0.04', style='List Bullet')
    doc.add_paragraph('β (beta) = 0.055', style='List Bullet')
    doc.add_paragraph('Source: "Norwalk virus: How infectious is it?" J. Med. Virol. 80:1468-1476', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Why Beta-Poisson is INVALID for Norovirus:')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph()

    doc.add_paragraph(
        'The Beta-Poisson approximation requires β >> 1 (beta much greater than 1) to be mathematically '
        'valid. For norovirus, β = 0.055 << 1, making this approximation inappropriate.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Consequences of Using Beta-Poisson (Incorrect):')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph('Underestimates risk by 2-4× at low doses (1-100 virions)', style='List Bullet')
    doc.add_paragraph('Produces scientifically indefensible results', style='List Bullet')
    doc.add_paragraph('Does not match reference calculations (David Wood, NIWA)', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Example Comparison (1 virion dose):')
    run1.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Beta-Binomial (CORRECT): ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run('42.1% infection probability')

    p = doc.add_paragraph(style='List Bullet')
    run1 = p.add_run('Beta-Poisson (WRONG): ')
    run1.bold = True
    run1.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('11.1% infection probability (3.8× underestimation)')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Validation:')
    run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'This application\'s Beta-Binomial implementation has been validated against reference Excel '
        'spreadsheets (David Wood, NIWA) and produces exact agreement for all test doses.'
    )

    add_colored_heading(doc, '6.4 Monte Carlo Simulation', level=2, color_rgb=(44, 62, 80))

    doc.add_paragraph(
        'The application uses Monte Carlo simulation to quantify uncertainty in risk estimates. This '
        'involves:'
    )

    doc.add_paragraph()

    doc.add_paragraph('Defining probability distributions for uncertain input parameters', style='List Number')
    doc.add_paragraph('Randomly sampling from these distributions for each iteration', style='List Number')
    doc.add_paragraph('Calculating risk for each sampled parameter set', style='List Number')
    doc.add_paragraph('Repeating 10,000 times to build output distribution', style='List Number')
    doc.add_paragraph('Summarizing results with mean, median, percentiles', style='List Number')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('Sources of Uncertainty Addressed:')
    run1.bold = True

    doc.add_paragraph('Pathogen concentration variability (log-normal distribution)', style='List Bullet')
    doc.add_paragraph('Dilution factor uncertainty (log-normal or uniform distribution)', style='List Bullet')
    doc.add_paragraph('Exposure volume and frequency (triangular or log-normal)', style='List Bullet')
    doc.add_paragraph('Treatment performance variation (normal or log-normal for LRV)', style='List Bullet')

    doc.add_paragraph()

    doc.add_paragraph(
        'Monte Carlo simulation provides a complete picture of risk uncertainty, enabling decision-makers '
        'to understand both typical risks (mean/median) and worst-case scenarios (95th percentile).'
    )

    doc.add_page_break()

    # ====================
    # 7. TROUBLESHOOTING
    # ====================
    add_colored_heading(doc, '7. Troubleshooting and Support', level=1)

    p = doc.add_paragraph()
    run1 = p.add_run('Common Issues and Solutions:')
    run1.bold = True

    doc.add_paragraph()

    add_bullet_with_bold(doc, 'Application won\'t start: ',
        'Verify Python and Streamlit are installed correctly. Check firewall settings aren\'t blocking port 8501.')

    add_bullet_with_bold(doc, 'File upload fails: ',
        'Ensure CSV files use comma delimiters and have correct column headers. Check for special characters in data.')

    add_bullet_with_bold(doc, 'Assessment runs slowly: ',
        'Monte Carlo with 10,000 iterations is computationally intensive. Reduce iteration count for testing (not recommended for final results).')

    add_bullet_with_bold(doc, 'Results seem unrealistic: ',
        'Verify input parameters (concentration, dilution, LRV) are in correct units and reasonable ranges.')

    add_bullet_with_bold(doc, 'Plots not displaying: ',
        'Update web browser to latest version. Clear browser cache. Try different browser.')

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('For Additional Support:')
    run1.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('NIWA - National Institute of Water & Atmospheric Research\n')
    run1.bold = True
    p.add_run('Hamilton, New Zealand\n')
    p.add_run('Website: www.niwa.co.nz\n')
    p.add_run('Email: info@niwa.co.nz')

    doc.add_page_break()

    # ====================
    # 8. REFERENCES
    # ====================
    add_colored_heading(doc, '8. References', level=1)

    references = [
        'Teunis, P.F.M., Moe, C.L., Liu, P., Miller, S.E., Lindesmith, L., Baric, R.S., Le Pendu, J., & Calderon, R.L. (2008). Norwalk virus: How infectious is it? Journal of Medical Virology, 80(8), 1468-1476. https://doi.org/10.1002/jmv.21237',

        'McBride, G.B. (2017). Bell Island Wastewater Treatment Plant Discharge QMRA. NIWA Client Report No: 2017350HN. Prepared for Wellington Water Limited.',

        'World Health Organization (2016). Quantitative Microbial Risk Assessment: Application for Water Safety Management. Geneva: WHO Press. ISBN: 978-92-4-156537-0',

        'Haas, C.N., Rose, J.B., & Gerba, C.P. (2014). Quantitative Microbial Risk Assessment (2nd ed.). John Wiley & Sons, Inc. ISBN: 978-1-118-14529-6',

        'National Research Council (1983). Risk Assessment in the Federal Government: Managing the Process. Washington, DC: The National Academies Press. https://doi.org/10.17226/366',

        'Haas, C.N. (2002). Conditional dose-response relationships for microorganisms: Development and application. Risk Analysis, 22(3), 455-463.',

        'Van Abel, N., Schoen, M.E., Kissel, J.C., & Meschke, J.S. (2017). Comparison of risk predicted by multiple norovirus dose-response models and implications for quantitative microbial risk assessment. Risk Analysis, 37(2), 245-264.',

        'WHO (2003). Guidelines for Safe Recreational Water Environments, Volume 1: Coastal and Fresh Waters. World Health Organization, Geneva.',

        'WHO (2011). Guidelines for Drinking-water Quality (4th ed.). World Health Organization, Geneva.',

        'Teunis, P.F.M., Havelaar, A.H., & Medema, G.J. (1996). Assessment of the risk of infection by Cryptosporidium or Giardia in drinking water from a surface water source. Water Research, 30(6), 1333-1346.',
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph(f'{i}. {ref}')
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ====================
    # APPENDIX / FINAL PAGE
    # ====================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    final_p = doc.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final_p.add_run('END OF USER GUIDE')
    final_run.font.size = Pt(14)
    final_run.font.bold = True
    final_run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()
    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        'For the latest version of this document and application updates,\n'
        'contact NIWA or visit the project repository.'
    )
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = Path('../../QMRA_User_Guide_Professional.docx')
    doc.save(str(output_path))

    print(f"\n{'='*80}")
    print("USER GUIDE CREATED SUCCESSFULLY")
    print(f"{'='*80}\n")
    print(f"File: {output_path.absolute()}")
    print(f"Size: {output_path.stat().st_size / 1024:.1f} KB")
    print(f"Pages: ~45-50 pages")
    print(f"\nDocument includes:")
    print("  - Professional title page")
    print("  - Executive summary")
    print("  - Complete table of contents")
    print("  - Comprehensive introduction to QMRA")
    print("  - Detailed workflow instructions (all 5 modes)")
    print("  - Input data requirements and formats")
    print("  - Results interpretation guide")
    print("  - Technical background (Beta-Binomial model)")
    print("  - Troubleshooting section")
    print("  - Complete references")
    print(f"\n{'='*80}\n")

    return output_path


if __name__ == '__main__':
    create_user_guide()
