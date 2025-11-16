#!/usr/bin/env python3
"""Create professional MS Word document for QMRA application overview."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('QMRA Batch Processing Web Application')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 119, 180)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Professional Overview')
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Footer info
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('Version 1.2.0 | November 2025\nNIWA Earth Sciences New Zealand')
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# Executive Summary
heading = doc.add_heading('Executive Summary', level=1)
heading.runs[0].font.color.rgb = RGBColor(31, 119, 180)

doc.add_paragraph(
    'The QMRA Batch Processing Web Application is a standalone tool for quantitative '
    'microbial risk assessment (QMRA). It evaluates health risks from pathogenic contamination '
    'in water using Monte Carlo simulation, providing evidence-based risk estimates for decision-making.'
)

doc.add_paragraph(
    'Key capabilities include multi-site analysis, temporal risk trends, treatment comparisons, '
    'and exposure-specific modeling for shellfish and swimming scenarios, with full uncertainty quantification.'
)

doc.add_page_break()

# What is QMRA
heading1 = doc.add_heading('1. What is QMRA?', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

heading2 = doc.add_heading('Definition', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)
doc.add_paragraph(
    'Quantitative Microbial Risk Assessment (QMRA) is a scientific framework for estimating '
    'the probability of infection or illness from exposure to pathogenic microorganisms (bacteria, '
    'viruses, protozoa) in water or food.'
)

heading2 = doc.add_heading('Core Concept', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)
doc.add_paragraph('QMRA combines four key components:')
doc.add_paragraph('Pathogen concentration in source water', style='List Bullet')
doc.add_paragraph('Treatment effectiveness (log reduction)', style='List Bullet')
doc.add_paragraph('Environmental dilution or bioaccumulation', style='List Bullet')
doc.add_paragraph('Dose-response relationship (dose to infection/illness)', style='List Bullet')

heading2 = doc.add_heading('Risk Calculation', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)
doc.add_paragraph(
    'QMRA produces annual infection and illness probabilities per person, which support '
    'compliance decisions against WHO guidelines (acceptable risk < 1 in 10,000 annually).'
)

heading2 = doc.add_heading('Why QMRA Matters', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)
doc.add_paragraph(
    'Traditional methods rely on indicator organisms or single-point assessments. QMRA accounts '
    'for scientific uncertainty and variability, enabling risk-based water safety management decisions.'
)

doc.add_page_break()

# Application Overview
heading1 = doc.add_heading('2. Application Overview', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

heading2 = doc.add_heading('Purpose', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)
doc.add_paragraph(
    'This web-based application enables users to run multiple QMRA scenarios simultaneously, '
    'visualize results, and generate professional reports—without requiring specialized programming knowledge.'
)

heading2 = doc.add_heading('Target Users', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)
doc.add_paragraph('Water utility managers', style='List Bullet')
doc.add_paragraph('Environmental health professionals', style='List Bullet')
doc.add_paragraph('Regulatory compliance officers', style='List Bullet')
doc.add_paragraph('Public health researchers', style='List Bullet')
doc.add_paragraph('Wastewater treatment engineers', style='List Bullet')

heading2 = doc.add_heading('Five Assessment Modes', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

modes = [
    ('Batch Scenarios', 'Run 15+ pre-configured scenarios to evaluate multiple conditions'),
    ('Spatial Assessment', 'Map risk across multiple sites with different dilution factors'),
    ('Temporal Assessment', 'Analyze seasonal or time-series risk trends from monitoring data'),
    ('Treatment Comparison', 'Compare multiple treatment technologies and their effectiveness'),
    ('Multi-Pathogen Assessment', 'Simultaneously evaluate 2–6 pathogens to identify dominant risks')
]

for mode, desc in modes:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(f'{mode}: ')
    p_run.bold = True
    p.add_run(desc)

doc.add_page_break()

# How It Works
heading1 = doc.add_heading('3. How It Works', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

heading2 = doc.add_heading('Step-by-Step Process', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

steps = [
    ('Input Data', 'Upload pathogen concentrations, dilution factors, and scenario parameters (or use examples)'),
    ('Configure Parameters', 'Select exposure routes, treatment LRV, population, and exposure frequency'),
    ('Monte Carlo Simulation', 'Run 10,000 iterations sampling from distributions for all uncertainties'),
    ('Risk Calculation', 'Calculate per-exposure and annual infection/illness risks'),
    ('Results & Visualization', 'View interactive charts, summary statistics, and compliance status'),
    ('Download', 'Export results as CSV, PDF, or complete ZIP package')
]

for i, (step, desc) in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p_run = p.add_run(f'{step}: ')
    p_run.bold = True
    p.add_run(desc)

heading2 = doc.add_heading('Exposure Routes (v1.2.0)', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

exposures = {
    'Shellfish Consumption': [
        'Meal size: 5–800 grams (truncated log-logistic distribution)',
        'Bioaccumulation Factor: 1–100× (shellfish concentrate pathogens)',
        'Total exposure: meal size × BAF'
    ],
    'Swimming/Primary Contact': [
        'Water ingestion rate: 5–200 mL/hour (truncated lognormal)',
        'Duration: 0.2–4 hours (triangular distribution)',
        'Total exposure: ingestion rate × duration'
    ],
    'Contaminated Water (Generic)': [
        'Fixed volume or range-based exposure',
        'Customizable via Method Harmonisation Factor'
    ]
}

for route, details in exposures.items():
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(f'{route}: ')
    p_run.bold = True
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet 2')

doc.add_page_break()

# Results Generated
heading1 = doc.add_heading('4. Results Generated', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

heading2 = doc.add_heading('Per-Scenario Output', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

results = [
    'Infection Risk (median, 5th, 95th percentile)',
    'Illness Risk (accounting for symptom development)',
    'Annual Risk (per person, accounting for exposure frequency)',
    'Population Impact (expected illness cases per year)',
    'WHO Compliance Status (COMPLIANT if < 1e-4; NON-COMPLIANT if >= 1e-4)',
    'Uncertainty Bounds (confidence intervals)'
]

for result in results:
    doc.add_paragraph(result, style='List Bullet')

heading2 = doc.add_heading('Visualizations', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

viz = [
    'Risk Overview (horizontal bar chart by scenario)',
    'Compliance Distribution (pie chart: COMPLIANT vs. NON-COMPLIANT)',
    'Risk Distribution (histogram of per-exposure probabilities)',
    'Population Impact (case counts per year)'
]

for v in viz:
    doc.add_paragraph(v, style='List Bullet')

heading2 = doc.add_heading('Download Formats', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph('Complete results CSV (all scenarios)', style='List Bullet')
doc.add_paragraph('Individual scenario details (CSV/Excel)', style='List Bullet')
doc.add_paragraph('High-resolution plots (PNG, 300 DPI)', style='List Bullet')
doc.add_paragraph('Professional PDF report (with charts and tables)', style='List Bullet')
doc.add_paragraph('Complete package (ZIP with all outputs)', style='List Bullet')

doc.add_page_break()

# Technical Features
heading1 = doc.add_heading('5. Technical Features', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

heading2 = doc.add_heading('Pathogens Supported', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

pathogens_data = [
    ('Norovirus', 'Virus', 'P(ill|inf)=0.60'),
    ('Campylobacter jejuni', 'Bacteria', 'P(ill|inf)=0.80'),
    ('Cryptosporidium parvum', 'Protozoa', 'P(ill|inf)=1.00'),
    ('E. coli O157:H7', 'Bacteria', 'P(ill|inf)=0.90'),
    ('Salmonella spp.', 'Bacteria', 'P(ill|inf)=0.75'),
    ('Rotavirus', 'Virus', 'P(ill|inf)=0.65')
]

for pathogen, ptype, param in pathogens_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{pathogen} ({ptype}): {param}')

heading2 = doc.add_heading('Dose-Response Models', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph('Beta-Poisson model (Norovirus, Rotavirus)', style='List Bullet')
doc.add_paragraph('Exponential model (Campylobacter, Cryptosporidium, E. coli)', style='List Bullet')
doc.add_paragraph('Alternative models available for comparison', style='List Bullet')

heading2 = doc.add_heading('Uncertainty Quantification', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph(
    'All results include 5th, 50th (median), and 95th percentiles to represent '
    'optimistic, most-likely, and conservative scenarios. Uncertainties in concentration, '
    'treatment effectiveness, dilution, and exposure volumes are automatically propagated.'
)

doc.add_page_break()

# Key Innovations
heading1 = doc.add_heading('6. Key Innovations (v1.2.0)', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

innovations = [
    ('Exposure-Specific Modeling', 'Route-specific distributions for shellfish meal size, BAF, swim rate, and duration'),
    ('Infection-to-Illness Conversion', 'Separate modeling of infection vs. clinical illness with WHO-based parameters'),
    ('Method Harmonisation Factor', 'Account for differences in measurement methods (water vs. shellfish sampling)'),
    ('Truncated Log-Logistic Distribution', 'Realistic right-skewed distribution for shellfish meal sizes'),
    ('Population Susceptibility', 'Age and population-specific factors affecting illness rates'),
    ('Population Case Estimates', 'Expected annual illness cases for public health planning')
]

for innovation, description in innovations:
    p = doc.add_paragraph(style='List Bullet')
    p_run = p.add_run(f'{innovation}: ')
    p_run.bold = True
    p.add_run(description)

doc.add_page_break()

# Getting Started
heading1 = doc.add_heading('7. Getting Started', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

heading2 = doc.add_heading('System Requirements', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph('Windows, Mac, or Linux', style='List Bullet')
doc.add_paragraph('Python 3.8+', style='List Bullet')
doc.add_paragraph('Web browser (Chrome, Firefox, Safari, Edge)', style='List Bullet')
doc.add_paragraph('~200 MB disk space', style='List Bullet')

heading2 = doc.add_heading('Quick Start', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

quickstart = [
    'Install dependencies: pip install -r requirements.txt',
    'Launch application: streamlit run web_app.py',
    'Select assessment mode from sidebar',
    'Use example data or upload your own',
    'Click "Run Assessment" and wait for results',
    'Download results in preferred format'
]

for i, step in enumerate(quickstart, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

doc.add_page_break()

# Support & References
heading1 = doc.add_heading('8. Support & References', level=1)
heading1.runs[0].font.color.rgb = RGBColor(31, 119, 180)

heading2 = doc.add_heading('Primary References', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

refs = [
    'WHO (2016). Quantitative Microbial Risk Assessment: Application for Water Safety Management',
    'Haas et al. (2014). Quantitative Microbial Risk Assessment, 2nd Edition',
    'U.S. EPA (2019). Method for Assessing Public Health Risk from Waterborne Pathogens',
    'Wood et al. (2023). From_David R Package – Exposure-Specific QMRA Methods'
]

for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

heading2 = doc.add_heading('Documentation', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph('DISTRIBUTION_PARAMETERS_GUIDE.md – Detailed parameter specifications', style='List Bullet')
doc.add_paragraph('BATCH_PROCESSING_INPUTS_GUIDE.md – Input file format reference', style='List Bullet')
doc.add_paragraph('README.md – Complete feature documentation', style='List Bullet')
doc.add_paragraph('Example files in input_data/ directory', style='List Bullet')

heading2 = doc.add_heading('Contact', level=2)
heading2.runs[0].font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph('NIWA Earth Sciences New Zealand')
doc.add_paragraph('For technical support, review example scenarios and documentation files.')

# Save
doc.save('QMRA_Application_Overview.docx')
print("[OK] Professional document created: QMRA_Application_Overview.docx")
print("[OK] Saved to current directory")
