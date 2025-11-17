#!/usr/bin/env python3
"""
Create QMRA User Guide Using NIWA Template Format
=================================================

Fills the NIWA client report template with QMRA user guide content.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
import shutil

def create_niwa_user_guide():
    """Create user guide using NIWA template."""

    # Copy template to new document
    template_path = Path("../../NIWA QMRA report.docx")
    output_path = Path("../../QMRA_User_Guide_NIWA_Format.docx")

    print(f"\nCreating NIWA-formatted QMRA User Guide...")
    print(f"Template: {template_path}")
    print(f"Output: {output_path}\n")

    # Load the template
    doc = Document(str(template_path))

    print(f"Loaded template with {len(doc.paragraphs)} paragraphs")
    print(f"Sections: {len(doc.sections)}")
    print(f"Tables: {len(doc.tables)}\n")

    # Find where to start adding content (after approval pages)
    # Look for "Executive summary" or start of main content
    start_index = 0
    for i, para in enumerate(doc.paragraphs):
        if "executive summary" in para.text.lower() or "introduction" in para.text.lower():
            start_index = i
            break

    print(f"Starting content insertion at paragraph {start_index}\n")

    # Clear existing content after approval pages (keep template structure)
    # We'll add new paragraphs instead of modifying existing ones

    # Add Executive Summary
    print("Adding Executive Summary...")
    para = doc.add_paragraph()
    para.style = 'z_Exec Summary heading'
    para.add_run('Executive Summary')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'The QMRA Batch Processing Web Application is a professional tool for quantitative '
        'microbial risk assessment (QMRA) of waterborne pathogens. It evaluates public health '
        'risks from pathogen exposure through shellfish consumption, recreational water contact, '
        'and other routes using validated dose-response models and Monte Carlo simulation.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'This application implements Production Mode as the default configuration, restricting '
        'pathogen selection to norovirus only with the validated Beta-Binomial dose-response model '
        '(α=0.04, β=0.055 from Teunis et al. 2008). This ensures all risk assessments meet '
        'scientific standards and align with project deliverables.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Key capabilities include:')
    run.bold = True

    bullets = [
        'Five specialized assessment workflows for different risk scenarios',
        'Validated Beta-Binomial dose-response model (exact formulation, not Beta-Poisson approximation)',
        'Monte Carlo simulation with 10,000 iterations for full uncertainty quantification',
        'WHO guideline compliance checking (10⁻⁴ annual infection risk target)',
        'Interactive visualizations and comprehensive statistical summaries',
        'Professional PDF reports and CSV data exports for regulatory submissions'
    ]

    for bullet in bullets:
        para = doc.add_paragraph(bullet, style='Normal')
        para.style = doc.styles['Normal']
        # Set bullet format
        para.paragraph_format.left_indent = Inches(0.5)

    # Add page break before main content
    doc.add_page_break()

    # Add main sections
    print("Adding main content sections...")

    # 1. INTRODUCTION
    heading = doc.add_heading('Introduction to QMRA', level=1)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'Quantitative Microbial Risk Assessment (QMRA) is a scientific framework for estimating '
        'the probability of infection or illness from exposure to pathogenic microorganisms in '
        'water, food, or environmental media. Unlike traditional water quality monitoring that '
        'relies solely on indicator organisms (e.g., E. coli), QMRA directly estimates health '
        'risks by modeling the entire exposure pathway from source to receptor.'
    )

    # 1.1 QMRA Framework
    heading = doc.add_heading('The QMRA Framework', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('The QMRA process integrates four key components:')

    components = [
        ('Hazard Identification', 'Identifying pathogenic microorganisms of concern (e.g., norovirus, Cryptosporidium, Campylobacter)'),
        ('Exposure Assessment', 'Quantifying pathogen concentrations in source water and calculating human exposure doses'),
        ('Dose-Response Modeling', 'Relating pathogen dose to probability of infection using validated mathematical models'),
        ('Risk Characterization', 'Calculating annual infection/illness probabilities and comparing to health-based guidelines')
    ]

    for title, desc in components:
        para = doc.add_paragraph()
        para.style = 'Body Text'
        run = para.add_run(f'{title}: ')
        run.bold = True
        para.add_run(desc)

    # 1.2 Application Purpose
    heading = doc.add_heading('Application Purpose and Scope', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'This web-based application enables water quality professionals to conduct comprehensive '
        'QMRA studies without requiring specialized programming skills or advanced statistical software. '
        'Primary use cases include:'
    )

    use_cases = [
        'Evaluating health risks from wastewater discharge to shellfish harvesting areas',
        'Assessing recreational water quality at swimming beaches and coastal locations',
        'Comparing treatment technology effectiveness for pathogen reduction',
        'Analyzing temporal risk trends from long-term monitoring programs',
        'Supporting regulatory compliance and water safety plan development',
        'Shellfish harvesting area classification (approved, conditionally approved, restricted)'
    ]

    for case in use_cases:
        para = doc.add_paragraph(case, style='Normal')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # 2. GETTING STARTED
    heading = doc.add_heading('Getting Started', level=1)

    heading2 = doc.add_heading('Launching the Application', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'The QMRA application runs as a local web server using Streamlit. '
        'To launch the application:'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Method 1: Using the Batch File (Windows)')
    run.bold = True

    steps = [
        'Navigate to: Batch_Processing_App/app/',
        'Double-click: launch_web_gui.bat',
        'Your default web browser will automatically open to http://localhost:8501'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Normal')
        para.paragraph_format.left_indent = Inches(0.5)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Method 2: Using Command Line')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Open a terminal and run:')

    para = doc.add_paragraph('cd Batch_Processing_App/app', style='Normal')
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    para = doc.add_paragraph('streamlit run web_app.py', style='Normal')
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # 2.2 Production Mode
    heading2 = doc.add_heading('Production Mode vs Research Mode', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('CRITICAL: Production Mode (Default Setting)')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'Production Mode is enabled by default and restricts pathogen selection to norovirus only. '
        'This ensures all risk assessments use the validated Beta-Binomial dose-response model and '
        'produce scientifically defensible results for regulatory submissions.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Production Mode provides:')

    features = [
        'Norovirus-only pathogen selection (highly infectious waterborne virus)',
        'Validated Beta-Binomial dose-response model (α=0.04, β=0.055)',
        'Complete verification against reference Excel calculations (David Wood, NIWA)',
        'Alignment with project contract deliverables and scope',
        'Scientifically defensible results for regulatory compliance'
    ]

    for feature in features:
        para = doc.add_paragraph(feature, style='Normal')
        para.paragraph_format.left_indent = Inches(0.75)

    doc.add_page_break()

    # 3. BETA-BINOMIAL MODEL (CRITICAL SECTION)
    heading = doc.add_heading('Dose-Response Model: Beta-Binomial (Norovirus)', level=1)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('THE CRITICAL DISTINCTION FOR NOROVIRUS')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(200, 0, 0)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'This application uses the EXACT Beta-Binomial dose-response model for norovirus, which is the '
        'mathematically correct formulation. This is NOT the same as the Beta-Poisson approximation.'
    )

    heading2 = doc.add_heading('Beta-Binomial Formula (CORRECT)', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('The exact mathematical formula implemented in this application:')

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run(
        'P(infection) = 1 - exp[ln Γ(β+dose) + ln Γ(α+β) - ln Γ(α+β+dose) - ln Γ(β)]'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 100, 0)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Parameters (Teunis et al. 2008):')
    run.bold = True

    params = [
        'α (alpha) = 0.04',
        'β (beta) = 0.055',
        'Source: Teunis et al. (2008) "Norwalk virus: How infectious is it?" J. Med. Virol. 80:1468-1476'
    ]

    for param in params:
        para = doc.add_paragraph(param, style='Normal')
        para.paragraph_format.left_indent = Inches(0.75)

    heading2 = doc.add_heading('Why Beta-Poisson is INVALID for Norovirus', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('CRITICAL WARNING: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    para.add_run(
        'The Beta-Poisson approximation is mathematically INVALID for norovirus and produces '
        'incorrect risk estimates.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Validity requirement: β >> 1 (beta must be much greater than 1)')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('For norovirus: β = 0.055 << 1 (violates validity requirement)')

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Consequences of using Beta-Poisson (incorrect):')
    run.bold = True

    consequences = [
        'Underestimates risk by 2-4× at low doses (1-100 virions)',
        'Produces scientifically indefensible results',
        'Does not match validated reference calculations',
        'Not acceptable for regulatory submissions'
    ]

    for consequence in consequences:
        para = doc.add_paragraph(consequence, style='Normal')
        para.paragraph_format.left_indent = Inches(0.75)

    heading2 = doc.add_heading('Validation and Verification', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'This application\'s Beta-Binomial implementation has been rigorously validated against '
        'reference Excel spreadsheets (David Wood, NIWA) and produces exact agreement for all test doses.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Example comparison (1 virion dose):')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('Beta-Binomial (CORRECT): 42.1% infection probability')
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('Beta-Poisson (WRONG): 11.1% infection probability')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('Underestimation factor: 3.8× (unacceptable error)')
    run.font.color.rgb = RGBColor(255, 0, 0)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('References:')
    run.bold = True

    references = [
        'McBride, G.B. (2017). Bell Island Wastewater Treatment QMRA. NIWA Client Report 2017350HN, Appendix B',
        'Teunis et al. (2008). Norwalk virus: How infectious is it? J. Med. Virol. 80(8):1468-1476',
        'Haas, C.N. (2002). Conditional dose-response relationships for microorganisms. Risk Analysis 22(3):455-463'
    ]

    for ref in references:
        para = doc.add_paragraph(ref, style='Normal')
        para.paragraph_format.left_indent = Inches(0.75)
        para.paragraph_format.font_size = Pt(10)

    doc.add_page_break()

    # 4. ASSESSMENT WORKFLOWS
    heading = doc.add_heading('Assessment Workflows', level=1)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'The application provides five specialized workflows, each designed for specific '
        'risk assessment scenarios.'
    )

    workflows = [
        ('Batch Scenarios', 'Run multiple pre-configured scenarios simultaneously (15+ scenarios)'),
        ('Spatial Assessment', 'Evaluate risk across multiple sites with different dilution factors'),
        ('Temporal Assessment', 'Analyze risk trends over time from monitoring data'),
        ('Treatment Comparison', 'Compare multiple treatment technologies side-by-side'),
        ('Multi-Pathogen Assessment', 'Evaluate multiple pathogens (Production Mode: norovirus only)')
    ]

    for workflow, desc in workflows:
        para = doc.add_paragraph()
        para.style = 'Body Text'
        run = para.add_run(f'{workflow}: ')
        run.bold = True
        para.add_run(desc)

    # Add detailed workflow instructions for each
    heading2 = doc.add_heading('Spatial Assessment (Detailed Example)', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run('Step-by-step procedure:')

    steps = [
        'Select "Spatial Assessment" from Assessment Mode dropdown in sidebar',
        'Configure parameters: Pathogen (norovirus), Exposure Route (Shellfish Consumption or Primary Contact)',
        'Upload dilution data CSV or check "Use example data" (6 sites with dilution factors 100-100,000)',
        'Set pathogen concentration (e.g., 1×10⁶ organisms/L for raw wastewater)',
        'Set treatment LRV (Log Reduction Value, typically 2-4 for secondary treatment)',
        'Set exposure frequency (e.g., 12 events/year for monthly shellfish consumption)',
        'Click "Run Spatial Assessment"',
        'Review results: site-by-site risk table, WHO compliance status, risk gradient visualization',
        'Download results as CSV and plots as PNG'
    ]

    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='Normal')
        para.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # 5. RESULTS INTERPRETATION
    heading = doc.add_heading('Understanding Results', level=1)

    heading2 = doc.add_heading('Risk Metrics', level=2)

    metrics = [
        ('Mean Dose (organisms per exposure)',
         'Average number of viable pathogen organisms ingested or inhaled during a single exposure event'),
        ('Infection Probability (per exposure)',
         'Probability of infection from a single exposure event, calculated using Beta-Binomial model'),
        ('Annual Infection Risk',
         'Probability of at least one infection over a year: 1 - (1 - P_inf)^n, where n = exposure frequency'),
        ('Illness Probability',
         'Probability of symptomatic illness: P_infection × P(ill|infected). For norovirus: P(ill|infected) = 0.60'),
        ('95% Confidence Interval',
         'Uncertainty range from Monte Carlo simulation (10,000 iterations)')
    ]

    for metric, description in metrics:
        para = doc.add_paragraph()
        para.style = 'Body Text'
        run = para.add_run(f'{metric}: ')
        run.bold = True
        para.add_run(description)

    heading2 = doc.add_heading('WHO Guideline Compliance', level=2)

    para = doc.add_paragraph()
    para.style = 'Body Text'
    para.add_run(
        'The World Health Organization (WHO) recommends a health-based target of 10⁻⁴ (0.0001 or 1 in 10,000) '
        'annual infection risk for drinking water and recreational water quality.'
    )

    para = doc.add_paragraph()
    para.style = 'Body Text'
    run = para.add_run('Compliance evaluation:')
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('PASS (Compliant): Annual infection risk < 10⁻⁴')
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.bold = True

    para = doc.add_paragraph()
    para.style = 'Normal'
    para.paragraph_format.left_indent = Inches(0.75)
    run = para.add_run('FAIL (Exceeds Guideline): Annual infection risk ≥ 10⁻⁴')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True

    doc.add_page_break()

    # 6. REFERENCES
    heading = doc.add_heading('References', level=1)

    references = [
        'Teunis, P.F.M., Moe, C.L., Liu, P., Miller, S.E., Lindesmith, L., Baric, R.S., Le Pendu, J., & Calderon, R.L. (2008). Norwalk virus: How infectious is it? Journal of Medical Virology, 80(8), 1468-1476.',

        'McBride, G.B. (2017). Bell Island Wastewater Treatment Plant Discharge QMRA. NIWA Client Report No: 2017350HN. Prepared for Wellington Water Limited.',

        'World Health Organization (2016). Quantitative Microbial Risk Assessment: Application for Water Safety Management. Geneva: WHO Press.',

        'Haas, C.N., Rose, J.B., & Gerba, C.P. (2014). Quantitative Microbial Risk Assessment (2nd ed.). John Wiley & Sons, Inc.',

        'Haas, C.N. (2002). Conditional dose-response relationships for microorganisms: Development and application. Risk Analysis, 22(3), 455-463.',

        'Van Abel, N., Schoen, M.E., Kissel, J.C., & Meschke, J.S. (2017). Comparison of risk predicted by multiple norovirus dose-response models and implications for quantitative microbial risk assessment. Risk Analysis, 37(2), 245-264.',
    ]

    for i, ref in enumerate(references, 1):
        para = doc.add_paragraph()
        para.style = 'Body Text'
        para.add_run(f'{i}. {ref}')
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.left_indent = Inches(0.5)

    # Save the document
    doc.save(str(output_path))

    print(f"\n{'='*80}")
    print("NIWA-FORMATTED USER GUIDE CREATED SUCCESSFULLY")
    print(f"{'='*80}\n")
    print(f"Output file: {output_path.absolute()}")
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB\n")
    print("Document structure:")
    print("  ✓ NIWA client report template format")
    print("  ✓ Executive summary")
    print("  ✓ Introduction to QMRA")
    print("  ✓ Getting Started (Production Mode emphasis)")
    print("  ✓ Beta-Binomial Model (CRITICAL SECTION)")
    print("  ✓ Assessment Workflows")
    print("  ✓ Results Interpretation")
    print("  ✓ WHO Guideline Compliance")
    print("  ✓ Complete References")
    print(f"\n{'='*80}\n")

    return output_path


if __name__ == '__main__':
    create_niwa_user_guide()
